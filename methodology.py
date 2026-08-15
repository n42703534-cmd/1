"""Generate methodology_v3.docx"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
OUT = r"C:\Users\帅美婷sweet baby\Desktop\network\methodology_v3.docx"
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.54); s.bottom_margin = Cm(2.54)
    s.left_margin = Cm(2.54); s.right_margin = Cm(2.54)
sty = doc.styles['Normal']
sty.font.name = 'Times New Roman'; sty.font.size = Pt(11)
sty.paragraph_format.space_after = Pt(6); sty.paragraph_format.line_spacing = 1.15
for i in range(1,4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Times New Roman'; h.font.color.rgb = None; h.font.bold = True
    h.font.size = {1:Pt(14), 2:Pt(12), 3:Pt(11)}[i]
def P(text, bold=False, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold; r.italic = italic; return p
def eq(text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11)

# ENGLISH
doc.add_heading('2. Methodology', level=1)
doc.add_heading('2.1 Network Model', level=2)
P('The metro station evacuation space is abstracted as a directed graph G = (V, E). Vertices V represent platform waiting zones, concourses, transfer corridors, ticket gates, staircases, escalators, and exits. Edges E represent traversable connections. Each edge e_ij carries length l_ij (m), capacity c_ij (persons/s), and a facility type tag. Edge lengths are measured as walking-path distances from architectural plans; for vertical connections (stairs, escalators), elevation differences are incorporated into the effective path length. Each vertex i has area A_i (m^2), service rate mu_i (persons/s), and instantaneous occupancy N_i(t).')
P('The graph topology is constructed from architectural drawings. Facility capacities are parameterised according to GB/T 33668-2017 (Code for safety evacuation of metro) and supplemented by design-code-based throughput values. The evacuation population is distributed across platform waiting zones, halls, and transfer corridors according to scenario-based demand assumptions, with additional train-load passengers under bidirectional full-load scenarios. Edge costs are recalculated at every simulation time step (Delta t = 0.5 s).')
P('Assumption: The current model assumes a homogeneous pedestrian profile with a uniform free-flow walking speed of v_max = 1.427 m/s. Heterogeneous population characteristics are deferred to future work.', italic=True)

doc.add_heading('2.2 Improved A* Baseline', level=2)
P('The improved A* algorithm by Meng et al. (2022) serves as the baseline:')
eq('f(n) = g(n) + h(n)')
eq('g(n) = alpha * l + beta * l / v(rho)')
eq('h(n) = gamma * d[P]')
P('Here l is edge length, and v(rho) follows the Fruin density-speed relationship (Fruin, 1971), implemented as a clipped piecewise function: v(rho) = 1.427 for rho <= 0.2; v(rho) = max(epsilon, 1.427 - 0.3549*rho) for 0.2 < rho <= 4.0; and edges with rho > 4.0 are treated as jammed. The term d[P] is the precomputed Dijkstra shortest geometric distance from each vertex to each exit. Weights are alpha = 0.15, beta = 0.85, gamma = 0.10. Edges with density exceeding 3.0 p/m^2 are blocked.')
P('Three limitations motivate our work. First, g(n) mixes units of length and time. Second, the binary cutoff at 3.0 p/m^2 creates a discontinuity. Third, the formulation omits bottleneck service rates.')

doc.add_heading('2.3 Proposed Algorithm: AdaptiveQueueAwareAStar', level=2)
P('We propose AdaptiveQueueAwareAStar, which reformulates the A* cost model around two principles: (i) all path costs are expressed in a single physical unit (seconds); (ii) bottleneck queuing delay is explicitly modelled using vertex-level service rates. The algorithm retains the A* framework f(n) = g(n) + h(n) and modifies only the edge cost and heuristic.')

doc.add_heading('2.3.1 Edge Cost Model', level=3)
P('The edge cost for traversing (u, v) is composed of two additive components in seconds:')
eq('g_edge = t_travel + t_wait')
P('Travel time captures the physical time to walk the edge under instantaneous density:')
eq('v_actual = min(v(rho), v_cap)')
eq('t_travel = l / v_actual')
P('where v_cap = 1.427 m/s for flat passageways, 0.75 m/s for stairs, and 0.50 m/s for escalators. Unlike the baseline, no premature binary cutoff at 3.0 p/m^2 is applied; instead, traversal cost increases continuously with density until the jam-density limit is reached. To prevent mathematical singularities at jam density, the edge cost tends to infinity as v(rho) approaches zero, effectively pruning blocked paths from the search space.')
P('Waiting time is activated exclusively at bottleneck vertices (gates, stairs, escalators). We adopt a deterministic point-queue approximation for bottleneck delay, following classical point-queue modelling in dynamic network loading (Jin, 2014). This is well-suited to pedestrian evacuation for three reasons. First, unlike the M/M/1 stochastic model, the point-queue model does not require Poisson arrivals or exponential service times, which are violated in coordinated mass evacuation. Second, unlike steady-state formulations, it captures transient queue build-up and dissipation. Third, it requires only two observable quantities (current queue length and service rate), both available from the network model without additional calibration. Li et al. (2022) provide a related application of queue-aware dynamic guidance for evacuation flow equilibrium.')
P('The derivation follows from conservation of pedestrians at a bottleneck. Let Q(t) be the queue length, mu the fixed service rate, and lambda(t) the arrival rate. The queue dynamics satisfy dQ/dt = lambda(t) - mu. In discrete time:')
eq('Q(t + Delta t) = max(0, Q(t) + (lambda(t) - mu) * Delta t)')
P('The expected waiting delay before service for a newly assigned pedestrian is approximated by:')
eq('t_wait = N(t) / mu')
P('where N(t) = Q(t) + I_reserved(t) is the effective queue length, combining physically present passengers with those already assigned to this bottleneck by the guidance system. In our implementation, I_reserved(t) is dynamically updated as passengers are routed in each simulation step, closing the prediction loop. This expression is consistent with Little\'s Law under saturated bottleneck service (L = N, lambda = mu, W = t_wait), rather than a steady-state queueing assumption. The waiting time t_wait represents the queueing delay before entering service; the physical traversal time of the facility itself is included in t_travel. The queueing delay is added only once when entering a bottleneck vertex, rather than on every adjacent edge, to avoid double-counting. The service rate mu is constrained by the fixed physical capacity of the facility. Embedding the point-queue model into the A* cost function is the central innovation of our approach.')

doc.add_heading('2.3.2 Heuristic Function', level=3)
P('To obtain the shortest path under the frozen snapshot of current edge weights, the heuristic must be admissible with respect to the instantaneous edge-cost snapshot (Hart et al., 1968). We adopt the free-flow lower bound:')
eq('h(n) = d_min(n, exit) / v_max')
P('where d_min is the precomputed Dijkstra shortest geometric distance and v_max = 1.427 m/s is the Fruin free-flow walking speed. The heuristic is computed once at initialisation and shares the same physical unit (seconds) as g(n).')

doc.add_heading('2.3.3 Ablation Variant', level=3)
P('To isolate the contribution of bottleneck queuing, we construct one ablation variant:')
P('NoWaitingTime (Density-Only). The waiting time is disabled (t_wait = 0), reducing the cost to g_edge = l / v_actual. Continuous Fruin speed and time-unified formulation are retained. This yields a clean three-way decomposition: AdaptiveQueueAwareAStar vs NoWaitingTime isolates the queueing term; NoWaitingTime vs ImprovedA* isolates the combined effect of continuous density modelling and time-unified costs; AdaptiveQueueAwareAStar vs ImprovedA* quantifies the net improvement of the complete algorithm.')
P('Forward-extrapolation of queue states and upstream-demand estimation were experimentally evaluated. Both variants degraded performance relative to the snapshot formulation reported here, a finding discussed as a limitation of decentralised A* search in Section 5.')

# CHINESE
doc.add_page_break()
doc.add_heading('2. 模型与算法', level=1)
doc.add_heading('2.1 网络模型', level=2)
P('将地铁车站疏散空间抽象为有向图 G = (V, E)。节点 V 表示站台等候区、站厅、换乘通道、闸机、楼梯、扶梯和出口。边 E 表示设施间可通行连接。每条边 e_ij 具有长度 l_ij (m)、通行能力 c_ij (人/s) 和设施类型标签。边长采用建筑图纸测量的行走路径距离；竖向连接 (楼梯、扶梯) 将高差纳入有效路径长度。每个节点 i 具有面积 A_i (m^2)、服务率 mu_i (人/s) 及当前人数 N_i(t)。')
P('图拓扑基于建筑图纸构建。设施通行能力按 GB/T 33668-2017 (地铁安全疏散规范) 参数化并辅以设计规范中的通行量值。疏散总人数按场景需求假设分配至各源节点，双向满载场景下叠加列车乘客。边成本在每个仿真步 (Delta t = 0.5 s) 重新计算。')
P('假设：当前模型假设同质行人群体，统一自由流步行速度 v_max = 1.427 m/s。异质性人群特征留待未来研究。', italic=True)

doc.add_heading('2.2 改进 A* 基线', level=2)
P('Meng 等 (2022) 针对邮轮疏散提出的改进 A* 算法作为本文基线：')
eq('f(n) = g(n) + h(n)')
eq('g(n) = alpha * l + beta * l / v(rho)')
eq('h(n) = gamma * d[P]')
P('其中 l 为边长，v(rho) 为 Fruin 密度-速度关系 (Fruin, 1971)，采用截断分段函数：v(rho) = 1.427 (rho <= 0.2)；v(rho) = max(epsilon, 1.427 - 0.3549*rho) (0.2 < rho <= 4.0)；rho > 4.0 时视为堵塞。d[P] 为 Dijkstra 预计算的节点到出口最短几何距离。权重 alpha = 0.15, beta = 0.85, gamma = 0.10。密度超过 3.0 p/m^2 的边被视为不可通行。')
P('该基线有三点局限：(1) g(n) 混合长度和时间两种量纲；(2) 3.0 p/m^2 处二元截断产生不连续；(3) 不包含瓶颈节点服务率。')

doc.add_heading('2.3 本文算法：AdaptiveQueueAwareAStar', level=2)
P('本文提出 AdaptiveQueueAwareAStar 算法，围绕两个原则重构 A* 成本模型：(i) 所有路径成本统一为秒；(ii) 通过节点级服务率显式建模瓶颈排队延迟。算法保留 A* 框架 f(n) = g(n) + h(n)，仅修改边代价和启发式的构成。')

doc.add_heading('2.3.1 边代价模型', level=3)
P('边 (u, v) 的通行代价由两个以秒为单位的可加成分构成：')
eq('g_edge = t_travel + t_wait')
P('旅行时间捕获在当前密度下步行边长的物理时间：')
eq('v_actual = min(v(rho), v_cap)')
eq('t_travel = l / v_actual')
P('其中 v_cap 取值：平地 1.427 m/s，楼梯 0.75 m/s，扶梯 0.50 m/s。与基线不同，不在 3.0 p/m^2 处设置提前二元封闭阈值，而是使通行代价随密度连续上升，直至达到堵塞密度极限。为防范堵塞密度处的数学奇点，边代价趋向无穷大。')
P('等待时间仅在瓶颈节点 (闸机、楼梯、扶梯) 启用。采用确定性点排队近似进行瓶颈延迟建模，遵循动态网络加载中的经典点排队模型 (Jin, 2014)。适用于行人疏散有三个原因：第一，不同于 M/M/1 随机模型，点排队不要求泊松到达或指数服务时间假设，这些假设在协调性大规模疏散中不成立；第二，不同于稳态公式，点排队能捕捉瞬态排队累积与消散过程；第三，仅需两个可观测量 (当前队列长度和服务率)，均可从网络模型直接获取。Li 等 (2022) 提供了排队感知动态引导在疏散流量均衡中的相关应用。')
P('推导源于瓶颈节点的行人守恒原理。令 Q(t) 为队列长度，mu 为固定服务率，lambda(t) 为到达率。队列动力学满足 dQ/dt = lambda(t) - mu。离散时间下：')
eq('Q(t + Delta t) = max(0, Q(t) + (lambda(t) - mu) * Delta t)')
P('新到达行人进入服务前的预期等待延迟近似为：')
eq('t_wait = N(t) / mu')
P('其中 N(t) = Q(t) + I_reserved(t) 为有效队列长度，融合物理在场乘客与已被引导系统分配至该瓶颈的预留流入量。在本文实现中，I_reserved(t) 在每个仿真步随乘客路径分配动态更新，形成闭环预测。该式与瓶颈饱和服务下 Little 定律一致 (L = N, lambda = mu, W = t_wait)，而非稳态排队假设。等待时间 t_wait 表示进入服务前的排队延迟；设施本身的物理通过时间包含在 t_travel 中。排队等待时间仅在行人进入瓶颈节点时计入一次，避免在瓶颈相邻边上重复累计。服务率 mu 受限于设施的固定物理容量。将点排队模型嵌入 A* 成本函数是本文算法的核心创新。')

doc.add_heading('2.3.2 启发式函数', level=3)
P('为在每一决策步获得当前冻结边权快照下的最短路径，启发式须相对于瞬时边权快照可采纳 (Hart et al., 1968)。采用自由流下界：')
eq('h(n) = d_min(n, exit) / v_max')
P('其中 d_min 为预计算几何最短距离，v_max = 1.427 m/s。初始化时计算一次，与 g(n) 同量纲。')

doc.add_heading('2.3.3 消融变体', level=3)
P('为分离瓶颈排队模型的贡献，构造一个消融变体：')
P('NoWaitingTime (Density-Only)。设 t_wait = 0，边代价退化为 g_edge = l / v_actual。保留连续 Fruin 速度和时间统一公式。形成三项清晰分解：AdaptiveQueueAwareAStar vs NoWaitingTime 分离排队项的贡献；NoWaitingTime vs ImprovedA* 分离连续密度建模和时间统一成本的综合效果；AdaptiveQueueAwareAStar vs ImprovedA* 量化完整算法的净改善。')
P('预测外推和上游需求估计已进行实验评估，均劣于当前快照版本，视为分布式 A* 搜索的固有局限，在第五章讨论。')

# REFERENCES
doc.add_page_break()
doc.add_heading('References', level=1)
for r in [
    'Bai, J., Lv, X., Nie, L., & Fang, M. (2025). Evacuation route determination in indoor architectural environments based on dynamic fire risk assessment. Buildings, 15(10), 1715.',
    'Fruin, J.J. (1971). Pedestrian Planning and Design. Metropolitan Association of Urban Designers and Environmental Planners, New York.',
    'Hart, P.E., Nilsson, N.J., & Raphael, B. (1968). A formal basis for the heuristic determination of minimum cost paths. IEEE Transactions on Systems Science and Cybernetics, 4(2), 100-107.',
    'Jin, W.-L. (2014). Point queue models: a unified approach. Transportation Research Part B, 77, 1-17.',
    'Li, M., Xu, C., Xu, Y., Ma, L., & Wei, Y. (2022). Dynamic sign guidance optimization for crowd evacuation considering flow equilibrium. Journal of Advanced Transportation, 2022, 2555350.',
    'Meng, D., Hu, Z., & Zhang, H. (2022). Simulation of multi-layer ship evacuation system based on improved A* algorithm. Journal of System Simulation, 34(6), 1375-1382. (in Chinese)',
    'Zuo, S., Mao, Z., Fan, C., Chen, X., Gong, M., Ren, J., Fan, X., & Guo, Y. (2024). Dynamic planning of crowd evacuation path for metro station based on Dynamic Avoid Smoke A-Star algorithm. Tunnelling and Underground Space Technology, 154, 106145.',
    'GB/T 33668-2017. (2017). Code for safety evacuation of metro. Standards Press of China, Beijing.',
]:
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.first_line_indent = Cm(-1.27)
    p.add_run(r).font.size = Pt(10)

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f'Saved: {OUT}')
