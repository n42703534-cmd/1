from __future__ import annotations

import math
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUTDIR = ROOT / "outputs"
ASSETDIR = OUTDIR / "algorithm_method_assets"
OUT = OUTDIR / "算法方法章节_TUST风格_20260706.docx"
FLOW = ASSETDIR / "AA算法动态更新流程.png"

# Design preset: narrative_proposal.
# Named override "academic_method_chapter": Times New Roman + SimSun, black
# hierarchy, compact academic spacing. Page/table geometry remains the exact
# Letter/9360-DXA preset geometry.
INK = "1A1A1A"
MUTED = "626262"
ACCENT = "1F4E79"
PALE = "EAF1F7"
PALE2 = "F4F6F9"
LINE = "9CA9B5"
WHITE = "FFFFFF"


def font_path(*names: str) -> str:
    base = Path(r"C:\Windows\Fonts")
    for name in names:
        p = base / name
        if p.exists():
            return str(p)
    return str(base / "arial.ttf")


def make_flowchart(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    W, H = 1800, 1700
    img = Image.new("RGB", (W, H), "white")
    d = ImageDraw.Draw(img)
    f_title = ImageFont.truetype(font_path("msyhbd.ttc", "simhei.ttf"), 56)
    f_box = ImageFont.truetype(font_path("msyh.ttc", "simsun.ttc"), 34)
    f_small = ImageFont.truetype(font_path("msyh.ttc", "simsun.ttc"), 29)

    def rounded(x0, y0, x1, y1, fill, outline=ACCENT, width=4, radius=24):
        d.rounded_rectangle((x0, y0, x1, y1), radius=radius, fill=fill, outline="#" + outline, width=width)

    def center_text(box, text, font=f_box, fill="#1A1A1A", spacing=8):
        x0, y0, x1, y1 = box
        lines = text.split("\n")
        heights = [d.textbbox((0, 0), line, font=font)[3] for line in lines]
        total = sum(heights) + spacing * (len(lines) - 1)
        y = (y0 + y1 - total) / 2
        for line, h in zip(lines, heights):
            bb = d.textbbox((0, 0), line, font=font)
            x = (x0 + x1 - (bb[2] - bb[0])) / 2
            d.text((x, y), line, font=font, fill=fill)
            y += h + spacing

    def arrow(x, y0, y1):
        d.line((x, y0, x, y1 - 18), fill="#607D93", width=5)
        d.polygon([(x, y1), (x - 13, y1 - 22), (x + 13, y1 - 22)], fill="#607D93")

    title = "AdaptiveQueueAwareAStar 的单步闭环更新"
    bb = d.textbbox((0, 0), title, font=f_title)
    d.text(((W - (bb[2] - bb[0])) / 2, 30), title, font=f_title, fill="#173B57")

    cx = W // 2
    boxes = [
        (260, 125, 1540, 230, "输入网络 G、当前状态 S(t) 与出口集合 D", PALE),
        (260, 300, 1540, 425, "处理在途到达\n更新边密度与服务节点到达率 EMA", PALE2),
        (260, 495, 1540, 630, "计算每条边的时间代价\n通行时间 + 到达时刻预测等待时间", PALE),
        (260, 700, 1540, 850, "对每个有客流节点、每个可达出口执行 A*\n比较候选完整路径的瞬时代价", PALE2),
        (260, 920, 1540, 1055, "应用路径切换惯性\n保持、常规切换、强制切换或退化触发", PALE),
        (260, 1125, 1540, 1275, "顺序预留引导流入并刷新受影响边权\n必要时启用 L2/L18 局部并行设施", PALE2),
        (260, 1345, 1540, 1485, "容量约束、整数化与下游储存约束\n生成移动量并推进 Δt = 0.5 s", PALE),
    ]
    for i, (x0, y0, x1, y1, txt, fill) in enumerate(boxes):
        rounded(x0, y0, x1, y1, "#" + fill)
        center_text((x0, y0, x1, y1), txt, f_box if i not in (5, 6) else f_small)
        if i < len(boxes) - 1:
            arrow(cx, y1 + 5, boxes[i + 1][1] - 5)

    # Feedback loop.
    xloop = 1660
    d.line((1540, 1415, xloop, 1415), fill="#607D93", width=5)
    d.line((xloop, 1415, xloop, 362), fill="#607D93", width=5)
    d.line((xloop, 362, 1545, 362), fill="#607D93", width=5)
    d.polygon([(1540, 362), (1563, 349), (1563, 375)], fill="#607D93")
    d.text((1680, 885), "未完成疏散", font=f_small, fill="#607D93", anchor="mm", stroke_width=0)
    img.save(path, dpi=(220, 220))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6) -> None:
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_table_geometry(table, widths_dxa: list[int], indent=120) -> None:
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), str(indent))
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for cell, w in zip(row.cells, widths_dxa):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(w))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_run_font(run, latin="Times New Roman", east="宋体", size=10.5,
                 bold=None, italic=None, color=INK) -> None:
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para(p, before=0, after=5, line=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_line=0.28, keep=False) -> None:
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.alignment = align
    pf.first_line_indent = Inches(first_line) if first_line else None
    pf.keep_with_next = keep


def add_body(doc, text: str, *, first=True, after=5, italic=False) -> None:
    p = doc.add_paragraph(style="Normal")
    set_para(p, after=after, first_line=0.28 if first else 0)
    r = p.add_run(text)
    set_run_font(r, italic=italic)


def add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)


def add_equation(doc, expr: str, num: int) -> None:
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [8500, 860], indent=120)
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tbl_pr.append(borders)
    for cell in table.rows[0].cells:
        set_cell_margins(cell, top=20, start=20, bottom=20, end=20)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = table.cell(0, 0).paragraphs[0]
    p.style = doc.styles["Equation"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(expr)
    set_run_font(r, latin="Cambria Math", east="Cambria Math", size=9.8)
    p = table.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    n = p.add_run(f"({num})")
    set_run_font(n, size=9.8)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)
    spacer.paragraph_format.line_spacing = 0.5


def add_caption(doc, text: str) -> None:
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=True)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int],
              font_size=9.2, first_col_left=True):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j, text in enumerate(headers):
        cell = hdr.cells[j]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, PALE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, size=font_size, bold=True, color=ACCENT)
    for i, row in enumerate(rows):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            cell = cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if i % 2 == 1:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if (first_col_left or j > 0) else WD_ALIGN_PARAGRAPH.CENTER
            if j == 0 and not first_col_left:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            r = p.add_run(text)
            set_run_font(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = paragraph.add_run("Page ")
    set_run_font(r, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, before, after in ((1, 16, 16, 8), (2, 13, 12, 6), (3, 11.5, 8, 4)):
        s = styles[f"Heading {level}"]
        s.font.name = "Times New Roman"
        s._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
        s._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(INK)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
        s.paragraph_format.keep_together = True
    cap = styles["Caption"]
    cap.font.name = "Times New Roman"
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    cap.font.size = Pt(9.5)
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(4)
    cap.paragraph_format.keep_with_next = True
    eq = styles.add_style("Equation", 1)
    eq.font.name = "Cambria Math"
    eq.font.size = Pt(10.5)
    eq.paragraph_format.space_before = Pt(3)
    eq.paragraph_format.space_after = Pt(5)
    eq.paragraph_format.line_spacing = 1.0
    eq.paragraph_format.keep_together = True


def build() -> Path:
    OUTDIR.mkdir(exist_ok=True)
    make_flowchart(FLOW)
    doc = Document()
    configure_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    # Running furniture: quiet technical-document treatment.
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("模型与算法  |  AdaptiveQueueAwareAStar")
    set_run_font(hr, size=8.5, color=MUTED)
    add_page_number(sec.footer.paragraphs[0])

    # First-page masthead without decorative border.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("2  模型与算法")
    set_run_font(r, east="黑体", size=22, bold=True, color="173B57")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run("Mesoscopic network model and AdaptiveQueueAwareAStar")
    set_run_font(r, size=11.5, italic=True, color=MUTED)

    add_body(doc, "本章面向复杂换乘地铁站的高负荷疏散，构建介观有向网络模型，并提出自适应排队感知 A* 算法（AdaptiveQueueAwareAStar，以下简称 AA）。算法的目标不是复现单个行人的连续轨迹，而是在 0.5 s 的离散时间步内，根据边上通行状态、瓶颈服务能力和预计到达时的队列长度，为各有客流节点确定下一跳及完整出口路径。方法章节按“网络抽象—基准算法—代价函数—动态更新—执行约束—伪代码”的顺序展开，这与 Tunnelling and Underground Space Technology 中动态疏散路径论文常用的写法一致（Zuo et al., 2024）。")

    add_heading(doc, "2.1 问题定义与建模边界", 2)
    add_body(doc, "将车站疏散空间抽象为有向图 G = (V, E)。节点集合 V 包含站台等候区、站厅或换乘连接点、闸机、楼梯、扶梯和出口；有向边集合 E 表示可通行连接。对边 (u, v) 记录长度 l_uv、有效宽度 w_uv、通行能力 c_uv 与设施类型；对节点 v 记录面积 A_v、当前人数 N_v(t) 与服务能力 μ_v。闸机、楼梯和扶梯被定义为容量服务节点。出口节点只接收流量，不再向网络内回流。")
    add_body(doc, "该模型属于介观模型：状态变量是节点或边上的整数客流包，决策对象是同一节点上的客流，而不是具有独立速度、反应时间和避碰行为的个体代理。因此，AA 与 Pathfinder 的比较应落在总疏散过程、出口和设施流量、排队与拥挤暴露等聚合指标上，不要求个体轨迹或逐人疏散时刻完全一致。当前龙阳路站实例包含 555 个节点、1381 条有向边和 16 个出口；高负荷双向满载场景共 17,905 人。网络规模只说明本研究的实现对象，不构成算法适用规模的理论上限。")
    add_body(doc, "模型采用以下假设：(1) 同一节点内的乘客具有相同的可达拓扑和引导信息；(2) 所有乘客遵从系统给出的节点级引导；(3) 设施几何与标称通行能力在一次仿真中保持不变；(4) 高负荷下允许排队向上游回堵，但不解析个体间的横向避碰；(5) 一般节点采用单一下一跳，只有经物理拓扑确认的 L2 竖向设施对和 L18 闸机对允许局部分流。上述边界决定了本算法是疏散流引导与负荷分配方法，而不是微观行人行为模型。")

    add_caption(doc, "表 1  主要符号及含义")
    add_table(doc,
              ["符号", "含义", "单位"],
              [
                  ["G = (V, E)", "车站介观有向网络", "—"],
                  ["l_uv, w_uv", "边 (u, v) 的长度与有效宽度", "m"],
                  ["N_v(t), A_v", "节点 v 在 t 时刻的人数与有效面积", "人；m²"],
                  ["R_v(t)", "同一时间步内已引导、尚未实际到达 v 的预留流入", "人"],
                  ["ρ_uv(t)", "边或其下游空间的有效密度", "人/m²"],
                  ["μ_v", "容量服务节点 v 的服务率", "人/s"],
                  ["λ̂_v(t)", "服务节点 v 的指数平滑到达率", "人/s"],
                  ["Q̂_v(t+τ)", "预计到达时刻的服务队列", "人"],
                  ["c_uv(t)", "AA 的边通行广义代价", "s"],
                  ["d(n, e)", "节点 n 至出口 e 的最短几何距离", "m"],
              ], [1300, 6660, 1400], font_size=9.0, first_col_left=False)

    add_heading(doc, "2.2 介观状态更新与通行约束", 2)
    add_heading(doc, "2.2.1 有效密度与通行速度", 3)
    add_body(doc, "路径搜索使用边上实时密度与下游占用共同构成的保守密度。边的代表面积取长度与宽度之积，并扣除显式障碍面积；当边上在途密度高于下游节点投影密度时，采用两者的较大值：")
    add_equation(doc, "ρᵤᵥ(t) = max{[Nᵥ(t)+Rᵥ(t)] / max[lᵤᵥwᵤᵥ−Aobs,ᵤᵥ, 0.1], ρlink,ᵤᵥ(t)}", 1)
    add_body(doc, "其中，ρ_uv^link(t) 为正在边 (u, v) 上运动的客流除以边有效面积。R_v(t) 只用于同一决策步内抑制后续客流继续选择已经承压的设施，并在该时间步的路径分配结束后清零，因此不会被当作实际在场人数重复累积。")
    add_body(doc, "为保持与 Improved A* 基线的一致可比性，路径搜索阶段沿用 Meng et al. (2022) 采用的 Fruin 型分段密度—速度关系：")
    add_equation(doc, "vF(ρ) = {1.427, ρ≤0.2;  max(1.427−0.3549ρ, 0), 0.2<ρ≤4.0;  0, ρ>4.0}", 2)
    add_equation(doc, "vᵤᵥ(t) = min[vF(ρᵤᵥ(t)), vcap,ᵤᵥ]", 3)
    add_body(doc, "设施速度上限 v_cap,uv 分别取平地 1.427 m/s、楼梯 0.75 m/s 和扶梯 0.50 m/s。式 (2) 中的 4.0 人/m²是路径可行性阈值；高负荷回堵模块另用 5.4 人/m²估计下游储存上限。两者服务于不同机制，不应混为同一标定参数。")

    add_heading(doc, "2.2.2 容量、整数客流与回堵", 3)
    add_body(doc, "每个时间步允许从 u 进入 v 的客流首先受边的有效服务能力约束。通道与楼梯在低于临界密度时使用标称能力与文献基本图能力的较小值；超过临界密度后，按 Weidmann 型速度—密度关系降低接收能力。闸机和扶梯保留设备标称能力。为避免 c_uvΔt 小于 1 时客流永久无法移动，代码累积小数容量信用，再将可移动量整数化。若同一节点存在多个已提出的物理并行边，则未被主边接收的剩余客流按各边剩余能力重新分配。")
    add_body(doc, "高负荷场景还施加下游储存约束。非出口节点的储存上限定义为有效面积容量与最小流量缓冲容量中的较大值：")
    add_equation(doc, "Kᵥ = max[Aeff,ᵥ ρjam,store, μout,ᵥ Tbuf, 1]", 4)
    add_equation(doc, "Sᵥ(t) = max{⌊Kᵥ−Nᵥ(t)−Itransit,ᵥ(t)⌋, 0}", 5)
    add_body(doc, "其中 ρ_jam^store = 5.4 人/m²，T_buf = 18 s，I_v^transit(t) 为已经从上游放行、正在前往 v 的在途人数。所有同时汇入 v 的边共享 S_v(t)，从而在出发时预留下游空间，使拥堵向上游传播，而不是让多个入口分别满额放行后在下游凭空叠加。")

    add_heading(doc, "2.3 Improved A* 基准算法", 2)
    add_body(doc, "基准算法根据蒙盾等（2022）的改进 A* 思路进行车站网络适配。其评价函数仍写为 f(n)=g(n)+h(n)，边代价同时考虑长度和密度折减后的旅行时间：")
    add_equation(doc, "gbase,ᵤᵥ(t) = αlᵤᵥ + βlᵤᵥ/vᵤᵥ(t)", 6)
    add_equation(doc, "hbase(n,e) = γd(n,e)", 7)
    add_body(doc, "取 α=0.15、β=0.85、γ=0.10。当节点空间密度或候选边有效密度超过 3.0 人/m²时，将其从本轮搜索空间中移除；只有当拥堵节点集合发生变化或原路径被阻断时才重规划，否则沿用已缓存的下一跳。该方法能够避开当前高密度区域，但没有显式描述闸机、楼梯或扶梯的排队服务过程。同时，式 (6) 的长度项与时间项量纲不同，因此其数值只具有加权排序意义。本文保留这一形式作为文献基准，不把它作为 AA 的物理时间代价。")

    add_heading(doc, "2.4 AdaptiveQueueAwareAStar", 2)
    add_heading(doc, "2.4.1 统一为时间的边代价", 3)
    add_body(doc, "AA 保留 A* 的搜索框架，但将每条边的代价统一为秒。对普通通道，代价仅为密度条件下的通行时间；当边的终点 v 是闸机、楼梯或扶梯时，再加入预计到达时刻的排队等待时间：")
    add_equation(doc, "τwalk,ᵤᵥ(t) = lᵤᵥ/vᵤᵥ(t)", 8)
    add_equation(doc, "cᵤᵥ(t) = τwalk,ᵤᵥ(t) + 𝟙(v∈Vs)τwait,ᵥ[t+τwalk,ᵤᵥ(t)]", 9)
    add_body(doc, "V_s 为容量服务节点集合。等待项只在进入服务节点时计入一次，设施本身的行走或乘行时间已经包含在 τ_uv^walk 中，因而不会在相邻边上重复计算。与基准算法在 3.0 人/m²处直接封闭不同，AA 在 0.2–4.0 人/m²范围内连续增加旅行时间；只有速度降至接近零时才令边代价为无穷大。")

    add_heading(doc, "2.4.2 到达率估计与到达时刻队列预测", 3)
    add_body(doc, "仅使用当前节点人数会低估“尚在路上但即将到达”的压力。为此，AA 在每个时间步统计实际完成入边行程并到达服务节点的人数 a_v^k，以指数移动平均估计到达率：")
    add_equation(doc, "λ̂ᵥᵏ = η(aᵥᵏ/Δt) + (1−η)λ̂ᵥᵏ⁻¹", 10)
    add_body(doc, "其中 η=0.30，Δt=0.5 s。随后采用确定性流体队列守恒，在从 u 到 v 的预计行程时间内外推队列（Jin, 2015）：")
    add_equation(doc, "Qnow,ᵥ(t) = Nᵥ(t)+Rᵥ(t)", 11)
    add_equation(doc, "Q̂ᵥ[t+τwalk,ᵤᵥ(t)] = max{Qnow,ᵥ(t)+[λ̂ᵥ(t)−μᵥ]τwalk,ᵤᵥ(t), 0}", 12)
    add_equation(doc, "τwait,ᵥ[t+τwalk,ᵤᵥ(t)] = Q̂ᵥ[t+τwalk,ᵤᵥ(t)]/μᵥ", 13)
    add_body(doc, "式 (12) 描述的是有限预测时域内的队列净增长，而不是 M/M/1 稳态随机排队。其所需变量均来自介观仿真：当前节点人数、当步预留流入、观察到的到达率和设施服务率。若到达率低于服务率，预计队列会在乘客到达前部分消散；若到达率高于服务率，等待代价相应增加。点排队模型可用于描述服务需求超过供给时的排队形成与消散（Jin, 2015），而显式考虑排队等待也与动态疏散流均衡研究的建模方向一致（Li et al., 2022）。")

    add_heading(doc, "2.4.3 启发式函数与多出口搜索", 3)
    add_body(doc, "对每个候选出口 e，AA 使用预先计算的最短几何距离构造自由流时间下界：")
    add_equation(doc, "h(n,e) = d(n,e)/vfree", 14)
    add_equation(doc, "f(n,e;t) = ∑(u,v)∈P(s,n)cᵤᵥ(t)+h(n,e)", 15)
    add_body(doc, "其中 v_free=1.427 m/s。由于任意实际边速度不超过 v_free，且等待时间非负，式 (14) 不高估冻结状态快照下的剩余代价，因而可作为可采纳启发式。算法分别对所有可达出口执行 A*，得到候选路径集合，再以完整路径的瞬时边代价之和排序并选择最小者。需要强调：当前实现没有额外启用一个独立的“出口 Gini 惩罚项”；出口负荷均衡来自服务队列、预留流入和下游容量对路径代价的共同反馈。这样写与代码一致，也避免把结果指标误写成优化目标。")

    add_heading(doc, "2.4.4 单步顺序预留", 3)
    add_body(doc, "如果同一时间步内所有有客流节点都根据同一份边权快照独立决策，多个来源可能同时选择边际代价最低的瓶颈。AA 因此采用顺序预留：某一来源确定路径后，将本步拟发送量记入路径第一跳的下游节点，并记入后续所有容量服务节点的 R_v(t)；随后立即重算这些节点所有入边的代价，再处理下一个来源。该预留只改变本步尚未决策客流看到的预测状态，不提前移动实际客流，也不改变设施物理能力。所有来源完成决策后，R_v(t) 被清空，真实移动仍由容量与在途传播模块执行。")

    add_heading(doc, "2.4.5 路径切换惯性", 3)
    add_body(doc, "动态最短路在两个代价接近的路径之间可能频繁振荡。AA 为每个有客流节点保存上次选择的完整路径、下一跳、切换时刻及选择时成本，并采用以下规则：新路径至少比原路径低 3% 才进行常规切换；持有时间不足 2.0 s 时原则上保持原路径；若新路径低 20% 以上，可忽略最小持有时间强制切换；若原路径成本相对上次选择时增长超过 50%，且最佳候选至少再低 2%，则触发退化切换。若最佳路径的下一跳未变化，仅更新当前成本而不记录一次新的切换。该机制改变的是引导稳定性，不改变边的物理通行能力。")

    add_heading(doc, "2.4.6 物理并行设施的局部分流", 3)
    add_body(doc, "单一下一跳是 AA 的基本决策形式，但聚合站台节点会掩盖相邻设施的物理并行关系：当两部设施连接同一上游空间和同一下游走向时，微小的代价差会使全部客流锁定到其中一部。为纠正这种介观聚合偏差，代码只对已核验的设施对开放局部分流：L2 的 Stair_L2_1/Stair_L2_2、Stair_L2_3/Escalator_L2_up1，以及 L18 的 E1/E2、S1/S2 闸机对。先按 AA 选定完整下游路线；若主设施本步能力不足以接收全部等待客流，再把剩余量送入保持相同下游节点的配对设施。该规则不改变出口选择，也不是任意多路径分流。")

    add_caption(doc, "图 1  AA 算法与介观客流传播的闭环更新流程")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(FLOW), width=Inches(5.5))

    add_heading(doc, "2.5 算法执行过程", 2)
    add_body(doc, "算法 1 给出 AA 在一个完整疏散过程中的执行顺序。其关键不是一次性计算静态路线，而是在每个仿真步重新形成状态快照、更新服务到达率与密度、对活动节点逐一搜索并预留，然后将决策结果交给容量和回堵模块执行。")
    add_caption(doc, "算法 1  AdaptiveQueueAwareAStar 与介观仿真的耦合过程")
    pseudo = [
        ["1", "输入 G=(V,E)、出口集合 D、初始客流、Δt；预计算各节点至各出口的几何最短距离 d(n,e)。"],
        ["2", "初始化 t=0、在途队列、服务节点到达率 EMA、路径状态和容量信用。"],
        ["3", "while 网络内或在途队列中仍有乘客 do"],
        ["4", "  处理 t 时刻到达：更新节点人数、出口累计量和服务节点实际到达人数 a_v^k。"],
        ["5", "  按式 (10) 更新 λ̂_v(t)，刷新边上在途人数与实时密度。"],
        ["6", "  清空本步预留 R_v(t)，按式 (1)–(3)、(8)–(13)计算所有边 c_uv(t)。"],
        ["7", "  识别当前有客流且非出口的活动节点集合 V_a(t)。"],
        ["8", "  for u ∈ V_a(t) do"],
        ["9", "    for e ∈ D 且 e 对 u 可达 do"],
        ["10", "      以 c_uv(t) 为边权、式 (14) 为启发式执行 A*，得到候选路径 P_ue。"],
        ["11", "    end for"],
        ["12", "    选择瞬时完整路径代价最小的候选，并按 2.4.5 节的惯性规则决定是否切换。"],
        ["13", "    计算第一跳主边可发送量；将该量顺序预留至下一跳和路径中的后续服务节点。"],
        ["14", "    若路径满足已核验的 L2/L18 并行设施映射且仍有剩余客流，则向配对设施提出附加流量。"],
        ["15", "  end for"],
        ["16", "  按边能力、整数容量信用与剩余能力对提出流量进行整数化。"],
        ["17", "  按式 (4)–(5)对同一目的节点的合流客流施加共享储存约束。"],
        ["18", "  将获准流量加入在途队列，实际到达时刻为 t+τ_uv；清除本步引导预留。"],
        ["19", "  t ← t+Δt。"],
        ["20", "end while；输出疏散时程、路径/出口分配、设施通过量和拥挤指标。"],
    ]
    add_table(doc, ["行", "操作"], pseudo, [650, 8710], font_size=8.7, first_col_left=False)

    add_heading(doc, "2.6 计算复杂度", 2)
    add_body(doc, "设当前时间步有 n_a 个活动节点、m 个出口，网络含 |V| 个节点和 |E| 条边。边权的统一刷新为 O(|E|)。当前实现对每个活动节点和每个可达出口分别执行一次 A*；使用二叉堆时，单次搜索的最坏复杂度为 O[(|E|+|V|)log|V|]，故单步最坏时间复杂度为 O{n_a m[(|E|+|V|)log|V|]+|E|}。顺序预留、惯性判断、局部分流和容量分配均不改变该主导阶。内存复杂度为 O(|V|+|E|)，另加在途客流记录。该复杂度说明当前实现适合车站级介观网络；若扩展到超大区域网络，可通过多目标最短路、出口反向势场或批量最短路降低重复搜索次数。")

    add_heading(doc, "2.7 参数与代码对应", 2)
    add_caption(doc, "表 2  AA 及高负荷介观传播的主要参数")
    add_table(doc,
              ["参数", "当前值", "作用与代码含义", "性质"],
              [
                  ["Δt", "0.5 s", "仿真、到达率更新与容量执行的时间步", "离散化参数"],
                  ["v_free", "1.427 m/s", "A* 启发式自由流速度；Fruin 分段速度上限", "文献/基线一致"],
                  ["ρ_free", "0.2 人/m²", "低于该密度时按自由流速度", "文献/基线一致"],
                  ["ρ_jam^route", "4.0 人/m²", "路径代价中的零速度阈值", "文献/基线一致"],
                  ["密度斜率", "0.3549", "0.2–4.0 人/m²内的速度线性衰减", "文献/基线一致"],
                  ["v_cap", "1.427 / 0.75 / 0.50 m/s", "平地/楼梯/扶梯的搜索速度上限", "设施参数"],
                  ["η", "0.30", "服务节点到达率 EMA 的新观测权重", "算法参数"],
                  ["最小保持时间", "2.0 s", "抑制频繁路径切换", "算法参数"],
                  ["常规切换阈值", "3%", "新路径至少降低 3% 才切换", "算法参数"],
                  ["强制切换阈值", "20%", "显著更优时忽略最小保持时间", "算法参数"],
                  ["退化触发", ">50% 且替代低 2%", "识别原路径相对选择时显著恶化", "算法参数"],
                  ["ρ_jam^store", "5.4 人/m²", "高负荷下游储存与回堵阈值", "基本图参数"],
                  ["T_buf", "18 s", "拓扑连接节点的最小流量缓冲", "数值稳定参数"],
              ], [1500, 1450, 4560, 1850], font_size=8.6, first_col_left=True)
    add_body(doc, "说明：η、切换阈值和 T_buf 是当前实现参数，不应写成已经由文献普遍验证的固定常数。正式投稿前应在验证章节报告敏感性或消融结果。特别是 4.0 与 5.4 人/m²分别属于路径搜索和储存回堵模块，任何统一或调整都需要重新标定，而不能只为改善结果而改动。", first=False, after=6, italic=True)

    add_caption(doc, "表 3  Improved A* 与 AA 的机制差异（按当前代码）")
    add_table(doc,
              ["比较项", "Improved A*", "AdaptiveQueueAwareAStar"],
              [
                  ["边代价量纲", "0.15×长度 + 0.85×时间，混合量纲", "通行时间 + 预测等待时间，统一为秒"],
                  ["高密度处理", ">3.0 人/m²移除节点/封闭边", "0.2–4.0 人/m²连续增大代价；速度近零时不可达"],
                  ["瓶颈服务", "未显式建模", "闸机、楼梯、扶梯使用 μ_v 和流体队列"],
                  ["未来流入", "不考虑", "EMA 到达率 + 同步内顺序预留"],
                  ["重规划", "拥堵集合改变或缓存路径失效", "每步计算候选，但受路径惯性约束"],
                  ["基本决策", "缓存的单一路径/下一跳", "动态单一路径/下一跳"],
                  ["并行设施", "无显式局部分流", "只对已核验 L2/L18 设施对补充分流"],
                  ["下游回堵", "与同一传播模块耦合", "与同一传播模块耦合"],
              ], [1550, 3500, 4310], font_size=8.7, first_col_left=True)

    add_heading(doc, "2.8 方法适用性与限制", 2)
    add_body(doc, "AA 的可解释优势是把路径选择直接关联到“到达瓶颈需要走多久”和“到达时预计还要等多久”，并通过本步预留降低多个来源同时涌向同一设施的同步偏差。其均衡效果是这些物理时间项作用后的结果，不是通过预设出口比例或直接优化 Gini 系数得到的。")
    add_body(doc, "现阶段仍有四项限制。第一，客流按节点聚合，无法表示个体超越、结伴、逆行与恐慌行为；第二，EMA 预测只外推到下一条边的预计到达时刻，不是全路径时变最短路；第三，并行设施修正依赖人工核验的物理配对，尚未自动从拓扑识别；第四，当前对每个活动节点逐出口执行 A*，网络规模进一步扩大时计算量会明显增加。因此，Pathfinder 更适合作为独立微观验证工具，而不应被写成 AA 必须逐项复现的“真值模型”。")

    add_heading(doc, "本章参考文献", 2)
    refs = [
        "Hart, P.E., Nilsson, N.J., Raphael, B., 1968. A formal basis for the heuristic determination of minimum cost paths. IEEE Transactions on Systems Science and Cybernetics 4(2), 100–107. https://doi.org/10.1109/TSSC.1968.300136.",
        "Jin, W.-L., 2015. Point queue models: A unified approach. Transportation Research Part B: Methodological 77, 1–16. https://doi.org/10.1016/j.trb.2015.02.015.",
        "Li, M., Xu, C., Xu, Y., Ma, L., Wei, Y., 2022. Dynamic sign guidance optimization for crowd evacuation considering flow equilibrium. Journal of Advanced Transportation 2022, 2555350. https://doi.org/10.1155/2022/2555350.",
        "Mandal, T., Rao, K.R., Tiwari, G., 2023. Evacuation of metro stations: A review. Tunnelling and Underground Space Technology 140, 105304. https://doi.org/10.1016/j.tust.2023.105304.",
        "蒙盾, 胡卓, 张华军, 2022. 基于改进 A* 算法的多层邮轮疏散系统仿真. 系统仿真学报 34(6), 1375–1382. https://doi.org/10.16182/j.issn1004731x.joss.21-0075.",
        "Weidmann, U., 1993. Transporttechnik der Fussgänger: Transporttechnische Eigenschaften des Fussgängerverkehrs. Schriftenreihe des IVT, ETH Zürich, No. 90.",
        "Zuo, S., Mao, Z., Fan, C., Chen, X., Gong, M., Ren, J., Fan, X., Guo, Y., 2024. Dynamic planning of crowd evacuation path for metro station based on Dynamic Avoid Smoke A-Star algorithm. Tunnelling and Underground Space Technology 154, 106145. https://doi.org/10.1016/j.tust.2024.106145.",
        "GB/T 33668-2017, 2017. 地铁安全疏散规范. 中国标准出版社, 北京.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(f"[{i}] {ref}")
        set_run_font(r, size=9.0)

    # Core properties and metadata.
    doc.core_properties.title = "模型与算法：AdaptiveQueueAwareAStar"
    doc.core_properties.subject = "龙阳路换乘站高负荷疏散介观算法章节"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "metro evacuation; mesoscopic; A-star; queue prediction"
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
