from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import build_algorithm_method_chapter as b


ROOT = Path(__file__).resolve().parent
OUTDIR = ROOT / "outputs"
ASSETDIR = OUTDIR / "algorithm_method_tust_v2_assets"
OUT = OUTDIR / "算法方法章节_TUST精修版_20260706.docx"
FRAMEWORK = ASSETDIR / "图1_研究框架.png"
LOOP = ASSETDIR / "图2_AA动态更新流程.png"

# TUST-manuscript override on narrative_proposal: black hierarchy, grey tables,
# minimal running furniture, no corporate accent colours.
b.INK = "111111"
b.MUTED = "666666"
b.ACCENT = "222222"
b.PALE = "ECEFF1"
b.PALE2 = "F7F7F7"
b.LINE = "9A9A9A"


def make_framework(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    width, height = 2000, 940
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    f_title = ImageFont.truetype(b.font_path("msyhbd.ttc", "simhei.ttf"), 50)
    f_head = ImageFont.truetype(b.font_path("msyhbd.ttc", "simhei.ttf"), 35)
    f_body = ImageFont.truetype(b.font_path("msyh.ttc", "simsun.ttc"), 29)
    f_note = ImageFont.truetype(b.font_path("msyh.ttc", "simsun.ttc"), 25)

    def box(x0, y0, x1, y1, title, lines):
        draw.rounded_rectangle((x0, y0, x1, y1), 22, fill="#F5F5F5", outline="#555555", width=4)
        draw.rectangle((x0, y0, x1, y0 + 86), fill="#E4E7E9", outline="#555555", width=3)
        bb = draw.textbbox((0, 0), title, font=f_head)
        draw.text(((x0 + x1 - bb[2] + bb[0]) / 2, y0 + 20), title, font=f_head, fill="#151515")
        y = y0 + 118
        for line in lines:
            bb = draw.textbbox((0, 0), line, font=f_body)
            draw.text(((x0 + x1 - bb[2] + bb[0]) / 2, y), line, font=f_body, fill="#222222")
            y += 54

    def arrow(x0, y, x1):
        draw.line((x0, y, x1 - 22, y), fill="#666666", width=5)
        draw.polygon([(x1, y), (x1 - 25, y - 14), (x1 - 25, y + 14)], fill="#666666")

    title = "介观疏散建模与自适应排队感知路径决策框架"
    bb = draw.textbbox((0, 0), title, font=f_title)
    draw.text(((width - (bb[2] - bb[0])) / 2, 32), title, font=f_title, fill="#111111")
    boxes = [
        (70, 150, 460, 680, "输入与标定", ["车站拓扑与几何", "设施能力与速度", "初始客流及来源组", "规范/图纸/场景数据"]),
        (550, 150, 940, 680, "介观状态模型", ["节点人数与在途客流", "边密度与接收能力", "服务节点到达率", "容量约束与回堵"]),
        (1030, 150, 1420, 680, "AA 路径决策", ["时间统一边代价", "到达时刻队列预测", "多出口 A* 搜索", "预留流入与切换惯性"]),
        (1510, 150, 1930, 680, "执行与验证", ["整数流量与并行设施", "疏散时程/排队/拥挤", "消融与敏感性分析", "Pathfinder 聚合层验证"]),
    ]
    for x0, y0, x1, y1, head, lines in boxes:
        box(x0, y0, x1, y1, head, lines)
    for i in range(3):
        arrow(boxes[i][2] + 18, 410, boxes[i + 1][0] - 18)
    draw.line((1720, 700, 1720, 790), fill="#777777", width=4)
    draw.line((1720, 790, 745, 790), fill="#777777", width=4)
    draw.line((745, 790, 745, 700), fill="#777777", width=4)
    draw.polygon([(745, 680), (730, 708), (760, 708)], fill="#777777")
    note = "每个 Δt 根据实际到达与拥堵状态闭环更新"
    bb = draw.textbbox((0, 0), note, font=f_note)
    draw.text(((width - (bb[2] - bb[0])) / 2, 815), note, font=f_note, fill="#555555")
    img.save(path, dpi=(240, 240))


def tune_styles(doc: Document) -> None:
    b.configure_styles(doc)
    normal = doc.styles["Normal"]
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(4)
    for level, size, before, after in ((1, 16, 14, 7), (2, 13, 10, 5), (3, 11.5, 7, 3)):
        style = doc.styles[f"Heading {level}"]
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(17, 17, 17)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_intro(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    b.set_para(p, after=6, line=1.15, first_line=0)
    run = p.add_run(text)
    b.set_run_font(run, size=10.5)


def build() -> Path:
    OUTDIR.mkdir(exist_ok=True)
    make_framework(FRAMEWORK)
    b.make_flowchart(LOOP)
    doc = Document()
    tune_styles(doc)
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.header_distance = Inches(0.45)
    sec.footer_distance = Inches(0.45)
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("模型与算法 | AdaptiveQueueAwareAStar")
    b.set_run_font(hr, size=8.2, color=b.MUTED)
    b.add_page_number(sec.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("2 模型与算法")
    b.set_run_font(run, east="黑体", size=17, bold=True, color="111111")
    add_intro(doc, "本节将多层换乘站高负荷疏散表述为介观有向网络上的时变路径分配问题，并提出自适应排队感知 A* 算法（AdaptiveQueueAwareAStar，简称 AA）。章节结构参考近年 Tunnelling and Underground Space Technology 算法类论文的常见写法：先界定研究问题并建立基础物理模型，再针对基准方法的具体缺陷给出改进机制及数学表达，最后独立说明求解流程、参数来源、验证设计与适用边界（Cai et al., 2022; Dong et al., 2022; Shen et al., 2022; Zhang et al., 2022; Zuo et al., 2024; Li et al., 2025）。")

    b.add_heading(doc, "2.1 研究框架与问题定义", 2)
    b.add_body(doc, "研究对象为多线路、多楼层、多出口地铁换乘站在高负荷突发疏散下的路径引导。给定车站有向网络 G=(V,E)、时变网络状态 S(t)、出口集合 D 以及设施服务能力，算法需要在每个离散决策时刻为存在待疏散客流的节点选择一条通往可达出口的路径。与一次性静态最短路不同，边代价随密度、服务队列和在途客流变化；与全局系统最优分配不同，AA 属于带短时预测和同步内预留的滚动最短时间决策。")
    b.add_equation(doc, "Pᵤ*(t) = arg min  C(P,t),   P∈℘ᵤₑ, e∈Dᵤ", 1)
    b.add_equation(doc, "C(P,t) = ∑(i,j)∈P cᵢⱼ(t)", 2)
    b.add_body(doc, "其中，℘ᵤₑ 为从活动节点 u 到出口 e 的可行路径集合，Dᵤ 为 u 的可达出口集合，cᵢⱼ(t) 为以秒计量的瞬时边代价。式 (1) 的目标是最小化当前状态快照下的预测路径时间，而不是直接最小化出口 Gini、总拥堵暴露或全体乘客总旅行时间。因此，后续观察到的负荷均衡是排队反馈和预留机制产生的结果，不能写成算法预先给定的全局目标。")
    b.add_caption(doc, "图 1  介观疏散模型与 AA 路径算法总体框架")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(FRAMEWORK), width=Inches(6.15))

    b.add_heading(doc, "2.2 介观疏散网络", 2)
    b.add_heading(doc, "2.2.1 网络表示与状态变量", 3)
    b.add_body(doc, "将站台等候区、站厅与换乘连接点、闸机、楼梯、扶梯和出口抽象为节点 V，将可通行连接抽象为有向边 E。边 (u,v) 具有长度 lᵤᵥ、有效宽度 wᵤᵥ、标称通行能力 Cᵤᵥ 和设施类型；节点 v 具有有效面积 Aᵥ、当前人数 Nᵥ(t) 以及节点服务率 μᵥ。闸机、楼梯和扶梯被定义为容量服务节点 Vˢ。乘客离开节点后进入在途队列，经过密度相关的边行程时间后才到达下游，因此不会在一个时间步内跨越多条边。")
    b.add_body(doc, "该模型属于介观模型。状态量是节点和边上的整数客流包，决策对象是处于同一节点的客流，而不是具有独立感知、加速度和避碰行为的微观代理。当前龙阳路实例包含 555 个节点、1381 条有向边和 16 个出口；高负荷双向满载场景为 17,905 人。上述数字描述案例实现，不构成算法适用规模的理论上限。")
    b.add_caption(doc, "表 1  主要符号及定义")
    b.add_table(doc, ["符号", "定义", "单位"], [
        ["G=(V,E)", "车站介观有向网络", "—"],
        ["lᵤᵥ, wᵤᵥ", "边 (u,v) 的长度与有效宽度", "m"],
        ["Nᵥ(t), Aᵥ", "节点 v 的人数与有效面积", "人；m²"],
        ["Cᵤᵥ(t)", "边 (u,v) 的有效接收能力", "人/s"],
        ["μᵥ", "闸机、楼梯或扶梯节点的服务率", "人/s"],
        ["λ̂ᵥ(t)", "服务节点 v 的 EMA 平滑观测到达率", "人/s"],
        ["Rᵥ(t)", "本时间步内分配至节点 v 的引导预留量", "人"],
        ["ρᵤᵥ(t)", "边 (u,v) 的路径决策有效密度", "人/m²"],
        ["Q̂ᵥ(t+τ)", "预计到达时刻的服务队列", "人"],
        ["cᵤᵥ(t)", "AA 使用的广义边时间代价", "s"],
    ], [1450, 6510, 1400], font_size=8.9, first_col_left=False)

    b.add_heading(doc, "2.2.2 流守恒与可执行约束", 3)
    b.add_body(doc, "设 Δt=0.5 s，Aᵥᵏ 和 Dᵥᵏ 分别表示第 k 个时间步实际到达节点 v 和从 v 获准出发的人数，则非出口节点满足离散流守恒：")
    b.add_equation(doc, "Nᵥᵏ⁺¹ = Nᵥᵏ + Aᵥᵏ − Dᵥᵏ", 3)
    b.add_equation(doc, "0 ≤ ∑ⱼ xᵥⱼᵏ = Dᵥᵏ ≤ Nᵥᵏ,    0 ≤ xᵤᵥᵏ ≤ CᵤᵥᵏΔt", 4)
    b.add_body(doc, "xᵤᵥᵏ 为从 u 向 v 提出的移动量。代码采用容量信用累计解决 CᵤᵥᵏΔt<1 时的整数流量问题，并在同一节点存在多个已提出物理边时，将主边未接收的剩余量分配至仍有能力的边。")
    b.add_body(doc, "高负荷场景进一步限制下游可接收人数。非出口节点的储存上限取面积容量和最小流量缓冲容量中的较大值：")
    b.add_equation(doc, "Kᵥ = max(Aᵥeff ρjam,store, μᵥout Tbuf, 1)", 5)
    b.add_equation(doc, "Sᵥᵏ = max{⌊Kᵥ−Nᵥᵏ−Iᵥtransit,k⌋,0}", 6)
    b.add_equation(doc, "∑ᵤ xᵤᵥᵏ ≤ Sᵥᵏ", 7)
    b.add_body(doc, "其中 Iᵥtransit,k 是已经从上游放行、正在前往 v 的在途人数。多个入口共享同一个 Sᵥᵏ，使下游空间在出发时被预占，从而形成可解释的上游回堵，而不是在下游节点无限叠加客流。")

    b.add_heading(doc, "2.2.3 密度相关运动模型", 3)
    b.add_body(doc, "路径决策使用边在途密度和下游空间投影密度的较大值。边代表面积由长度与有效宽度估计，并扣除显式障碍面积：")
    b.add_equation(doc, "ρᵤᵥ(t)=max{[Nᵥ(t)+Rᵥ(t)]/max(lᵤᵥwᵤᵥ−Aobs,ᵤᵥ,0.1), ρlink,ᵤᵥ(t)}", 8)
    b.add_body(doc, "路径搜索阶段沿用 Improved A* 基线采用的 Fruin 型分段密度–速度关系，以维持对比的一致性（Meng et al., 2022）：")
    b.add_equation(doc, "vF(ρ)={1.427, ρ≤0.2; max(1.427−0.3549ρ,0), 0.2<ρ≤4.0; 0, ρ>4.0}", 9)
    b.add_equation(doc, "vᵤᵥ(t)=min[vF(ρᵤᵥ(t)), vcap,ᵤᵥ],    τᵤᵥwalk(t)=lᵤᵥ/vᵤᵥ(t)", 10)
    b.add_body(doc, "vcap,ᵤᵥ 对平地、楼梯和扶梯分别取 1.427、0.75 和 0.50 m/s。4.0 人/m²是路径搜索中的零速度阈值；5.4 人/m²仅用于高负荷储存回堵。二者属于不同子模型，正式论文需分别报告敏感性，不应为了获得更好的结果而任意统一。")
    b.add_heading(doc, "2.3 基准算法：Improved A*", 2)
    b.add_body(doc, "基准算法根据蒙盾等（2022）的改进 A* 思路进行车站网络适配。其评价函数为 f=g+h，边代价同时包含长度和密度折减后的旅行时间：")
    b.add_equation(doc, "gᵤᵥbase(t)=αlᵤᵥ+βlᵤᵥ/vᵤᵥ(t),    hbase(n,e)=γd(n,e)", 11)
    b.add_body(doc, "取 α=0.15、β=0.85、γ=0.10。当节点空间密度或候选边有效密度超过 3.0 人/m²时，将其从本轮搜索空间中移除；只有拥堵节点集合改变或缓存路径失效时才重规划。该方法能够绕开当前高密度区域，但未显式表示瓶颈的服务队列，而且长度项与时间项量纲不同。本文保留这一实现作为文献基准，而不将其混合代价解释为物理时间。")

    b.add_heading(doc, "2.4 AdaptiveQueueAwareAStar 算法", 2)
    b.add_heading(doc, "2.4.1 改进思路", 3)
    b.add_body(doc, "TUST 中的算法论文通常将“基础模型”和“为解决特定缺陷而加入的机制”分开陈述。例如，Zuo et al. (2024) 先建立火灾风险模型，再修改 A* 评价函数和动态更新机制；Dong et al. (2022) 先定义目标函数与约束，再说明 NSGA-II 的求解流程；Li et al. (2025) 则在初始社会力模型上增加洪水作用项并先完成模型验证。按同一逻辑，AA 不重新定义介观传播模型，而针对基准算法的三个缺陷进行改造：混合量纲、忽略瓶颈服务队列、同一步多来源同步涌向同一设施。")
    b.add_caption(doc, "表 2  基准缺陷与 AA 改进机制的对应关系")
    b.add_table(doc, ["基准缺陷", "AA 机制", "直接数学作用", "预期作用"], [
        ["长度与时间混合", "时间统一的边代价", "全部边代价均以秒计", "提高物理可解释性"],
        ["只观测当前密度", "到达率 EMA 与流体队列预测", "在预计到达时刻评价队列", "提前识别瓶颈延误"],
        ["各来源基于同一快照决策", "同步内顺序预留", "后处理来源可观测已分配流入", "降低同步集中"],
        ["动态代价可能引起振荡", "路径切换惯性", "最短保持期与改善阈值", "稳定路径引导"],
        ["聚合节点掩盖并行设施", "经核验的局部并行修正", "剩余流量仅使用配对设施", "消除介观锁定偏差"],
    ], [2050, 2450, 2650, 2210], font_size=8.5, first_col_left=True)

    b.add_heading(doc, "2.4.2 到达率估计与队列预测", 3)
    b.add_body(doc, "每个时间步统计实际完成入边行程并到达容量服务节点 v 的人数 aᵥᵏ，以指数移动平均估计到达率：")
    b.add_equation(doc, "λ̂ᵥᵏ=η(aᵥᵏ/Δt)+(1−η)λ̂ᵥᵏ⁻¹", 12)
    b.add_body(doc, "η=0.30。对将从 u 进入 v 的客流，当前有效队列由实际在场人数和本步已经分配的引导预留构成，并采用确定性流体队列守恒外推至预计到达时刻（Jin, 2015）：")
    b.add_equation(doc, "Qᵥnow(t)=Nᵥ(t)+Rᵥ(t)", 13)
    b.add_equation(doc, "Q̂ᵥ[t+τᵤᵥwalk]=max{Qᵥnow(t)+[λ̂ᵥ(t)−μᵥ]τᵤᵥwalk,0}", 14)
    b.add_equation(doc, "τᵥwait[t+τᵤᵥwalk]=Q̂ᵥ[t+τᵤᵥwalk]/μᵥ", 15)
    b.add_body(doc, "该预测不是 M/M/1 稳态随机排队，也没有假定泊松到达；它只描述一条边行程时间内的队列净增长或消散。等待项只在进入闸机、楼梯或扶梯节点时计入一次，设施内部的物理通过时间仍包含在 τᵤᵥwalk 中。")

    b.add_heading(doc, "2.4.3 时间统一代价与多出口 A* 搜索", 3)
    b.add_equation(doc, "cᵤᵥ(t)=τᵤᵥwalk(t)+𝟙(v∈Vˢ)τᵥwait[t+τᵤᵥwalk(t)]", 16)
    b.add_body(doc, "对每个候选出口 e，启发式采用预计算几何最短距离除以自由流速度：")
    b.add_equation(doc, "h(n,e)=d(n,e)/vfree", 17)
    b.add_equation(doc, "f(n,e;t)=∑(u,v)∈P(s,n)cᵤᵥ(t)+h(n,e)", 18)
    b.add_body(doc, "由于实际边速度不超过 vfree=1.427 m/s 且等待时间非负，式 (17) 对冻结状态快照下的剩余代价不作高估，可作为可采纳启发式（Hart et al., 1968）。算法对所有可达出口分别执行 A*，最终按照完整路径边代价之和排序，而不是用启发式值直接比较出口。")

    b.add_heading(doc, "2.4.4 顺序预留与引导稳定性", 3)
    b.add_body(doc, "在同一时间步内，AA 依次处理有客流节点。某一来源确定路径后，将拟发送量计入第一跳下游节点以及路径中后续容量服务节点的 Rᵥ(t)，并立即刷新这些节点入边的代价；随后处理下一个来源。Rᵥ(t) 只是同步内的预测量，时间步结束后清零，真实客流仍由式 (3)–(7) 的传播与容量约束执行。该机制类似于对同步决策偏差的轻量修正，并不等价于求解动态用户均衡或系统最优。")
    b.add_body(doc, "为避免相近路径之间频繁振荡，算法保存上次路径、下一跳、切换时刻和选择时成本。新路径至少降低 3% 才常规切换；路径最少保持 2.0 s；若新路径降低 20% 以上，可强制切换；若原路径相对选择时成本增长超过 50%，且替代路径至少再低 2%，触发退化切换。上述阈值是当前实现参数，需在敏感性分析中报告，而不能表述为通用经验常数。")

    b.add_heading(doc, "2.4.5 经核验并行设施的局部修正", 3)
    b.add_body(doc, "一般节点仍采用单一下一跳。只有当两部设施已由实际拓扑确认连接同一上游空间和同一下游走向时，才允许局部并行修正。当前映射包括 L2 的 Stair_L2_1/Stair_L2_2、Stair_L2_3/Escalator_L2_up1，以及 L18 的 E1/E2、S1/S2 闸机对。算法先确定完整下游路径；若主设施本步能力不能接收全部等待客流，再将剩余量送入保持同一下游节点的配对设施。该修正用于消除介观聚合造成的“微小代价差锁定”，不改变出口选择，也不能概括成任意多路径分流。")
    b.add_heading(doc, "2.5 动态求解流程", 2)
    b.add_caption(doc, "图 2  AA 与介观传播模型的闭环更新流程")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    p.add_run().add_picture(str(LOOP), width=Inches(4.70))
    doc.add_page_break()
    b.add_caption(doc, "算法 1  AdaptiveQueueAwareAStar 与介观传播模型的耦合求解")
    pseudo = [
        ["1", "输入 G=(V,E)、出口 D、初始需求与 Δt；预计算 d(n,e)"],
        ["2", "初始化在途记录、λ̂ᵥ、路径状态与容量信用"],
        ["3", "当网络中或在途队列中仍有乘客时，执行循环"],
        ["4", "处理实际到达；更新 Nᵥ、出口计数和 aᵥᵏ"],
        ["5", "按式 (12) 更新 λ̂ᵥ；刷新边占用与密度"],
        ["6", "清零 Rᵥ；按式 (8)–(16) 计算全部边的 cᵤᵥ"],
        ["7", "依次处理每个有客流的非出口节点 u"],
        ["8", "以式 (17) 为启发式，对每个可达出口执行 A*"],
        ["9", "按 C(P,t) 排序完整候选路径，并应用路径切换惯性"],
        ["10", "提出第一跳流量，并沿下游服务节点进行引导预留"],
        ["11", "若存在经核验的局部配对且仍有剩余需求，提出配对流量"],
        ["12", "结束节点循环"],
        ["13", "在边能力与剩余能力约束下对提出流量整数化"],
        ["14", "按式 (5)–(7) 执行共享接收空间约束"],
        ["15", "将获准流量加入在途队列，并清除引导预留"],
        ["16", "令 t←t+Δt"],
        ["17", "结束循环；输出疏散、排队、拥堵、设施和出口统计"],
    ]
    b.add_table(doc, ["行", "操作"], pseudo, [650, 8710], font_size=8.55, first_col_left=False)

    b.add_heading(doc, "2.6 参数来源与可复现性", 2)
    b.add_caption(doc, "表 3  参数、来源及报告要求")
    b.add_table(doc, ["参数", "取值", "作用", "来源/报告要求"], [
        ["Δt", "0.5 s", "状态更新与容量执行", "数值设置；需报告敏感性"],
        ["vfree", "1.427 m/s", "自由流速度与 A* 下界", "Fruin/Meng 基准"],
        ["ρfree", "0.2 人/m²", "自由流密度阈值", "Fruin/Meng 基准"],
        ["ρjam,route", "4.0 人/m²", "路径搜索零速度阈值", "Fruin/Meng 基准"],
        ["密度斜率", "0.3549", "线性速度折减", "Fruin/Meng 基准"],
        ["vcap", "1.427/0.75/0.50 m/s", "平地/楼梯/扶梯速度上限", "设施实现参数"],
        ["η", "0.30", "到达率 EMA 权重", "算法设置；需做敏感性"],
        ["最短保持时间", "2.0 s", "路径最短保持时长", "算法设置；需做敏感性"],
        ["切换阈值", "3% / 20%", "常规/强制路径切换", "算法设置；需做敏感性"],
        ["退化触发", ">50%，替代路径再低 2%", "识别已选路径恶化", "算法设置；需做敏感性"],
        ["ρjam,store", "5.4 人/m²", "接收空间与回堵", "Weidmann 基本图"],
        ["Tbuf", "18 s", "连接节点最小缓冲", "数值稳定设置；需论证"],
    ], [1450, 1800, 3100, 3010], font_size=8.45, first_col_left=True)
    b.add_body(doc, "设施标称能力依据 GB/T 33668-2017 与模型中记录的设备数量或有效宽度换算。表 3 明确区分了文献参数、物理设施参数和当前算法/数值参数。正式投稿时，η、切换阈值、Δt 与 Tbuf 至少需要给出单因素敏感性或消融结果；若没有标定证据，应称为 implementation setting，而不是 empirically validated constant。", first=False, italic=True)

    b.add_caption(doc, "表 4  基准算法与 AA 的机制对比")
    b.add_table(doc, ["对比方面", "Improved A*", "AdaptiveQueueAwareAStar"], [
        ["代价量纲", "0.15×长度 + 0.85×时间", "通行时间 + 预测等待时间"],
        ["拥堵处理", ">3.0 人/m²时阻断", "到路径拥挤阈值前连续增大代价"],
        ["瓶颈服务", "未显式表示", "服务率与流体队列"],
        ["近未来流入", "未表示", "观测到达率 EMA + 同步内预留"],
        ["重规划", "拥堵集合或路径变化时", "每步更新，但受切换惯性约束"],
        ["路径形式", "缓存单一路径", "动态单一路径"],
        ["并行设施", "无显式修正", "仅修正经核验的 L2/L18 配对设施"],
        ["空间回堵", "共用传播模块", "共用传播模块"],
    ], [1600, 3550, 4210], font_size=8.55, first_col_left=True)
    b.add_heading(doc, "2.7 复杂度与验证设计", 2)
    b.add_heading(doc, "2.7.1 计算复杂度", 3)
    b.add_body(doc, "设当前步有 nₐ 个活动节点、m 个出口。边权统一更新为 O(|E|)。当前实现对每个活动节点和每个可达出口执行一次 A*；使用二叉堆时，单步最坏复杂度为 O{nₐm[(|E|+|V|)log|V|]+|E|}，内存复杂度为 O(|V|+|E|)，另加在途记录。该复杂度适用于车站级介观网络；扩展至区域级网络时，可用出口反向势场或批量多目标最短路减少重复搜索。")
    b.add_heading(doc, "2.7.2 机制验证", 3)
    b.add_body(doc, "为避免只用最终疏散时间宣称算法有效，验证应与改进机制一一对应。当前代码已提供 NoWaitingTime 变体：保留时间统一的密度相关旅行时间，但令 τwait=0。由此，AA 与 NoWaitingTime 的差异用于识别排队预测项的贡献；NoWaitingTime 与 Improved A* 的差异用于识别时间统一和连续密度代价的影响；AA 与 Improved A* 给出完整算法的净效果。预留机制、路径惯性和局部并行修正若要分别声称贡献，还需增加可独立关闭的开关并报告消融结果，现阶段不能根据最终结果反推每个机制都有效。")

    b.add_heading(doc, "2.8 适用范围与局限性", 2)
    b.add_body(doc, "AA 的核心贡献是将路径决策转化为“到达瓶颈前的通行时间 + 到达时仍需等待的时间”，并用同步内预留减少多个来源同时选择同一低成本设施的偏差。它不直接优化出口比例或 Gini 系数，也不保证全局系统最优。")
    b.add_body(doc, "模型仍有四项边界：(1) 节点聚合无法表示个体超越、结伴、逆行和恐慌；(2) 队列只外推至下一条边的预计到达时刻，不是全路径时变最短路；(3) 并行设施映射依赖人工核验，尚未自动识别；(4) 逐活动节点、逐出口搜索的计算量随网络扩大而增长。因此，Pathfinder 应作为独立微观验证工具，用于检验聚合疏散时程、出口/设施流量和拥挤位置是否处于相同数量级，而不要求逐人轨迹与介观模型完全一致。")
    b.add_heading(doc, "本节参考文献", 2)
    refs = [
        "Cai, Z., Zhou, R., Cui, Y., Wang, Y., Jiang, J., 2022. Influencing factors for exit selection in subway station evacuation. Tunnelling and Underground Space Technology 125, 104498. https://doi.org/10.1016/j.tust.2022.104498.",
        "Dong, Y.-H., Peng, F.-L., Zha, B.-H., Qiao, Y.-K., Li, H., 2022. An intelligent layout planning model for underground space surrounding metro stations based on NSGA-II. Tunnelling and Underground Space Technology 128, 104648. https://doi.org/10.1016/j.tust.2022.104648.",
        "Hart, P.E., Nilsson, N.J., Raphael, B., 1968. A formal basis for the heuristic determination of minimum cost paths. IEEE Transactions on Systems Science and Cybernetics 4(2), 100–107. https://doi.org/10.1109/TSSC.1968.300136.",
        "Jin, W.-L., 2015. Point queue models: A unified approach. Transportation Research Part B: Methodological 77, 1–16. https://doi.org/10.1016/j.trb.2015.02.015.",
        "Li, Y., Xu, D., Wang, J., Liu, J., Wang, Y., Jiang, J., 2025. Simulation of subway flood evacuation based on modified social force model. Tunnelling and Underground Space Technology 156, 106244. https://doi.org/10.1016/j.tust.2024.106244.",
        "Mandal, T., Rao, K.R., Tiwari, G., 2023. Evacuation of metro stations: A review. Tunnelling and Underground Space Technology 140, 105304. https://doi.org/10.1016/j.tust.2023.105304.",
        "蒙盾, 胡卓, 张华军, 2022. 基于改进 A* 算法的多层邮轮疏散系统仿真. 系统仿真学报 34(6), 1375–1382. https://doi.org/10.16182/j.issn1004731x.joss.21-0075.",
        "Shen, Y., Ma, J., Fang, H., Lo, S.M., Shi, C., 2022. Deep reinforcement learning based train door adaptive control in metro tunnel evacuation optimization. Tunnelling and Underground Space Technology 128, 104636. https://doi.org/10.1016/j.tust.2022.104636.",
        "Weidmann, U., 1993. Transporttechnik der Fussgänger: Transporttechnische Eigenschaften des Fussgängerverkehrs. Schriftenreihe des IVT, ETH Zürich, No. 90.",
        "Yang, X., Zhu, H., Wan, J., Li, Y., Chen, Z., 2025. Adaptability-enhanced evacuation path optimization and safety assessment for subway station passengers in floods: From uncertain challenge to reliable escape. Tunnelling and Underground Space Technology 163, 106683. https://doi.org/10.1016/j.tust.2025.106683.",
        "Zhang, Y., Li, W., Rui, Y., Wang, S., Zhu, H., Yan, Z., 2022. A modified cellular automaton model of pedestrian evacuation in a tunnel fire. Tunnelling and Underground Space Technology 130, 104673. https://doi.org/10.1016/j.tust.2022.104673.",
        "Zuo, S., Mao, Z., Fan, C., Chen, X., Gong, M., Ren, J., Fan, X., Guo, Y., 2024. Dynamic planning of crowd evacuation path for metro station based on Dynamic Avoid Smoke A-Star algorithm. Tunnelling and Underground Space Technology 154, 106145. https://doi.org/10.1016/j.tust.2024.106145.",
        "GB/T 33668-2017, 2017. 地铁安全疏散规范. 中国标准出版社, 北京.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(f"[{i}] {ref}")
        b.set_run_font(run, size=8.7)

    doc.core_properties.title = "模型与算法：AdaptiveQueueAwareAStar"
    doc.core_properties.subject = "参照 TUST 算法类论文结构撰写的介观疏散方法章节"
    doc.core_properties.author = ""
    doc.core_properties.keywords = "metro evacuation; mesoscopic; A-star; predictive queue"
    settings = doc.settings.element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build())
