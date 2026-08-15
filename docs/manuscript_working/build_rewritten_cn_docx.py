from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
WORK = Path(__file__).resolve().parent
FIG = WORK / "figures_revised"
OUT = WORK / "43_tust_manuscript_rewritten_cn_two_loads.docx"

BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(92, 97, 101)
LIGHT = "E9ECEF"


def set_font(run, east="宋体", latin="Times New Roman", size=10.5, bold=False, italic=False, color=BLACK):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, header_bottom=True):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:color"), "333333")
        borders.append(el)
    for edge in ("left", "right", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    inside_h = OxmlElement("w:insideH")
    inside_h.set(qn("w:val"), "single")
    inside_h.set(qn("w:sz"), "2")
    inside_h.set(qn("w:color"), "D7D7D7")
    borders.append(inside_h)


def configure(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.2)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2)
    sec.right_margin = Cm(2.2)
    sec.header_distance = Cm(1.25)
    sec.footer_distance = Cm(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.first_line_indent = Cm(0.74)
    pf.space_before = Pt(0)
    pf.space_after = Pt(5)
    pf.line_spacing = 1.35

    specs = {
        "Heading 1": (14, 12, 6),
        "Heading 2": (12, 9, 4),
        "Heading 3": (11, 7, 3),
    }
    for name, (size, before, after) in specs.items():
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run("面向多线换乘站共享瓶颈协调的到达时刻队列感知疏散路径规划")
    set_font(r, east="黑体", size=17, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("Tunnelling and Underground Space Technology 中文重写稿 · 高负荷结果版")
    set_font(r, east="宋体", size=9.2, color=GREY)


def add_p(doc: Document, text: str, *, indent=True, italic=False, size=10.5, color=BLACK):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.35
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    r = p.add_run(text)
    set_font(r, size=size, italic=italic, color=color)
    return p


def _m_run(text):
    run = OxmlElement("m:r")
    node = OxmlElement("m:t")
    node.text = str(text)
    run.append(node)
    return run


def _m_arg(tag, *children):
    node = OxmlElement(f"m:{tag}")
    for child in children:
        if isinstance(child, str):
            node.append(_m_run(child))
        else:
            node.append(child)
    return node


def _m_sub(base, sub):
    node = OxmlElement("m:sSub")
    node.append(OxmlElement("m:sSubPr"))
    node.append(_m_arg("e", base))
    node.append(_m_arg("sub", sub))
    return node


def _m_sup(base, sup):
    node = OxmlElement("m:sSup")
    node.append(OxmlElement("m:sSupPr"))
    node.append(_m_arg("e", base))
    node.append(_m_arg("sup", sup))
    return node


def _m_subsup(base, sub, sup):
    node = OxmlElement("m:sSubSup")
    node.append(OxmlElement("m:sSubSupPr"))
    node.append(_m_arg("e", base))
    node.append(_m_arg("sub", sub))
    node.append(_m_arg("sup", sup))
    return node


def _m_frac(num, den):
    node = OxmlElement("m:f")
    node.append(OxmlElement("m:fPr"))
    node.append(_m_arg("num", num))
    node.append(_m_arg("den", den))
    return node


def _equation_nodes(number):
    r, sub, sup, subsup, frac = _m_run, _m_sub, _m_sup, _m_subsup, _m_frac
    if number == 1:
        return [r("ρ:E→ℛ∪{∅}")]
    if number == 2:
        return [
            subsup("c", "ij", "Imp"), r("(t)=α"), sub("l", "ij"), r("+β"),
            subsup("t", "ij", "move"), r("(t),    h(n)=γd(n,𝒳)")
        ]
    if number == 3:
        return [
            sub("Q̂", "r"), r("("), sub("τ", "r"), r(")=ℱ("),
            sub("Q", "r"), r("("), sub("t", "0"), r("), {("),
            sub("a", "k"), r(","), sub("m", "k"), r(")}, "),
            sub("μ", "r"), r(","), sub("τ", "r"), r(")")
        ]
    if number == 4:
        return [
            sub("τ", "j"), r("="), sub("τ", "i"), r("+"),
            subsup("t", "ij", "move"), r("+"),
            frac(_m_arg("e", sub("Q̂", "r"), r("("), sub("τ", "i"), r(")")), _m_arg("e", sub("μ", "r"))),
            r("+"), subsup("t", "r", "batch"), r("+"), subsup("t", "j", "space")
        ]
    if number == 5:
        return [
            sub("C", "j"), r("="), sub("C", "i"), r("+("), sub("τ", "j"),
            r("−"), sub("τ", "i"), r(")+λΔ"), sub("R", "ij")
        ]
    if number == 6:
        return [
            sub("T", "p"), r("=min{t:"), sup("E", "out"), r("(t)≥pN},    p∈{0.50,0.80,0.95,0.99,1.00}")
        ]
    if number == 7:
        return [
            sup("W", "stat"), r("="), sub("∑", "t"), sup("N", "stationary"), r("(t)Δt")
        ]
    if number == 8:
        return [
            r("F̂(t)="), sup("N", "−1"), subsup("∑", "i=1", "N"), r("1("),
            sub("T", "i"), r("≤t)")
        ]
    if number == 9:
        return [
            r("Δ"), sub("T", "i"), r("="), subsup("T", "i", "Improved"), r("−"), subsup("T", "i", "AA")
        ]
    raise ValueError(number)


def add_equation(doc: Document, latex: str, number: int):
    # `latex` is retained as the manuscript's canonical equation source.  The
    # rendered DOCX uses native Word OMML so the formula remains editable and
    # does not expose raw backslash commands in print/PDF output.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(8.1), WD_TAB_ALIGNMENT.CENTER)
    p.paragraph_format.tab_stops.add_tab_stop(Cm(16.2), WD_TAB_ALIGNMENT.RIGHT)
    p.add_run("\t")
    math = OxmlElement("m:oMath")
    for node in _equation_nodes(number):
        math.append(node)
    p._p.append(math)
    num = p.add_run(f"\t({number})")
    set_font(num, east="Times New Roman", latin="Times New Roman", size=10.0)


def add_figure(doc: Document, filename: str, caption: str, width_cm=16.2):
    path = FIG / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    shape = p.add_run().add_picture(str(path), width=Cm(width_cm))
    shape._inline.docPr.set("title", filename)
    shape._inline.docPr.set("descr", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.first_line_indent = None
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption)
    set_font(r, size=8.6)


def add_table(doc: Document, title: str, headers, rows, widths_cm):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, size=9.2, bold=True)

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    total_twips = int(sum(widths_cm) / 2.54 * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for i, (header, width) in enumerate(zip(headers, widths_cm)):
        cell = table.rows[0].cells[i]
        cell.width = Cm(width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd"))
        cell._tc.tcPr[-1].set(qn("w:fill"), LIGHT)
        set_cell_margins(cell)
        cell.text = str(header)
        for run in cell.paragraphs[0].runs:
            set_font(run, size=8.2, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, (value, width) in enumerate(zip(row, widths_cm)):
            cell = cells[i]
            cell.width = Cm(width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            cell.text = str(value)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                set_font(run, size=8.0)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def heading(doc: Document, text: str, level: int):
    return doc.add_heading(text, level=level)


def build():
    doc = Document()
    configure(doc)
    add_title(doc)

    heading(doc, "摘要", 1)
    add_p(doc, "多线换乘站的疏散路径不仅决定乘客移动距离，也决定不同客流批次到达共享楼扶梯、通道和出口的先后次序。若路径代价只读取决策时刻的设施状态，尚未到达但已经被分配至同一瓶颈的客流不会进入当前判断，可能导致未来服务需求被持续低估。针对这一状态—到达错位，本文提出到达时刻队列感知 A*（arrival-time queue-aware A*，AA*）。该方法把容量受限设施表示为共享服务资源，按已承诺客流的到达事件推进队列，并在时变多标签搜索中逐段累计移动时间、预计资源等待、空间接纳等待和密度暴露代价。研究以龙阳路多线换乘站为案例，站体几何来自项目 CAD；比较 2,187 人基准需求与叠加两列各线路列车客流后的 17,905 人高负荷需求。AA* 与 Improved A* 共用相同的物理加载层，并通过逐模块消融识别性能来源；两组路径进一步映射到 Pathfinder 连续空间模型，同时设置 Pathfinder Goto Any Exit 作为软件原生自主选出口参照。高负荷网络仿真中，AA* 相比 Improved A* 将 T95 从 1125 s 降至 889 s、T100 从 1486 s 降至 1222 s，累计静止暴露降低 32.4%，但总移动距离增加 15.5%，计算时间增至 16.3 倍；低负荷下 T95 从 289 s 降至 258 s，但 T100 均为 325 s。消融显示，资源队列等待代价是识别服务瓶颈的基础，到达时刻预测提供主要增益；密度暴露和多标签搜索主要影响最末端清空。Pathfinder 高负荷复现中，AA* 路径的平均完成时间和 T100 分别比 Improved A* 路径降低 9.1% 和 8.2%；低负荷下二者 T95 和 T100 相同，而 AA* 平均完成时间略高。Pathfinder Goto Any Exit 在两种负荷下均取得较短的典型乘客完成时间，但保留更长或相近的最终清空尾部，表明典型乘客效率与残余尾部清空并非同一目标。研究由此把复杂换乘站疏散路径规划表述为共享服务能力上的到达时刻协调问题，并给出负荷分层的尾部收益、等待暴露、绕行与计算代价证据。")
    add_p(doc, "关键词：地铁站疏散；动态路径规划；到达时刻队列；共享瓶颈；Pathfinder；多线换乘站", indent=False)

    heading(doc, "1 引言", 1)
    intro = [
        "多线换乘站在有限空间内连接多个站台、站厅、换乘通道和地面出口。全站疏散时，站台候车客流、列车到达客流、站厅客流和换乘客流并非沿相互独立的通道离开，而是在不同位置和时刻竞争同一组楼梯、自动扶梯、闸机区和出口。局部容量不足由此不只造成设施附近的瞬时拥堵，还会沿客流链向上游传播，并延迟与该瓶颈相距较远的客流批次。把站内空间表示为由活动区域、竖向设施和出口组成的有向设施网络，可以描述这种跨线路耦合；若进一步把设施服务过程与路径选择联立，则可以分析路径分配如何改变拥堵形成和消散[1,2]。",
        "既有地铁站疏散研究已经从不同角度连接空间建模、路径优化和微观仿真。面向实际地铁站的仿真—优化研究通常依次给出站体建模、场景构造、优化过程、疏散快照或密度分布，再以完成时间、拥堵和出口流量评价方案[3–5]。Yang 等在洪水疏散研究中进一步将水动力、结伴行为、道路中断和路径优化连接到 Pathfinder 执行评价[6]。这些研究共同说明，疏散路径不能只以几何长度衡量，结果也不能只由一个全站完成时间支撑。",
        "对于多个批次共享容量设施的情形，仍存在一个直接影响路径选择的时间错位。乘客在时刻 t0 做出决策，但其真正进入楼梯、闸机区或出口的时刻为未来的 τ。在 t0 与 τ 之间，设施会继续服务当前队列，也会接收其他已经分配但尚未到达的客流。如果候选路径仍以 t0 的可见密度或队列计算代价，就会忽略这些已承诺的未来到达量，使当前看似通畅的设施在乘客抵达前转化为瓶颈。因此，本文关注的不是再次调整几何最短路权重，而是把共享设施状态推进到候选客流的预计到达时刻。",
        "基于此，本文提出到达时刻队列感知 A*（AA*），并围绕一个主问题展开：在无火灾、建模设施均可用的多线换乘站疏散中，把当前队列、在途承诺流量和设施服务过程推进至预计到达时刻，能否减少共享瓶颈的等待暴露并压缩 T95–T100 尾部？为回答该问题，研究进一步考察三个方面：其一，AA* 相比 Improved A* 是否改变站级完成分布和线路清空过程；其二，变化如何通过关键设施及来源—出口流量重分配形成，并对应多少额外移动与计算代价；其三，资源等待、到达时刻预测、空间接纳、多标签与密度暴露机制分别贡献什么，以及网络模型中的差异能否在 Pathfinder 连续空间运动中复现。",
        "本文的贡献体现在三个层次。第一，在路径代价中显式区分决策时刻和预计到达时刻，把共享资源当前队列、服务消化和已承诺到达事件纳入同一队列推进过程。第二，在共同物理执行层下，通过完整 AA*、Improved A* 和五个单模块变体建立“总体结果—设施机制—模块贡献”的证据链。第三，将网络模型输出的两组路径映射到同一 Pathfinder 站体模型，并以 Pathfinder Goto Any Exit 补充自主选出口参照，从平均完成、上尾清空、拥堵暴露和移动距离四个维度检验不同运行点。",
    ]
    for text in intro:
        add_p(doc, text)

    heading(doc, "2 文献基础与研究定位", 1)
    heading(doc, "2.1 设施网络、动态加载与疏散路径优化", 2)
    add_p(doc, "Shen 等把地铁站设施及其流量联系表示为层级网络，分析设施过载后的级联传播和流量重分配[1]；Wen 等将动态路径选择、节点排队和客流加载耦合，用于描述综合客运枢纽的时空分布[2]。这两类工作确立了本文的基本对象：换乘站不是一组彼此独立的走廊，而是由容量受限设施构成的服务链。Guo 和 Zhang 的实际地铁站研究则采用“仿真场景—代理评价—多目标优化—回代仿真”的闭环，并同时展示站体、疏散过程热图、敏感性和优化前后结果[3]。")
    add_p(doc, "路径规划研究进一步把拥堵、风险、引导和出口负荷纳入目标。Yang 等针对引导乘客建立多目标路径模型，并用现场—仿真对照、疏散快照、轨迹、竖向设施密度和出口流量组成证据链[4]；其后续 TUST 研究在洪水条件下连接站体网络、时间预测、多目标路径和 Pathfinder 对照[5]。Yang 等 2025 年研究又把结伴比例、道路中断和洪水风险纳入路径优化，并在 Pathfinder 中执行自然疏散、ACO 和 ETACO 路径[6]。这些工作为“真实站体—计算路径—微观执行”的论文结构提供了直接依据，但其洪水、结伴和引导变量不进入本文无火灾、设施均可用的研究场景。")
    heading(doc, "2.2 Pathfinder 原生路径选择的比较角色", 2)
    add_p(doc, "Hua 等通过全尺寸多楼梯选择实验区分最短与局部最快选择，并将 Pathfinder 用于行为复现[7]。Pathfinder 技术手册将其内部路径规划表述为 Locally Quickest，Goto Any Exit 行为允许乘客前往任一可达出口[8,9]。因此，本文将 Pathfinder Goto Any Exit 定义为软件原生自主选出口参照，而不是 AA* 的第三种实现或经验真值。Improved A* 与 AA* 的主要比较在共同网络物理层完成；Pathfinder 负责检验两组预分配路径在连续空间运动中的结果方向，并额外展示软件原生自主选择的平均效率—尾部清空取舍。")

    heading(doc, "3 方法", 1)
    heading(doc, "3.1 站体空间与设施—通道网络", 2)
    add_p(doc, "研究对象为无火灾、无烟气且所有建模设施均可用条件下的全站疏散。公共空间表示为有向图 G=(V,E)：节点包括站台与候车区、站厅、换乘通道、楼梯、自动扶梯、闸机排队区、出口前区和安全出口；有向边表示相邻空间间的可行移动。当前计算模型包含 572 个节点、1414 条有向边和 16 个最终出口。图 1 将 CAD 与 Pathfinder 对象抽象为楼层—线路—换乘关系，以替代难以辨识的原始全节点网络图。")
    add_figure(doc, "fig_station_spatial_structure_cn.png", "图1 龙阳路多线换乘站的空间结构与模型对象。（a）依据 CAD 与 Pathfinder 建模对象抽象的楼层—线路—出口关系；（b）活动区域、竖向设施和最终出口在连续空间模型与网络模型中的对应角色。示意图不按比例绘制。")
    add_p(doc, "同一实体设施可能在图中连接多条边。本文设置资源集合 R，并以映射 ρ 把所有共享同一服务能力的边或节点归入同一资源，从而防止一个楼梯、闸机组或出口的容量因网络连接数而被重复计算。")
    add_equation(doc, r"\rho:E\rightarrow\mathcal{R}\cup\{\varnothing\}", 1)

    heading(doc, "3.2 共同物理加载层", 2)
    add_p(doc, "Improved A* 与 AA* 共用相同的 1 s 时间步长、密度相关移动、设施服务队列、下游有限接纳能力和溢回规则。路径算法只决定客流批次的后续路线；一旦路线被接受，移动、服务与接纳均由同一执行器推进。该设置使两方法之间的差异可归因于路径决策，而不是不同的速度或容量规则。Improved A* 的代价沿用参考实现中的长度—时间结构[10]；本文参数作为模型输入使用，不表述为龙阳路现场标定值。")
    add_equation(doc, r"c_{ij}^{Imp}(t)=\alpha l_{ij}+\beta t_{ij}^{move}(t),\qquad h(n)=\gamma d(n,\mathcal{X})", 2)

    heading(doc, "3.3 到达时刻队列与广义路径代价", 2)
    add_p(doc, "设 t0 为决策时刻，Qr(t0) 为共享资源 r 的当前队列，μr 为服务率，{(ak,mk)} 为已经分配至该资源但将在未来到达的事件。AA* 先按 ak 排序，在相邻事件间扣除可完成的服务量，再叠加 mk，从而得到候选批次于预计到达时刻 τr 面对的队列。该事件推进避免把所有在途客流一次性叠加，也避免只读取 t0 的瞬时状态。")
    add_equation(doc, r"\widehat Q_r(\tau_r)=\mathcal F\!\left(Q_r(t_0),\{(a_k,m_k)\},\mu_r,\tau_r\right)", 3)
    add_p(doc, "沿候选路径从节点 i 推进至 j 时，预计到达时间由物理移动、预计资源等待、批处理等待和空间接纳等待组成；广义代价再累计等待与移动期间的密度暴露。")
    add_equation(doc, r"\tau_j=\tau_i+t_{ij}^{move}+\widehat Q_r(\tau_i)/\mu_r+t_r^{batch}+t_j^{space}", 4)
    add_equation(doc, r"C_j=C_i+(\tau_j-\tau_i)+\lambda\Delta R_{ij}", 5)
    add_p(doc, "由于同一节点的后续代价依赖到达时刻，AA* 以 (τ,C) 保存互不支配的时间标签；较晚但当前代价较低的标签不会自动支配较早标签。已接受路线所产生的未来资源到达事件回写到事件索引，供后续批次查询。图 2 概括候选路径共享瓶颈、事件推进和逐段代价的关系。")
    add_figure(doc, "fig_aa_method_revised_cn.png", "图2 AA* 的到达时刻队列预测与路径代价。（a）几何长度相近的候选路径共享后续容量瓶颈；（b）当前队列、服务消化与已承诺到达事件被推进到预计到达时刻；（c）搜索逐段累计物理移动、资源等待、空间接纳和密度暴露，并保留互不支配的时间标签。", width_cm=14.2)

    heading(doc, "4 案例与实验设计", 1)
    heading(doc, "4.1 需求场景", 2)
    add_p(doc, "2,187 人基准需求采用研究实施阶段冻结的仿真输入，由各线路站台候车、站厅与换乘空间人口组成。人口输入的线路构成依据作者提供的人数设计记录确定，该记录使用上海地铁刷卡客流研究和龙阳路进站客流完成线路尺度换算[11]。本文按实际仿真含义称其为‘基准需求’，不使用节假日、周末或工作日标签强化场景含义。")
    add_p(doc, "高负荷需求在基准人口上叠加两列各线路列车客流：2 号线每列 2400 人、7 号线每列 1620 人、16 号线每列 1230 人、18 号线每列 1650 人、磁浮每列 959 人；列车客流共 15,718 人，场景总人数为 17,905 人。图 3 同时展示基准需求的空间来源和高负荷的线路构成。")
    add_figure(doc, "fig_demand_revised_cn.png", "图3 需求场景构成。（a）代码中冻结的 2,187 人基准需求在站台候车、站厅与换乘空间的分布；（b）两列各线路列车客流叠加后的 17,905 人高负荷需求。")
    add_table(doc, "表1 需求场景的人口构成（人）", ["线路", "站台", "站厅", "换乘", "列车1", "列车2", "总计"], [
        ("2号线", 236, 350, 526, 2400, 2400, 5912),
        ("7号线", 219, 112, 169, 1620, 1620, 3740),
        ("16号线", 42, 15, 27, 1230, 1230, 2544),
        ("18号线", 178, 125, 188, 1650, 1650, 3791),
        ("磁浮", 0, 0, 0, 959, 959, 1918),
        ("合计", 675, 602, 910, 7859, 7859, 17905),
    ], [2.4, 2.0, 2.0, 2.0, 2.1, 2.1, 2.0])

    heading(doc, "4.2 网络比较与评价指标", 2)
    add_p(doc, "网络主比较只改变路径方法。两组运行使用相同站体、人口、密度相关移动、资源服务率、空间接纳、溢回和完全疏散终止条件。站级结果由平均完成时间和 Tp 描述，Tp 表示累计完成比例首次达到 p 的时刻；累计静止暴露 Wstat 统计每个时间步仍在站内但未移动的人数。机制指标包括线路清空时间、来源—出口分配、总移动距离、出口和关键设施 Jain 指数以及墙钟运行时间。")
    add_equation(doc, r"T_p=\min\{t:E^{out}(t)\geq pN\},\quad p\in\{0.50,0.80,0.95,0.99,1.00\}", 6)
    add_equation(doc, r"W^{stat}=\sum_t N^{stationary}(t)\Delta t", 7)

    heading(doc, "4.3 单模块消融", 2)
    add_p(doc, "消融采用 leave-one-component-out 设计，以完整 AA* 为参照，每次仅关闭资源队列等待代价、到达时刻队列预测、空间接纳等待、时变多标签或密度暴露代价中的一项。所有变体使用同一高负荷需求和共同物理执行层。由于当前输出为每个确定性配置的一次完整运行，本文报告相对变化和机制一致性，不构造不存在的重复样本或显著性检验。")

    heading(doc, "4.4 Pathfinder 跨模型微观复现", 2)
    add_p(doc, "网络模型负责生成客流批次的路径和出口分配，Pathfinder 负责在连续空间中执行这些分配。高负荷复现包含 Improved A*、AA* 和 Pathfinder Goto Any Exit 三个协议。前两组把网络模型输出的完整路径合并为 Pathfinder 行为分配；第三组使用模型文件中的 `goto exit any` 行为，由软件自主选择可达出口。三组协议均使用 17,905 人，三个 `.geom` 文件的 SHA-256 完全一致。")
    add_p(doc, "Improved A* 与 AA* 共用同一站体几何、场景人口与连续空间运行条件。本节对三种协议统一比较全体乘客的经验累计完成曲线、完成时间分位数、拥堵暴露和移动距离；Goto Any Exit 作为 Pathfinder 原生自主选出口的场景级参照。")
    add_equation(doc, r"\widehat F(t)=N^{-1}\sum_{i=1}^{N}\mathbf{1}(T_i\leq t)", 8)

    heading(doc, "5 结果", 1)
    heading(doc, "5.1 高负荷站级完成与线路清空", 2)
    add_p(doc, "Improved A* 与 AA* 均完成 17,905 人疏散并满足人口守恒。AA* 将平均完成时间从 438.9 s 降至 335.5 s，T95 从 1125 s 降至 889 s，T99 从 1309 s 降至 1097 s，T100 从 1486 s 降至 1222 s。两条完成时间剖面从平均值到完全清空均保持间隔，且 T95–T100 的绝对差为 236–264 s，说明改善延伸至残余尾部。")
    add_p(doc, "线路结果进一步表明，全站改善通过线路间重新分配能力实现。AA* 将 2、7、16 和 18 号线的完全清空时间分别由 1483、1486、874 和 1309 s 降至 993、1222、344 和 418 s；磁浮由 590 s 增至 651 s。7 号线在两种方法下均为最后清空线路。该结果不要求每条线路同时改善，而显示 AA* 以局部让渡换取全站尾部缩短。")
    add_figure(doc, "fig_network_high_load_revised_cn.png", "图4 高负荷网络仿真的站级、线路级及代价结果。（a）平均值至 T100 的完成时间剖面；（b）各来源线路完全清空时间；（c）累计移动与静止人·秒构成；（d）总移动距离、累计静止暴露与墙钟运行时间的联合取舍。")
    add_table(doc, "表2 高负荷网络仿真主要结果", ["指标", "Improved A*", "AA*", "AA* 相对变化"], [
        ("平均完成时间 (s)", "438.9", "335.5", "−23.5%"),
        ("T95 (s)", "1125", "889", "−21.0%"),
        ("T100 (s)", "1486", "1222", "−17.8%"),
        ("累计静止 (人·s)", "6,552,858", "4,431,142", "−32.4%"),
        ("总移动距离 (m)", "1,481,121", "1,710,553", "+15.5%"),
        ("出口 Jain", "0.626", "0.711", "+13.6%"),
        ("关键设施 Jain", "0.162", "0.409", "+152.4%"),
        ("墙钟运行时间 (s)", "37.7", "615.3", "×16.3"),
    ], [5.2, 3.6, 3.6, 3.8])

    heading(doc, "5.2 负荷分层的网络层结果", 2)
    add_p(doc, "低负荷网络运行包含 2,187 人，Improved A* 与 AA* 均完成全站疏散，最终清空时间均为 325 s，且 7 号线均为最后清空线路。AA* 将 T95 从 289 s 降至 258 s（−10.7%），平均完成时间从 138.9 s 降至 135.8 s（−2.3%），累计静止暴露从 108,228 人·s 降至 91,853 人·s（−15.1%）。与此同时，总移动距离由 213,389 m 增至 221,539 m（+3.8%），墙钟运行时间由 8.0 s 增至 10.9 s（+35.8%）。线路层面，L2 和 L16 的完全清空时间分别提前 57 s 和 44 s，而 L7 仍保持 325 s 的全站尾部，因此低负荷下 AA* 的收益主要表现为中前段清空和等待减少，尚不足以改变最后一条线路的 T100。")
    add_figure(doc, "fig_load_stratified_network_cn.png", "图5 负荷分层的网络层结果。（a）2,187 人与 17,905 人条件下 Improved A* 和 AA* 的 T95/T100；（b）人均静止暴露；（c）AA* 相对 Improved A* 的绕行—等待取舍；（d）墙钟计算代价。低负荷和高负荷使用同一指标定义。")
    add_table(doc, "表3 两种负荷下的网络层主要结果（完成时间三元组依次为平均/T95/T100，单位：s）", ["负荷", "方法", "完成时间：平均/T95/T100", "静止人·s", "距离 (m)"], [
        ("低负荷", "Improved A*", "139 / 289 / 325", "108,228", "213,389"),
        ("低负荷", "AA*", "136 / 258 / 325", "91,853", "221,539"),
        ("高负荷", "Improved A*", "439 / 1125 / 1486", "6,552,858", "1,481,121"),
        ("高负荷", "AA*", "336 / 889 / 1222", "4,431,142", "1,710,553"),
    ], [2.2, 3.1, 4.4, 2.8, 2.6])

    heading(doc, "5.3 来源—出口流量重分配", 2)
    add_p(doc, "来源—出口矩阵显示，路径差异集中在少数大流量重分配。18 号线客流从 Exit 17 向 Exit 12 转移 1650 人；16 号线减少分配至 Exit 10 的 428 人，并增加分配至 Exit 11 西侧的 446 人；磁浮客流在 18–21 号出口间重新分布。2 号线则在多个本线出口间进行数百人的调整。每个来源线路总人数保持不变，因而这些差值直接表示出口路径组合的重构，而不是人口变化。")
    add_figure(doc, "fig_flow_redistribution_revised_cn.png", "图6 高负荷来源—出口流量重分配。（a）AA* 相对 Improved A* 的来源线路—最终出口人数差；（b）绝对变化最大的九个流向。蓝色表示 AA* 减少分配，红色表示增加分配。")

    heading(doc, "5.4 模块消融", 2)
    add_p(doc, "完整 AA* 的消融基准为 T95=889 s、T100=1222 s。去除资源队列等待代价后，T50、T95 和 T100 分别增加 73.2%、196.3% 和 182.7%，累计静止暴露增加 164.6%；关键设施 Jain 指数相对完整 AA* 下降 69.4%。该变体产生最大退化，说明显式资源等待是方法识别容量服务瓶颈的基础。")
    add_p(doc, "保留资源等待但关闭到达时刻队列预测后，T95 增至 1060 s、T100 增至 1427 s，累计静止暴露增加 45.2%，同时运行时间减少约 11.7%。去除密度暴露项和改为单标签均不改变 T95，但分别将 T100 增至 1254 s 和 1251 s，说明其独立作用集中在最末端残余客流。关闭规划层空间接纳等待后，各项汇总结果与完整 AA* 相同；共同物理层中的有限接纳和溢回仍保持开启。")
    add_figure(doc, "fig_ablation_revised_cn.png", "图7 AA* 的高负荷单模块消融。（a）相对完整 AA* 的结果退化，Jain 指数按降低方向计为退化；（b）T100 变化与墙钟运行时间变化。每个变体只关闭一个规划机制。")

    heading(doc, "5.5 Pathfinder 高负荷微观复现", 2)
    add_p(doc, "在 Pathfinder 连续空间运动中，AA* 路径相对 Improved A* 路径将平均完成时间从 435.5 s 降至 396.0 s，T95 从 1024.5 s 降至 985.1 s，T100 从 1414.6 s 降至 1298.3 s；平均拥堵时间由 311.3 s 降至 276.6 s。三项完成时间指标和拥堵暴露尾部均显示，AA* 的站级尾部改善能够在连续空间执行中复现。")
    add_p(doc, "Pathfinder Goto Any Exit 的平均完成时间为 362.5 s，低于两种预分配路径；其平均移动距离也最短，为 114.6 m。然而，该协议的 T95=1034.1 s、T100=1458.5 s，均高于 AA* 路径的 985.1 s 和 1298.3 s。经验累计曲线在上尾交叉，表明 Goto Any Exit 使较大比例乘客较早完成，却留下更长的最终清空阶段。")
    add_figure(doc, "fig_pathfinder_high_load_revised_cn.png", "图8 Pathfinder 高负荷完成时间与拥堵暴露分布。（a）全部乘客的经验累计完成曲线；（b）完成时间分位剖面；（c）乘客移动距离的经验累计分布；（d）累计拥堵暴露的尾部分布。Goto Any Exit 只做场景级比较。")
    add_p(doc, "拥堵组成进一步解释了三种协议的取舍。Goto Any Exit 的人均水平区域拥堵最低，但楼梯拥堵高于 AA*；AA* 的楼梯拥堵最低，而其移动距离分布高于 Goto Any Exit。全站平均完成、竖向设施暴露和完全清空因此不能由单一路径长度指标同时代表。")
    add_figure(doc, "fig_pathfinder_tradeoff_revised_cn.png", "图9 Pathfinder 高负荷的拥堵—距离—尾部取舍。（a）水平区域与楼梯拥堵构成；（b）乘客移动距离分布，白点为中位数，粗线为四分位区间；（c）人均完成时间与 T100 的联合运行点。")
    add_table(doc, "表4 Pathfinder 高负荷场景级结果", ["协议", "平均 (s)", "T95 (s)", "T100 (s)", "人均拥堵 (s)", "人均距离 (m)"], [
        ("Improved A*", "435.5", "1024.5", "1414.6", "311.3", "131.8"),
        ("AA*", "396.0", "985.1", "1298.3", "276.6", "131.7"),
        ("Pathfinder Goto Any Exit", "362.5", "1034.1", "1458.5", "253.0", "114.6"),
    ], [4.6, 2.5, 2.5, 2.6, 2.8, 2.8])

    heading(doc, "5.6 Pathfinder 低负荷微观复现", 2)
    add_p(doc, "低负荷 Pathfinder 三组协议均使用 2,187 名乘客和同一 SHA-256 几何文件。Improved A* 的平均完成时间、T95 和 T100 分别为 117.6、254.6 和 319.0 s；AA* 分别为 120.8、254.6 和 319.0 s。AA* 的平均拥堵时间为 9.2 s/人，略低于 Improved A* 的 9.4 s/人，但平均移动距离由 120.7 m 增至 124.7 m。由此可见，低负荷下两种预分配路径的 T95 和 T100 相同，AA* 的差异主要体现在略低的平均拥堵暴露和略高的移动距离，而没有转化为更短的连续空间平均完成时间。")
    add_p(doc, "Pathfinder Goto Any Exit 的平均完成时间为 102.2 s，T95 为 224.8 s，均低于两组预分配路径；但 T100 为 320.0 s，略高于二者的 319.0 s。其平均移动距离为 97.3 m，而人均拥堵时间为 14.8 s，高于 Improved A* 和 AA*。因此，低负荷下软件原生自主选出口更有利于大多数乘客较早完成，却没有改善最终一名乘客的清空时间，也没有降低总体拥堵暴露。")
    add_figure(doc, "fig_load_stratified_pathfinder_cn.png", "图10 Pathfinder 两种负荷下的连续空间复现。（a）三种协议的经验累计完成曲线；（b）完成时间分位剖面；（c）平均完成时间与尾部完成时间的联合比较；（d）拥堵暴露尾部。颜色区分协议，实线为低负荷，虚线为高负荷。")

    heading(doc, "6 讨论", 1)
    heading(doc, "6.1 到达时刻协调产生“绕行换等待”", 2)
    add_p(doc, "AA* 的主要收益不是几何缩短，而是减少未来共享服务竞争。网络主比较中，总移动距离与移动人·秒上升，而静止人·秒和完成尾部下降；到达预测消融又呈现“距离缩短、等待增加”的反向变化。负荷分层结果进一步表明，这一机制并非在所有负荷下都同样改变 T100：低负荷下 AA* 已降低 T95 和静止暴露，但 7 号线仍控制最终清空。两组证据共同说明，到达时刻队列是连接路径选择与瓶颈服务过程的关键状态。")
    heading(doc, "6.2 平均效率与尾部清空的分离", 2)
    add_p(doc, "Pathfinder Goto Any Exit 揭示了平均效率与尾部清空的分离，而且这一现象在两种负荷下均出现。自主选出口能够让大量乘客快速完成，但局部选择累积后，少量乘客可能进入持续更久的残余阶段。AA* 在高负荷下同时降低平均完成时间、T95、T100 和拥堵暴露；在低负荷下，AA* 与 Improved A* 的 T95、T100 相同，仅平均拥堵暴露略低而平均完成时间略高。因而本文同时报告平均值、尾部指标、移动代价和计算代价，避免用单一指标概括不同协议。")
    heading(doc, "6.3 批次级引导的实施含义", 2)
    add_p(doc, "AA* 的运行代价明显高于 Improved A*，高负荷时差异尤其明显，更适合转化为列车、站台或客流批次级引导策略，通过动态标志、广播或现场组织实施，而不是要求每名乘客进行高频个体重规划。低负荷结果显示，资源竞争减弱后，网络层的等待暴露仍可降低，但连续空间中的完成时间指标未进一步缩短，因此实施评价应同时关注批次级清空和连续空间运动结果。")

    heading(doc, "7 结论", 1)
    add_p(doc, "本文将多线换乘站疏散路径规划表述为共享服务能力上的到达时刻协调，并提出在时变多标签搜索中推进当前队列、在途承诺流量与资源服务过程的 AA*。高负荷网络仿真表明，AA* 相比 Improved A* 同时缩短 T95–T100 尾部并减少累计静止暴露，其实现机制是增加部分移动以避开未来共享瓶颈；低负荷下，AA* 仍降低 T95 和静止暴露，但不改变 T100。消融将资源队列等待识别为基础机制，将到达时刻预测识别为主要增益来源，并显示多标签与密度暴露主要收紧最末端清空。Pathfinder 连续空间复现表明，AA* 的平均和尾部改善在高负荷下可以复现；在低负荷下，两种预分配路径的 T95、T100 相同，AA* 主要表现为略低的平均拥堵暴露和略高的移动距离。软件原生 Goto Any Exit 在两种负荷下均取得较低平均或 T95，却保留略长的 T100，进一步区分了典型乘客效率与全站残余尾部。")

    heading(doc, "参考文献", 1)
    refs = [
        "[1] Shen Y, Yang H, Ren G, Ran B. Model cascading overload failure and dynamic vulnerability analysis of facility network of metro station. Reliability Engineering & System Safety 242 (2024) 109711. https://doi.org/10.1016/j.ress.2023.109711.",
        "[2] Wen X, Si B, Xu M, Zhao F, Jiang R. A passenger flow spatial-temporal distribution model for a passenger transit hub considering node queuing. Transportation Research Part C 163 (2024) 104640. https://doi.org/10.1016/j.trc.2024.104640.",
        "[3] Guo K, Zhang L. Simulation-based passenger evacuation optimization in metro stations considering multi-objectives. Automation in Construction 133 (2022) 104010. https://doi.org/10.1016/j.autcon.2021.104010.",
        "[4] Yang X, Yang Y, Li Y, Yang X. Path planning for guided passengers during evacuation in subway station based on multi-objective optimization. Applied Mathematical Modelling 111 (2022) 777–801. https://doi.org/10.1016/j.apm.2022.07.024.",
        "[5] Yang X, Dai W, Li Y, Yang X. An efficient evacuation path optimization for passengers in subway stations under floods. Tunnelling and Underground Space Technology 144 (2024) 105473. https://doi.org/10.1016/j.tust.2023.105473.",
        "[6] Yang X, Wan J, Zhu H, Xie C-Z, Zhang B. Optimization of passenger evacuation path in flood scenarios considering companion behaviors. Simulation Modelling Practice and Theory 145 (2025) 103212. https://doi.org/10.1016/j.simpat.2025.103212.",
        "[7] Hua Y, Zhao J, Li H-T, Duan L. Shortest or locally quickest? A prediction-based approach for evacuation choice simulation between multiple staircases. Journal of Safety Science and Resilience 5 (2024) 281–294. https://doi.org/10.1016/j.jnlssr.2024.04.001.",
        "[8] Thunderhead Engineering. Pathfinder Technical Reference Manual: Path Planning. https://www.thunderheadeng.com/docs/2024-2/pathfinder/technical-reference-manual/.",
        "[9] Thunderhead Engineering. Pathfinder User Manual: Behaviors—Goto Any Exit. https://www.thunderheadeng.com/docs/2026-1/pathfinder/behaviors/.",
        "[10] 蒙盾, 胡志强, 张洪雨. 基于改进 A* 算法的多层邮轮疏散系统仿真. 系统仿真学报 34(6) (2022) 1375–1382. https://doi.org/10.16182/j.issn1004731x.joss.21-0075.",
        "[11] Yu L, Liu H, Fang Z, Ye R, Huang Z, You Y. A new approach on passenger flow assignment with multi-connected agents. Physica A 628 (2023) 129175. https://doi.org/10.1016/j.physa.2023.129175.",
    ]
    for ref in refs:
        p = add_p(doc, ref, indent=False, size=8.8)
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.65)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(3)

    doc.core_properties.title = "面向多线换乘站共享瓶颈协调的到达时刻队列感知疏散路径规划"
    doc.core_properties.subject = "TUST 中文重写稿（低—高负荷结果版）"
    doc.core_properties.keywords = "metro evacuation; arrival-time queue; Pathfinder; AA*"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
