from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "manuscript_working"
FIG_DIR = OUT_DIR / "figures"
EQ_DIR = FIG_DIR / "equations"
DOCX_PATH = OUT_DIR / "34_tust_manuscript_working_draft_en_review_copy.docx"
FIG_PATH = FIG_DIR / "fig2_pathfinder_high_load_validation.png"
PF_SUMMARY_PATH = FIG_DIR / "table_pathfinder_high_load_summary.csv"

BLUE = "1769AA"
DARK = "1F2933"
MID = "58636D"
LIGHT = "EAF2F8"
PALE = "F5F7F8"
WHITE = "FFFFFF"


def set_run_font(run, name="Calibri", size=10.5, bold=None, italic=None, color=DARK):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), name)
    rpr.rFonts.set(qn("w:hAnsi"), name)
    rpr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=95, bottom=70, end=95):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "CDD5DB")
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_paragraph_keep(paragraph, keep_next=False, keep_lines=True):
    ppr = paragraph._p.get_or_add_pPr()
    if keep_next:
        node = OxmlElement("w:keepNext")
        ppr.append(node)
    if keep_lines:
        node = OxmlElement("w:keepLines")
        ppr.append(node)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def add_text_paragraph(doc, text, *, style=None, first_indent=True, italic=False, color=DARK):
    p = doc.add_paragraph(style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.14
    p.paragraph_format.space_after = Pt(5)
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.25)
    run = p.add_run(text)
    set_run_font(run, size=10.5, italic=italic, color=color)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=10.2)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_run_font(run, size=10.2)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    run = p.add_run(text)
    set_run_font(
        run,
        size={1: 14, 2: 11.5, 3: 10.5}.get(level, 10.5),
        bold=True,
        color=BLUE if level <= 2 else DARK,
    )
    set_paragraph_keep(p, keep_next=True)
    return p


def add_equation(doc, latex, number=None):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Inches(5.8)
    table.columns[1].width = Inches(0.5)
    left, right = table.rows[0].cells
    left.width = Inches(5.8)
    right.width = Inches(0.5)
    for cell in (left, right):
        set_cell_margins(cell, top=20, start=20, bottom=20, end=20)
        tc_pr = cell._tc.get_or_add_tcPr()
        borders = tc_pr.find(qn("w:tcBorders"))
        if borders is None:
            borders = OxmlElement("w:tcBorders")
            tc_pr.append(borders)
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement(f"w:{edge}")
            e.set(qn("w:val"), "nil")
            borders.append(e)
    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_path = EQ_DIR / f"eq{int(number):02d}.png" if number else None
    if eq_path and eq_path.exists():
        run = p.add_run()
        shape = run.add_picture(str(eq_path), width=Inches(5.75))
        shape._inline.docPr.set("title", f"Equation {number}")
        shape._inline.docPr.set("descr", f"LaTeX source: {latex}")
    else:
        run = p.add_run(latex)
        set_run_font(run, name="Cambria Math", size=10.3, italic=True, color=DARK)
    q = right.paragraphs[0]
    q.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if number:
        run = q.add_run(f"({number})")
        set_run_font(run, name="Cambria Math", size=10.0, color=DARK)
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    run = p.add_run(text)
    set_run_font(run, size=9.2, italic=True, color=MID)
    set_paragraph_keep(p)
    return p


def add_note_box(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT)
    set_cell_margins(cell, top=110, start=140, bottom=110, end=140)
    p = cell.paragraphs[0]
    run = p.add_run(title + "  ")
    set_run_font(run, bold=True, size=10.3, color=BLUE)
    run = p.add_run(text)
    set_run_font(run, size=10.0, color=DARK)
    return table


def add_table(doc, headers, rows, widths, caption=None, font_size=8.7):
    if caption:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(caption)
        set_run_font(run, size=9.5, bold=True, color=DARK)
        set_paragraph_keep(p, keep_next=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        shade_cell(cell, BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(header))
        set_run_font(run, size=font_size, bold=True, color=WHITE)
    for r_idx, row_data in enumerate(rows):
        row = table.add_row()
        if r_idx % 2 == 1:
            for cell in row.cells:
                shade_cell(cell, PALE)
        for idx, value in enumerate(row_data):
            p = row.cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=DARK)
    set_table_geometry(table, widths)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.68)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.30)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(5)

    for name, size, color in (
        ("Heading 1", 14, BLUE),
        ("Heading 2", 11.5, BLUE),
        ("Heading 3", 10.5, DARK),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 7)
        style.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("TUST MANUSCRIPT WORKING DRAFT  |  ARRIVAL-TIME QUEUE-AWARE ROUTING")
    set_run_font(run, size=8.2, bold=True, color=MID)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Working draft  •  ")
    set_run_font(run, size=8.0, color=MID)
    add_page_field(footer)


def title_page(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(36)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Arrival-time queue-aware evacuation routing for complex multi-line metro transfer stations")
    set_run_font(run, size=22, bold=True, color=BLUE)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("A network–microsimulation study")
    set_run_font(run, size=15, bold=True, color=DARK)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    run = p.add_run("Target journal")
    set_run_font(run, size=9, bold=True, color=BLUE)
    p.add_run("\n")
    run = p.add_run("Tunnelling and Underground Space Technology")
    set_run_font(run, size=11.5, bold=True, color=DARK)

    add_note_box(
        doc,
        "DRAFT STATUS",
        "Research positioning, manuscript flow, full methods and experimental protocol are drafted. "
        "High-load network and Pathfinder results are inserted. Low-load Pathfinder results remain reserved for later insertion; no value is inferred or pre-written.",
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    run = p.add_run("Scope locked for this version")
    set_run_font(run, size=10, bold=True, color=BLUE)
    for item in [
        "Emergency evacuation without fire; all modeled facilities remain available.",
        "Reference-demand scenario: 2,187 occupants.",
        "Train-arrival-augmented scenario: 17,905 occupants.",
        "Primary methods: Improved A* and arrival-time queue-aware A* (AA*).",
        "Pathfinder Any Exit retained only as a native locally-quickest contextual reference (PF-LQ).",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def executive_positioning(doc):
    add_heading(doc, "Manuscript positioning", 1)
    add_heading(doc, "Research purpose", 2)
    add_text_paragraph(
        doc,
        "This study investigates how emergency evacuation routes should be organised when passenger groups from several metro lines reach the same stairs, escalators, passages, fare-gate areas and exits at different times. The purpose is not to prove that one route-search algorithm is universally superior. It is to determine whether anticipating the queue that a group will encounter upon arrival can coordinate competing flows, reduce waiting exposure and accelerate the clearance of the evacuation tail under different demand levels, while explicitly reporting route length and computational cost.",
    )
    add_heading(doc, "Primary research question", 2)
    add_note_box(
        doc,
        "RQ",
        "How can evacuation routes be organised by anticipating the queues that passenger groups will encounter upon arrival at shared egress facilities in a complex multi-line metro transfer station under different demand levels?",
    )
    add_heading(doc, "Operational sub-questions", 2)
    for item in [
        "RQ1. How can a multi-line station be represented as a time-dependent network in which multiple passenger sources compete for shared facility service capacity?",
        "RQ2. How should current queues, committed future arrivals and service release before arrival be combined in the path cost used by a rolling route-search procedure?",
        "RQ3. Relative to an Improved A* reference, how does arrival-time queue awareness change completion-time quantiles, stationary waiting, bottleneck loading, route length and computation under reference and train-arrival-augmented demand?",
        "RQ4. When the prescribed route allocations are executed in Pathfinder, do the upper-tail and congestion patterns remain distinct from both the Improved A* allocation and Pathfinder’s native locally-quickest Any Exit behaviour?",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Contribution statement", 2)
    for item in [
        "A service-chain network representation that prevents multiple graph links from duplicating the capacity of the same physical facility.",
        "An arrival-time queue prediction that combines the visible queue with already committed future arrivals and subtracts service completed before the evaluated group reaches the facility.",
        "A time-dependent multi-label A* search coupled to a shared-capacity loading layer, so route evaluation and route execution refer to the same physical resources.",
        "A two-level evidence design: controlled comparison in the network model, followed by cross-model microscopic testing in Pathfinder using the same physical geometry and passenger scale.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Recommended paper flow", 2)
    flow_rows = [
        ("1. Introduction", "Shared facility queues, arrival-time mismatch and research gap", "RQ and contributions"),
        ("2. Literature background", "Station facility networks; dynamic loading and queues; evacuation guidance; microsimulation", "Position AA* without a catalogue-style review"),
        ("3. Methodology", "Station network → shared resources → queue prediction → AA* → common loading layer", "Reproducible model"),
        ("4. Case and experiments", "CAD geometry; demand construction; comparison matrix; Pathfinder protocol; metrics", "Traceable experiment"),
        ("5. Results", "Prediction mechanism → station performance → facility mechanism → Pathfinder cross-model results", "Evidence chain"),
        ("6. Discussion", "Why arrival-time coordination changes waiting and the tail; operational meaning; computation", "Interpretation"),
        ("7. Conclusions", "Answer RQ with low/high demand evidence", "Compact takeaways"),
    ]
    add_table(doc, ["Section", "Function", "Output"], flow_rows, [1.35, 3.35, 1.75], "Table 1. Evidence-led manuscript architecture", 8.5)
    doc.add_page_break()


def manuscript_text(doc):
    add_heading(doc, "Draft manuscript text", 1)
    add_heading(doc, "Abstract", 2)
    add_text_paragraph(
        doc,
        "Emergency evacuation in a multi-line metro transfer station is governed not only by walking distance but also by the time at which passenger groups reach shared stairs, passages, gate areas and exits. A route that appears uncongested at the decision time may already have substantial future demand committed to the same downstream facility. This study develops an arrival-time queue-aware A* method (AA*) that represents these facilities as shared-capacity resources and estimates the queue encountered at the expected arrival time from the visible queue, committed arrivals and intervening service. The method is evaluated in a five-line transfer-station model reconstructed from project CAD drawings under a 2,187-person reference-demand scenario and a 17,905-person train-arrival-augmented scenario. Improved A* is used as the primary reference in a common network-loading layer, and the resulting high-load route allocations are further executed in Pathfinder 2023.3. In the high-load network model, AA* reduced T95 and total clearance time by 21.0% and 17.8%, respectively, and reduced cumulative stationary exposure by 32.4%, while increasing travelled distance and computation. In Pathfinder, AA* reduced mean completion time by 9.1% and T100 by 8.2% relative to the prescribed Improved A* allocation. Pathfinder’s native locally-quickest Any Exit behaviour produced the lowest mean completion time but the longest T100, revealing a mean–tail trade-off rather than universal dominance. These findings show that arrival-time queue awareness is most consequential for coordinating shared bottlenecks and clearing the upper evacuation tail.",
    )
    p = doc.add_paragraph()
    run = p.add_run("Keywords: ")
    set_run_font(run, bold=True, size=10.2, color=BLUE)
    run = p.add_run("metro station evacuation; dynamic route guidance; arrival-time queue; shared bottleneck; Pathfinder; transfer station")
    set_run_font(run, size=10.2)

    add_heading(doc, "1. Introduction", 1)
    add_text_paragraph(
        doc,
        "Large underground transfer stations connect passenger streams from several lines within a vertically and horizontally constrained public space. During a station-wide evacuation, platform occupants, passengers released from trains, concourse occupants and cross-line transfer streams do not move through independent corridors. They successively share vertical circulation, passages, gate areas and exits. A local capacity shortfall can therefore propagate upstream and influence the clearance of passenger groups that originated far from the visible bottleneck. Facility-network studies have shown the value of mapping station functions and passenger flows into a directed, hierarchical system, while hub-loading research has demonstrated that node queuing and dynamic route choice must be represented together [1,2].",
    )
    add_text_paragraph(
        doc,
        "Most evacuation routing approaches recognise distance, current density, congestion, risk or multiple objectives, and microscopic simulation has become an established means of testing guided routes in detailed station geometry [3–6]. The remaining timing problem is more specific. A group does not encounter the state visible when its route is selected; it encounters the facility state after travelling through the preceding service chain. If several groups are assigned to the same facility in the meantime, a currently attractive route can become congested before the evaluated group arrives. Current-state impedance consequently understates future waiting and can repeatedly direct demand to a bottleneck whose service capacity is already committed.",
    )
    add_text_paragraph(
        doc,
        "This paper formulates the problem as arrival-time coordination of multi-source evacuation flows. Shared facilities are represented as resources with one conserved service capacity, regardless of how many network links refer to them. For each candidate path, the queue expected at an evaluated facility is advanced to the passenger group’s estimated arrival time using the current waiting load, confirmed in-transit arrivals and the service that can be completed before arrival. The predicted waiting time is then included in a time-dependent, multi-label A* search. The resulting method is termed arrival-time queue-aware A* (AA*).",
    )
    add_text_paragraph(
        doc,
        "The method is tested at Longyang Road station in Shanghai under emergency evacuation without fire and with all modeled facilities available. Two demand levels are retained: a 2,187-person reference-demand scenario derived from the passenger-number design record, and a 17,905-person scenario that augments the reference demand with passenger releases from arriving trains. The primary comparison is AA* versus an Improved A* reference implemented in the same network-loading layer. High-load allocations are additionally executed in Pathfinder. Pathfinder’s native Goto Any Exit behaviour is retained as a locally-quickest contextual reference because it answers a different but operationally relevant question: how a self-organised microscopic fastest-route rule trades average completion against the clearance tail [7–9].",
    )
    add_text_paragraph(
        doc,
        "The study contributes a shared-resource representation for a multi-line station, an arrival-time queue prediction linked to route search, and a cross-model experimental workflow that separates route-allocation effects from microscopic movement. The analysis is organised around the formation and release of shared bottlenecks rather than around a claim of universal algorithmic dominance.",
    )

    add_heading(doc, "2. Literature background", 1)
    add_heading(doc, "2.1 Station facility networks and dynamic passenger loading", 2)
    add_text_paragraph(
        doc,
        "Within-station movement has increasingly been represented as a network of functional spaces and facilities rather than as a single aggregate room. Shen et al. mapped station facilities and passenger flows into a directed hierarchical network and showed that flow redistribution changes network vulnerability [1]. Wen et al. coupled dynamic route choice and dynamic passenger loading, with instantaneous passage and service-facility travel times affected by congestion and queuing [2]. These studies support the network-and-service-chain representation adopted here. The present study shifts the focus from failure propagation or general passenger distribution to the route-decision consequence of demand that is already committed but has not yet reached a shared facility.",
    )
    add_heading(doc, "2.2 Evacuation routing and microscopic testing", 2)
    add_text_paragraph(
        doc,
        "Metro evacuation studies have combined path optimisation with detailed station simulation to evaluate time, density, congestion, risk and intervention measures [3–6]. Yang et al., in Tunnelling and Underground Space Technology, organised a complete chain from station-network failure and node travel-time prediction to multi-objective route optimisation and Pathfinder comparison [3]. Yang et al. also treated guided-passenger allocation and route planning as coupled decisions [4], while Guo and Zhang used simulation-based surrogate modelling to evaluate conflicting evacuation objectives [5]. These workflows motivate a two-stage experiment here: first test the route mechanism under a common network executor; then transfer the resulting route allocations into an independent microscopic movement environment.",
    )
    add_heading(doc, "2.3 Why retain Pathfinder Any Exit", 2)
    add_text_paragraph(
        doc,
        "The shortest route and the locally quickest route are not interchangeable. Hua et al. examined staircase choice by explicitly contrasting shortest and locally quickest alternatives and building a prediction-based choice model from full-scale experiments [7]. Pathfinder uses a locally quickest, hierarchical door-choice procedure that considers local queues and downstream travel knowledge [8]. Its default Goto Any Exit behaviour sends occupants to any available exit by the fastest route [9]. PF-LQ is therefore retained as a native behavioural reference, not as a third implementation of AA* and not as ground truth. Its benefit is interpretive: it reveals whether a self-organised microscopic rule that is strong on typical completion also controls the final clearance tail.",
    )

    add_heading(doc, "3. Methodology", 1)
    add_heading(doc, "3.1 Scope and station service-chain network", 2)
    add_text_paragraph(
        doc,
        "The study considers station-wide emergency evacuation without fire, smoke or facility outage. All modeled exits and internal circulation facilities remain available. The public circulation system is represented by a directed graph G=(V,E). Nodes denote platforms, concourses, transfer connectors, decision areas, stairs, escalators, fare-gate areas, pre-exit areas and safe exits. Directed edges represent feasible movement between adjacent spaces and contain length, effective width, direction and facility type. Passenger groups g∈𝒢 represent origins that share an initial location or train-release logic.",
    )
    add_text_paragraph(
        doc,
        "A physical facility can appear on several graph links. To avoid duplicating its capacity, a resource set ℛ and a mapping ρ:E→ℛ∪{∅} are defined. All links mapped to the same resource compete for one service rate μ_r. This representation makes flows from different lines compete at the physical facility level rather than within independent graph links.",
    )
    add_equation(doc, r"\rho:E\rightarrow\mathcal{R}\cup\{\varnothing\}", 1)

    add_heading(doc, "3.2 Common movement and loading layer", 2)
    add_text_paragraph(
        doc,
        "Both route methods use the same discrete loading layer with a 1-s time step, density-dependent walking, downstream receiving constraints, facility service queues and spillback. The movement-time relationship reproduces the Improved A* reference implementation adapted from Meng et al. [10]. The literature parameters are used as model inputs, not as field-calibrated measurements for Longyang Road station.",
    )
    add_equation(doc, r"v(k)=\begin{cases}1.427,&k\leq0.2\\\max(1.427-0.3549k,0),&0.2<k\leq4.0\\0,&k>4.0\end{cases}", 2)
    add_text_paragraph(
        doc,
        "The reference method uses α=0.15, β=0.85 and γ=0.10 in the current implementation. Its accumulated cost combines edge length and density-constrained travel time, while its heuristic is a scaled remaining distance. It observes current network conditions but does not include AA*’s future committed arrivals or arrival-time service waiting.",
    )
    add_equation(doc, r"c^{\mathrm{Imp}}_e(t)=\alpha l_e+\beta\,t^{\mathrm{move}}_e(t),\qquad h(n)=\gamma d(n,\mathcal{X})", 3)

    add_heading(doc, "3.3 Arrival-time queue prediction", 2)
    add_text_paragraph(
        doc,
        "Let Q_r(t) be the current waiting queue at shared resource r. Let A_r(t,τ) contain confirmed passenger batches already assigned and moving toward r that will arrive no later than τ. The queue expected when the evaluated batch arrives is the current queue plus those committed arrivals, advanced through the service completed over the intervening event sequence. In compact form:",
    )
    add_equation(doc, r"\widehat{Q}_r(\tau)=\max\left\{0,\;Q_r(t)+A_r(t,\tau)-\mu_r(\tau-t)\right\}", 4)
    add_text_paragraph(
        doc,
        "The implementation processes confirmed arrivals as time-ordered events, so service is applied between events before each arriving batch is added. The waiting contribution evaluated at the resource is:",
    )
    add_equation(doc, r"w_r(\tau)=\widehat{Q}_r(\tau)/\mu_r", 5)
    add_text_paragraph(
        doc,
        "For a candidate movement, AA* adds predicted queue waiting, the evaluated batch’s own mean gate-service contribution where applicable, physical movement time and any downstream spatial receiving wait. Arrival time is propagated along the path, making the edge cost time dependent.",
    )
    add_equation(doc, r"\tau_{j}=\tau_i+w_{\rho(e)}(\tau_i)+s^{\mathrm{batch}}_{\rho(e)}+t^{\mathrm{move}}_e(\tau_i)+w^{\mathrm{space}}_j(\tau_j)", 6)

    add_heading(doc, "3.4 Time-dependent AA* search", 2)
    add_text_paragraph(
        doc,
        "Because the downstream cost depends on arrival time, AA* does not retain only one static label per node. It uses a multi-label state (node, expected arrival time, accumulated cost) and removes a label only when another label reaches the same node no later and with no greater cost. A free-flow remaining-time lower bound guides the search. Routes are recalculated within the rolling loading process, and confirmed movements are recorded as future resource-arrival events before subsequent groups are evaluated.",
    )
    add_equation(doc, r"L_1\prec L_2\iff \tau_1\leq\tau_2\;\land\;g_1\leq g_2", 7)
    add_text_paragraph(
        doc,
        "This search-execution coupling is central to the method: a route decision changes the committed-arrival state that later decisions evaluate, and all accepted moves consume the same physical service and receiving capacity used by the reference method.",
    )

    add_heading(doc, "4. Case study and experimental design", 1)
    add_heading(doc, "4.1 Case station and geometry", 2)
    add_text_paragraph(
        doc,
        "Longyang Road station is modelled as a multi-line complex connecting Metro Lines 2, 7, 16 and 18 and the Maglev. Walkable spaces, floor elevations, vertical circulation, passages, gate areas and exits were reconstructed from the project CAD drawing supplied for the study. The graph model abstracts these spaces into service-chain nodes and links; the Pathfinder model preserves the navigable geometry used for microscopic motion. Facility dimensions taken from CAD and model configuration will be tabulated separately from literature parameters and algorithmic settings in the submission version.",
    )

    add_heading(doc, "4.2 Demand construction", 2)
    add_text_paragraph(
        doc,
        "The 2,187-person reference demand follows the passenger-number design record supplied for the project, which assigns passengers to platform, concourse and transfer spaces by line. The value is presented as a traceable scenario input rather than being labelled by a calendar-day category. The train-arrival-augmented scenario adds the passenger loads of the modeled arriving trains to the same reference distribution. The added loads are read directly from the frozen experiment code. Passenger-flow assignment research based on Shanghai Metro AFC data is used to support the general logic of reconstructing station and platform flows at fine temporal resolution [11]; it is not used as the source of the project-specific 2,187-person total.",
    )
    demand_rows = [
        ("Line 2", "1,112", "2 × 2,400", "5,912"),
        ("Line 7", "500", "2 × 1,620", "3,740"),
        ("Line 16", "84", "2 × 1,230", "2,544"),
        ("Line 18", "491", "2 × 1,650", "3,791"),
        ("Maglev", "0", "2 × 959", "1,918"),
        ("Total", "2,187", "15,718", "17,905"),
    ]
    add_table(doc, ["Source", "Reference", "Added train arrivals", "Augmented total"], demand_rows, [1.25, 1.25, 2.05, 1.45], "Table 2. Implemented demand construction (persons)", 8.6)

    add_heading(doc, "4.3 Experiment matrix", 2)
    matrix_rows = [
        ("Network model", "Reference demand", "Improved A*; AA*", "Mechanism and low-load performance", "Results pending insertion"),
        ("Network model", "Train-arrival augmented", "Improved A*; AA*", "Primary high-load comparison", "Complete"),
        ("Pathfinder", "Reference demand", "P-Improved; P-AA; PF-LQ", "Cross-model low-load pattern", "Run/processing pending"),
        ("Pathfinder", "Train-arrival augmented", "P-Improved; P-AA; PF-LQ", "Cross-model high-load pattern", "Complete"),
    ]
    add_table(doc, ["Executor", "Demand", "Route protocol", "Purpose", "Status"], matrix_rows, [1.05, 1.35, 1.6, 2.0, 0.95], "Table 3. Two-demand, two-executor experiment matrix", 8.2)

    add_heading(doc, "4.4 Network-model evaluation", 2)
    add_text_paragraph(
        doc,
        "The network experiment changes only the route method. Geometry, demand, density–speed relationship, facility capacities, queue release, receiving constraints, spillback rules and termination criteria remain common. Conservation is checked by requiring evacuated plus remaining passengers to equal the input population at every run and by confirming complete evacuation. Primary outcomes are T95, T100 and cumulative stationary person-seconds. Mechanism outcomes are key-facility queue exposure, exit and facility load balance, line-specific clearance and route-distance change. Wall-clock time reports the computational price of arrival-time prediction.",
    )
    add_equation(doc, r"T_p=\min\{t:E^{\mathrm{out}}(t)\geq pN\},\qquad p\in\{0.50,0.80,0.90,0.95,0.99,1.00\}", 8)
    add_equation(doc, r"W^{\mathrm{stat}}=\sum_t N^{\mathrm{stationary}}(t)\,\Delta t", 9)

    add_heading(doc, "4.5 Pathfinder cross-model microscopic protocol", 2)
    add_text_paragraph(
        doc,
        "Pathfinder 2023.3.1206 is used in Steering mode. The high-load test contains 17,905 occupants in each protocol and uses byte-identical .geom files (identical SHA-256 digest), ensuring the same physical geometry. Three route protocols are evaluated: P-Improved, in which exported Improved A* route/exit allocations are prescribed; P-AA, in which exported AA* allocations are prescribed; and PF-LQ, in which occupants use Pathfinder’s native Goto Any Exit locally-quickest behaviour. P-Improved and P-AA form the primary cross-model route-allocation comparison. PF-LQ is a contextual self-organised reference because its route choice is generated within Pathfinder rather than by the network algorithms.",
    )
    protocol_rows = [
        ("P-Improved", "Prescribed", "Improved A* exported assignment", "Primary reference"),
        ("P-AA", "Prescribed", "AA* exported assignment", "Tests transfer of proposed allocation"),
        ("PF-LQ", "Self-organised", "Pathfinder Goto Any Exit / locally quickest", "Contextual mean–tail reference"),
    ]
    add_table(doc, ["Protocol", "Path control", "Origin of route choice", "Analytical role"], protocol_rows, [1.15, 1.25, 2.55, 1.65], "Table 4. Pathfinder route protocols", 8.5)
    add_text_paragraph(
        doc,
        "The audit confirms equal geometry, equal population and equal sets of output occupant ids across the three runs. P-Improved and P-AA share the same 17,905 occupant names and differ principally in assigned behaviours; initial orientation differs for 5,685 occupants, alongside seven coordinate pairs and one room/elevation record, and these differences are retained in the audit file. PF-LQ was generated as a separate native scenario and re-sampled many individual positions and body parameters. Consequently, P-AA versus P-Improved supports a matched-name descriptive comparison, whereas PF-LQ is interpreted at the scenario-distribution level. This is why PF-LQ is not used as ground truth or included in paired individual tests.",
    )
    add_text_paragraph(
        doc,
        "Each run exports occupant completion time, congestion time, level and stair congestion, and travelled distance. The empirical cumulative completion curve and T50–T100 quantiles are calculated from all 17,905 occupants without sampling. For P-AA versus P-Improved, occupant names provide a paired descriptive difference. Because the current evidence consists of a fixed high-load realisation, the paper reports complete distributions and effect magnitudes rather than p-values across simulated occupants.",
    )

    add_heading(doc, "5. Results available for the current draft", 1)
    add_heading(doc, "5.1 High-load network-model comparison", 2)
    add_text_paragraph(
        doc,
        "Both network runs evacuated all 17,905 passengers. Relative to Improved A*, AA* reduced T95 from 1,125 to 889 s and T100 from 1,486 to 1,222 s. Cumulative stationary exposure decreased from 6.553×10^6 to 4.431×10^6 person-s, and mean evacuation time decreased from 438.9 to 335.5 s. The result is consistent with the proposed mechanism: waiting was reduced more strongly than movement. Mean stationary time decreased by 32.4%, whereas total travelled distance increased by 15.5% and mean moving time increased by 20.5%. AA* therefore improved completion by accepting longer movement where it avoided more persistent shared-facility waiting.",
    )
    network_rows = [
        ("T95 (s)", "1,125", "889", "−21.0%"),
        ("T100 (s)", "1,486", "1,222", "−17.8%"),
        ("Mean evacuation time (s/person)", "438.85", "335.54", "−23.5%"),
        ("Stationary exposure (person-s)", "6,552,858", "4,431,142", "−32.4%"),
        ("Mean stationary time (s/person)", "365.98", "247.48", "−32.4%"),
        ("Total distance (m)", "1,481,121", "1,710,553", "+15.5%"),
        ("Exit-load Jain index", "0.626", "0.711", "+13.6%"),
        ("Key-facility Jain index", "0.162", "0.409", "+152.4%"),
        ("Wall-clock runtime (s)", "37.65", "615.28", "16.3×"),
    ]
    add_table(doc, ["Metric", "Improved A*", "AA*", "AA* change"], network_rows, [2.65, 1.25, 1.25, 1.15], "Table 5. Train-arrival-augmented network-model results", 8.5)
    add_text_paragraph(
        doc,
        "The redistribution was not uniform across source lines. AA* accelerated the clearance of Lines 2, 7, 16 and 18, with the largest absolute change for Line 18 (1,309 to 418 s), while Maglev clearance increased from 590 to 651 s. Line 7 remained the final clearing line in both methods. The station-wide gain therefore arose from redistributing shared capacity across sources, rather than improving every source in isolation.",
    )

    add_heading(doc, "5.2 High-load Pathfinder comparison", 2)
    pf = pd.read_csv(PF_SUMMARY_PATH, encoding="utf-8-sig")
    pf_rows = []
    for _, row in pf.iterrows():
        pf_rows.append(
            (
                row["method"],
                f"{row['mean_exit_time_s']:.1f}",
                f"{row['T50_s']:.1f}",
                f"{row['T95_s']:.1f}",
                f"{row['T99_s']:.1f}",
                f"{row['T100_s']:.1f}",
                f"{row['mean_congestion_time_s']:.1f}",
                f"{row['mean_distance_m']:.1f}",
            )
        )
    add_table(
        doc,
        ["Protocol", "Mean", "T50", "T95", "T99", "T100", "Congestion", "Distance"],
        pf_rows,
        [1.05, 0.72, 0.65, 0.72, 0.72, 0.75, 0.9, 0.8],
        "Table 6. Pathfinder completion, congestion and distance results (seconds except distance in metres)",
        7.7,
    )
    add_text_paragraph(
        doc,
        "P-AA reduced mean completion time from 435.5 to 396.0 s (9.1%) and T100 from 1,414.6 to 1,298.3 s (8.2%) relative to P-Improved. It also reduced mean congestion time from 311.3 to 276.6 s. Across the 17,905 occupants matched by name, P-AA was faster for 10,073 (56.3%), slower for 7,825 and equal for seven; the mean and median time saved were 39.5 and 12.2 s, respectively. These distributions show that the station-level gain is produced by heterogeneous individual changes rather than by uniformly shifting every trajectory.",
    )
    add_text_paragraph(
        doc,
        "PF-LQ produced the lowest mean completion time (362.5 s), T50 (292.8 s) and T90 (702.7 s), but its curve crossed the P-AA curve in the upper tail. PF-LQ reached T95 at 1,034.1 s and completed at 1,458.5 s, compared with 985.1 and 1,298.3 s for P-AA. In other words, Pathfinder’s native locally-quickest response served the typical occupant efficiently but left a longer residual tail. P-AA occupied a distinct operating point: it did not minimise the mean across all three protocols, but it delivered the earliest high-quantile and final clearance.",
    )
    if FIG_PATH.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.keep_with_next = True
        p.add_run().add_picture(str(FIG_PATH), width=Inches(6.65))
        add_caption(
            doc,
            "Fig. 2. Cross-model microscopic results under train-arrival-augmented demand. (a) Empirical completion curves for all occupants. (b) Median-to-tail quantiles. (c) Mean–T100 trade-off. (d) Mean travelled distance and congestion time. P-Improved and P-AA use prescribed exported allocations; PF-LQ uses Pathfinder Goto Any Exit.",
        )

    add_heading(doc, "5.3 Integrated interpretation", 2)
    add_text_paragraph(
        doc,
        "The two executors support the same primary route-allocation comparison. In the network model, AA* shortened T100 by 17.8% relative to Improved A*; after the allocations were transferred to Pathfinder, P-AA shortened T100 by 8.2% relative to P-Improved. The magnitude changed because the network and microscopic executors resolve motion, local conflict and spatial occupancy differently, but the direction of the prescribed-allocation contrast remained the same. The Pathfinder result also refined the claim: AA* should be presented as a method for coordinating bottlenecks and suppressing the evacuation tail, not as the fastest rule for every occupant or every summary statistic.",
    )

    add_heading(doc, "5.4 Low-load insertion contract", 2)
    add_note_box(
        doc,
        "TO BE INSERTED AFTER THE RUNS ARE FROZEN",
        "The 2,187-person Pathfinder section will use the identical three-protocol table and completion figure. It will report observed values only. The discussion should test whether lower overlap in facility-arrival times reduces the practical difference between AA* and Improved A*, rather than assuming in advance that the direction or magnitude must match the high-load case.",
    )

    add_heading(doc, "6. Discussion draft", 1)
    add_heading(doc, "6.1 Arrival-time coordination rather than shortest-path improvement", 2)
    add_text_paragraph(
        doc,
        "The results distinguish the proposed contribution from a conventional path-length improvement. AA* increased network-model movement distance while substantially reducing stationary exposure, and its Pathfinder distance was almost unchanged relative to P-Improved. The decisive quantity was therefore not geometric length alone but the service condition encountered at shared facilities. By recording already committed arrivals, AA* prevented a temporarily uncongested facility from being repeatedly treated as unclaimed capacity.",
    )
    add_heading(doc, "6.2 The mean–tail distinction", 2)
    add_text_paragraph(
        doc,
        "PF-LQ demonstrates why mean completion and final clearance must be reported together. A locally quickest rule can evacuate most occupants rapidly because it reacts to nearby door queues and downstream travel, yet the resulting self-organisation can leave a smaller set of occupants on long or persistently congested residual routes. P-AA instead imposed a system-level allocation generated from predicted shared-resource commitments. Its advantage appeared most clearly at T95–T100. For emergency management, this distinction separates rapid service for the majority from the ability to release the final occupied branches of a large station.",
    )
    add_heading(doc, "6.3 Operational interpretation and computation", 2)
    add_text_paragraph(
        doc,
        "The route output should be interpreted as an information and guidance policy for passenger groups, not as individual turn-by-turn optimisation. Candidate applications include dynamic signs, staff instructions and line- or platform-specific exit allocation. The high-load network run required approximately 16.3 times the runtime of Improved A*, reflecting event-indexed queue prediction and multi-label search. This cost is reported alongside the evacuation benefit because deployment depends on the update interval available to the station control system.",
    )

    add_heading(doc, "7. Conclusions draft", 1)
    add_text_paragraph(
        doc,
        "This study frames complex transfer-station evacuation as an arrival-time coordination problem over shared facility capacities. The proposed AA* method predicts the queue a passenger group will encounter when it reaches a facility and embeds that wait in a time-dependent search coupled to common capacity execution. Under train-arrival-augmented demand, the method shortened the upper-tail and final clearance times, reduced stationary waiting and distributed load more evenly in the network model, at the price of higher computation and additional movement. The Pathfinder experiment preserved the prescribed AA* advantage over prescribed Improved A* in both mean and final clearance. Pathfinder Any Exit was faster on average but slower in the upper tail, confirming that average and tail performance describe different routing outcomes. The final paper will use the completed 2,187-person simulations to determine how strongly this mechanism persists when shared-facility competition is lower.",
    )


def add_reference_section(doc):
    add_heading(doc, "References used in this working draft", 1)
    refs = [
        "[1] Shen, Y., Yang, H., Ren, G., Ran, B. (2024). Model cascading overload failure and dynamic vulnerability analysis of facility network of metro station. Reliability Engineering & System Safety 242, 109711. https://doi.org/10.1016/j.ress.2023.109711.",
        "[2] Wen, X., Si, B., Xu, M., Zhao, F., Jiang, R. (2024). A passenger flow spatial–temporal distribution model for a passenger transit hub considering node queuing. Transportation Research Part C 163, 104640. https://doi.org/10.1016/j.trc.2024.104640.",
        "[3] Yang, X., Dai, W., Li, Y., Yang, X. (2024). An efficient evacuation path optimization for passengers in subway stations under floods. Tunnelling and Underground Space Technology 144, 105473. https://doi.org/10.1016/j.tust.2023.105473.",
        "[4] Yang, X., Yang, Y., Li, Y., Yang, X. (2022). Path planning for guided passengers during evacuation in subway station based on multi-objective optimization. Applied Mathematical Modelling 111, 777–801. https://doi.org/10.1016/j.apm.2022.07.024.",
        "[5] Guo, K., Zhang, L. (2022). Simulation-based passenger evacuation optimization in metro stations considering multi-objectives. Automation in Construction 133, 104010. https://doi.org/10.1016/j.autcon.2021.104010.",
        "[6] Xu, H., Wei, Y., Tan, Y. (2024). Optimization of emergency evacuation in complex rail transit station. Journal of Building Engineering 98, 110321. https://doi.org/10.1016/j.jobe.2024.110321.",
        "[7] Hua, Y., Zhao, J., Li, H.-T., Duan, L. (2024). Shortest or locally quickest? A prediction-based approach for evacuation choice simulation between multiple staircases. Journal of Safety Science and Resilience 5, 281–294. https://doi.org/10.1016/j.jnlssr.2024.04.001.",
        "[8] Thunderhead Engineering Consultants, Inc. (2021). Pathfinder 2021.3 Technical Reference Manual: Path planning and locally quickest door choice. https://support.thunderheadeng.com/docs/pathfinder/2021-3/technical-reference-manual/.",
        "[9] Thunderhead Engineering Consultants, Inc. (2022). Pathfinder 2022.3 User Manual: Behaviors and Goto Any Exit. https://support.thunderheadeng.com/docs/pathfinder/2022-3/user-manual/.",
        "[10] Meng, D., Hu, Z., Zhang, H. (2022). Simulation of a multi-layer cruise-ship evacuation system based on an improved A* algorithm. Journal of System Simulation 34(6), 1375–1382. https://doi.org/10.16182/j.issn1004731x.joss.21-0075. (in Chinese).",
        "[11] Yu, L., Liu, H., Fang, Z., Ye, R., Huang, Z., You, Y. (2023). A new approach on passenger flow assignment with multi-connected agents. Physica A 628, 129175. https://doi.org/10.1016/j.physa.2023.129175.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(ref)
        set_run_font(run, size=8.9, color=DARK)

    add_heading(doc, "Author-facing source and output audit", 1)
    audit_rows = [
        ("Station geometry", "Project DWG supplied by author", "Model input; geometry statement"),
        ("Reference demand", "Author-supplied passenger-number design record", "2,187-person source distribution"),
        ("Train augmentation", "Frozen experiment code", "Line-specific train loads; total 17,905"),
        ("Improved A*", "Full local PDF of Meng et al. (2022) + implementation", "Reference parameters and method adaptation"),
        ("AA*", "Frozen local source code", "Equations and algorithm workflow"),
        ("Network results", "mode4 formal report", "Table 5 and Results 5.1"),
        ("Pathfinder results", "Three occupant CSV exports", "Table 6 and Fig. 2"),
        ("Pathfinder equivalence", "SHA-256 and parameter audit", "Protocol wording in Section 4.5"),
        ("Literature", "Full-text Zotero/PDF corpus and publisher full text", "Introduction, method context, figure logic"),
    ]
    add_table(doc, ["Component", "Direct source", "Use in manuscript"], audit_rows, [1.45, 2.55, 2.55], "Table A1. Claim-to-source audit for this draft", 8.3)


def main():
    if not FIG_PATH.exists() or not PF_SUMMARY_PATH.exists():
        raise FileNotFoundError("Run build_pathfinder_high_load_figure.py before building the manuscript.")
    doc = Document()
    configure_document(doc)
    title_page(doc)
    executive_positioning(doc)
    manuscript_text(doc)
    add_reference_section(doc)
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
