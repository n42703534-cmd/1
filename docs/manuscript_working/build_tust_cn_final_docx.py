from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
WORK = Path(__file__).resolve().parent
FIG = WORK / "figures_tust_cn_final"
OUT = WORK / "45_tust_manuscript_AA_load_conditioned_cn.docx"

BLACK = RGBColor(25, 30, 35)
GREY = RGBColor(92, 97, 101)
BLUE = RGBColor(60, 112, 150)
LIGHT = "EAF0F4"


def set_font(run, east="宋体", latin="Times New Roman", size=10.5,
             bold=False, italic=False, color=BLACK):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), east)
    rpr.rFonts.set(qn("w:ascii"), latin)
    rpr.rFonts.set(qn("w:hAnsi"), latin)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def configure(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.15)
    sec.right_margin = Cm(2.15)
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Cm(0.74)
    normal.paragraph_format.line_spacing = 1.32
    normal.paragraph_format.space_after = Pt(5)

    for name, size, before, after in [
        ("Heading 1", 14, 11, 5),
        ("Heading 2", 12, 8, 3),
        ("Heading 3", 11, 6, 2),
    ]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.append(fld)


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("面向多线换乘站共享瓶颈协调的到达时刻队列感知疏散路径规划")
    set_font(r, east="黑体", size=17, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(9)
    r = p.add_run("中文研究稿")
    set_font(r, east="宋体", size=9.5, color=GREY)


def add_p(doc, text, *, indent=True, size=10.5, color=BLACK, italic=False,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.32
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.first_line_indent = Cm(0.74) if indent else None
    r = p.add_run(text)
    set_font(r, size=size, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def _omml_run(text):
    r = OxmlElement("m:r")
    t = OxmlElement("m:t")
    t.text = text
    r.append(t)
    return r


def _omml_sub(base, sub):
    node = OxmlElement("m:sSub")
    node.append(OxmlElement("m:sSubPr"))
    e = OxmlElement("m:e"); e.append(_omml_run(base))
    s = OxmlElement("m:sub"); s.append(_omml_run(sub))
    node.extend([e, s])
    return node


def _omml_sup(base, sup):
    node = OxmlElement("m:sSup")
    node.append(OxmlElement("m:sSupPr"))
    e = OxmlElement("m:e"); e.append(_omml_run(base))
    s = OxmlElement("m:sup"); s.append(_omml_run(sup))
    node.extend([e, s])
    return node


def _omml_frac(num, den):
    node = OxmlElement("m:f")
    node.append(OxmlElement("m:fPr"))
    n = OxmlElement("m:num"); n.append(_omml_run(num))
    d = OxmlElement("m:den"); d.append(_omml_run(den))
    node.extend([n, d])
    return node


def add_equation(doc, number, kind):
    """Insert editable OMML. The accompanying source is retained in the caption text."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    math = OxmlElement("m:oMath")
    if kind == "resource":
        math.extend([_omml_run("r(u,v)="), _omml_run("{")])
        math.extend([_omml_run("(facility,v),  v∈{stair,escalator,gate}"), _omml_run(";  ")])
        math.extend([_omml_run("(edge,u,v),  otherwise")])
    elif kind == "queue_minus":
        math.extend([_omml_sub("Q", "k"), _omml_run("⁻=max["), _omml_sub("Q", "k−1"),
                     _omml_run("⁺−μ"), _omml_sub("r", "("), _omml_run("a"), _omml_run("k−a"),
                     _omml_run("k−1"), _omml_run("),0]")])
    elif kind == "queue_plus":
        math.extend([_omml_sub("Q", "k"), _omml_run("⁺="), _omml_sub("Q", "k"),
                     _omml_run("⁻+m"), _omml_sub("k", "")])
    elif kind == "queue_target":
        math.extend([_omml_sub("Q̂", "r"), _omml_run("(τ)=max["), _omml_run("Q"),
                     _omml_run("⁺−μ(τ−a"), _omml_run("K"), _omml_run("),0]")])
    elif kind == "batch":
        math.extend([_omml_sub("b", "r"), _omml_run("(m)="), _omml_frac("m", "2μ"),
                     _omml_run("  for facility resources;  0  otherwise")])
    elif kind == "arrival":
        math.extend([_omml_sub("τ", "j"), _omml_run("="), _omml_sub("τ", "i"),
                     _omml_run("+t"), _omml_sub("ij", "move"), _omml_run("+"),
                     _omml_sub("Q̂", "r"), _omml_run("(τ"), _omml_sub("i", ""),
                     _omml_run(")/μ"), _omml_sub("r", ""), _omml_run("+b"),
                     _omml_sub("r", ""), _omml_run("(m)+s"), _omml_sub("j", "space")])
    elif kind == "risk":
        math.extend([_omml_sub("R", "j"), _omml_run("="), _omml_sub("R", "i"),
                     _omml_run("+R"), _omml_sub("ij", "wait"), _omml_run("+R"),
                     _omml_sub("ij", "move")])
    elif kind == "objective":
        math.extend([_omml_sub("C", "j"), _omml_run("=(τ"), _omml_sub("j", ""),
                     _omml_run("−τ"), _omml_sub("0", ""), _omml_run(")+λR"),
                     _omml_sub("j", "")])
    elif kind == "tp":
        math.extend([_omml_sub("T", "p"), _omml_run("=min{t:E"), _omml_sub("out", ""),
                     _omml_run("(t)≥pN}")])
    elif kind == "wstat":
        math.extend([_omml_run("W"), _omml_sub("stat", ""), _omml_run("=Σ"),
                     _omml_sub("t", ""), _omml_run("N"), _omml_sub("stationary", ""),
                     _omml_run("(t)Δt")])
    elif kind == "jain":
        math.extend([_omml_run("J(x)="), _omml_frac("(Σxᵢ)²", "nΣxᵢ²")])
    elif kind == "ecdf":
        math.extend([_omml_run("F̂(t)=N⁻¹Σ"), _omml_sub("i=1", "N"),
                     _omml_run("1(Tᵢ≤t)")])
    else:
        math.append(_omml_run(kind))
    p._p.append(math)
    n = p.add_run(f"    ({number})")
    set_font(n, east="Times New Roman", latin="Times New Roman", size=9.5)
    return p


def add_latex_note(doc, text):
    # Keep LaTeX source in the builder script; the manuscript displays editable
    # Word OMML to avoid exposing raw source code in the submitted DOCX.
    return None


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tag = OxmlElement("w:tblHeader")
    tag.set(qn("w:val"), "true")
    tr_pr.append(tag)


def set_cell_margins(cell, top=70, start=110, bottom=70, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn("w:" + name))
        if node is None:
            node = OxmlElement("w:" + name)
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def add_table(doc, title, headers, rows, widths=None, font_size=8.1):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, size=9.3, bold=True)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for cell, width in zip(table.rows[0].cells, widths):
            cell.width = Cm(width)
    set_repeat_header(table.rows[0])
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), LIGHT)
        cell._tc.get_or_add_tcPr().append(shd)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_font(run, size=font_size, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for run in cells[i].paragraphs[0].runs:
                set_font(run, size=font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_figure(doc, filename, caption, width_cm=16.2):
    path = FIG / filename
    if not path.exists():
        path = WORK / "figures_revised" / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    shape = run.add_picture(str(path), width=Cm(width_cm))
    shape._inline.docPr.set("title", filename)
    shape._inline.docPr.set("descr", caption)
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    c.paragraph_format.first_line_indent = None
    c.paragraph_format.space_after = Pt(6)
    r = c.add_run(caption)
    set_font(r, size=8.5)


def add_code_block(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    set_font(r, east="等线", latin="Consolas", size=8.7, color=BLACK)
    return p


def build():
    doc = Document()
    configure(doc)
    add_title(doc)

    add_heading(doc, "摘要", 1)
    add_p(doc, "多线换乘站疏散中的共享楼梯、自动扶梯、闸机和通道构成容量受限的服务链。若路径规划只读取决策时刻的设施状态，已经分配但尚未到达瓶颈的客流不会进入当前判断，因而可能低估候选路线在未来时刻面对的竞争。本文以龙阳路多线换乘站为对象，提出到达时刻队列感知 A*（arrival-time queue-aware A*，AA*）。AA* 将楼梯、自动扶梯和闸机表示为设施资源，将普通移动边表示为边资源；对当前物理队列和已接受的未来到达事件按时间排序，在候选批次预计到达资源时推进服务过程，并在时变多标签搜索中累计物理移动、资源等待、设施批次服务、空间接纳等待和密度暴露。研究比较 2,187 人基准需求与在其上叠加列车到站客流形成的 17,905 人高负荷需求，并在共同网络执行层、单模块消融和 Pathfinder 连续空间复现中评估 AA*。高负荷下，AA* 将网络层平均完成时间从 438.9 s 降至 335.5 s，T95 从 1125 s 降至 889 s，T100 从 1486 s 降至 1222 s，累计静止暴露降低 32.4%，但总移动距离增加 15.5%，墙钟运行时间增加至 615.3 s。低负荷下，AA* 将 T95 从 289 s 降至 258 s，累计静止暴露降低 15.1%，但 T100 保持 325 s。Pathfinder 复现显示，高负荷下 AA* 的平均完成时间和 T100 改善方向得以保持；低负荷下两种预分配路径的 T95 和 T100 相同，AA* 平均完成时间略高。结果表明，AA* 的收益受负荷条件调节：其主要作用是协调共享瓶颈的未来到达竞争，而不是在所有场景中缩短几何路径。", indent=True)
    add_p(doc, "关键词：地铁站疏散；到达时刻队列；共享瓶颈；动态路径规划；Pathfinder；多线换乘站", indent=False)

    add_heading(doc, "1 引言", 1)
    add_p(doc, "多线换乘站在有限地下空间内连接多个站台、站厅、换乘通道和地面出口。疏散开始后，站台候车客流、站厅客流、换乘客流以及列车到站客流并不是相互独立地离开，而是在楼梯、自动扶梯、闸机和出口前区等设施上产生时空竞争。局部服务能力不足会沿客流链向上游传播，使距离瓶颈较远的批次也受到影响。地铁设施网络、节点排队和客流分配研究已经说明，站内疏散需要同时考虑空间连通关系与设施服务过程[1,2]。")
    add_p(doc, "现有路径规划通常在某一决策时刻读取密度、距离或当前队列，并据此比较候选路线。该做法在客流较少或瓶颈竞争弱时可能足够，但在多批次到达共享设施的场景中存在一个状态错位：乘客在时刻 t₀ 做出路径决策，而真正进入楼梯、闸机或出口前区的时刻是未来的 τ。t₀ 与 τ 之间，设施会继续服务当前队列，也会接收已经分配但尚未到达的其他批次。如果这些未来到达事件未被纳入路径代价，当前看似较短的路线可能在候选批次到达时形成新的拥堵。")
    add_p(doc, "地铁站疏散研究已经形成了“真实站体—计算网络—微观仿真”的常见证据链。Guo 和 Zhang 将仿真场景、代理评价和多目标优化连接起来[3]；Yang 等将引导客流路径、出口分配和拥堵评价结合起来[4]；其后续研究进一步将站体网络、时间预测、多目标路径与 Pathfinder 复现连接起来[5]。结伴行为、道路中断和洪水风险等因素也被用于扩展疏散路径模型[6]。这些研究为跨模型评估提供了基础，但本文关注的不是火灾、洪水或结伴行为，而是无火灾、设施均可用条件下，多线换乘站共享瓶颈中的未来到达竞争。")
    add_p(doc, "本文提出到达时刻队列感知 A*（AA*），围绕以下研究问题展开：RQ1，在低负荷和高负荷条件下，未来到达客流是否会改变共享瓶颈的路径分配与清空过程？RQ2，将当前队列、已接受未来到达事件和服务过程推进到候选批次的预计到达时刻，能否减少静止暴露并压缩尾部清空？RQ3，网络层的路径改善能否在 Pathfinder 连续空间中保持，其表现是否随负荷条件变化？")
    add_p(doc, "本文的贡献有三点。第一，给出与代码实际实现一致的共享资源表示：目的节点为楼梯、自动扶梯或闸机时使用设施资源，其他移动使用边资源。第二，将到达事件队列、设施服务和空间接纳等待嵌入时变多标签搜索，并明确区分物理状态、已接受承诺和规划意图。第三，建立共同网络执行层、单模块消融和 Pathfinder 连续空间复现组成的证据链，报告平均效率、尾部清空、静止或拥堵暴露、绕行距离和计算代价，而不预设 AA* 在所有条件下全面优于基线。")

    add_heading(doc, "2 站体、网络表示与需求场景", 1)
    add_heading(doc, "2.1 龙阳路站体与网络对象", 2)
    add_p(doc, "研究对象为无火灾、无烟气且建模设施均可用条件下的龙阳路多线换乘站。依据 CAD 几何和 Pathfinder 建模对象，站体被抽象为有向设施—通道网络 G=(V,E)。节点包括站台与候车区、站厅、换乘通道、楼梯、自动扶梯、闸机及其排队区、出口前区和最终安全出口；有向边表示相邻空间或设施之间的可行移动，并保存长度、容量、宽度、障碍和运行状态等属性。当前网络包含 572 个节点、1,414 条有向边和 16 个最终出口。图 1 不再直接展示难以辨识的全节点云图，而是展示 CAD、网络对象和 Pathfinder 连续空间之间的层级对应。")
    add_figure(doc, "fig1_station_structure_cn.png", "图 1  龙阳路多线换乘站的空间结构与模型对象。左侧依据 CAD 与 Pathfinder 建模对象抽象楼层—线路—换乘关系；右侧说明活动区域、竖向设施和最终出口在连续空间模型与网络模型中的对应角色。该图用于说明对象类型和连通关系，不按比例表示实际几何。")
    add_p(doc, "资源队列与普通空间占用分开记录。对有向边 (u,v)，代码采用以下资源映射：当目的节点 v 为楼梯、自动扶梯或闸机节点时，边消耗该设施节点对应的服务资源；其他有向边使用自身的边资源。由此，同一楼梯或闸机连接多条边时，所有相关进入动作共享一个服务能力，而普通通道边保持独立容量。资源服务率由设施节点容量或边容量给出；资源队列不等同于目的节点的几何占用人数。")
    add_equation(doc, 1, "resource")
    add_latex_note(doc, r"r(u,v)=\{(facility,v),\ v\in\{stair,escalator,gate\};\ (edge,u,v),\ otherwise\}")
    add_p(doc, "网络模型和 Pathfinder 模型承担不同角色。网络模型负责路径选择、批次规模和出口分配；Pathfinder 在相同站体几何中执行预分配路线或软件原生行为。Pathfinder 的 Goto Any Exit 作为原生自主选出口参照，不作为 AA* 的第三种实现，也不被解释为经验真值。")

    add_heading(doc, "2.2 需求场景", 2)
    add_p(doc, "2,187 人基准需求来自项目实施阶段冻结的人数设计记录，按线路站台候车、站厅和换乘空间输入网络模型。高负荷需求在同一基准需求上叠加代码中两列列车的到站客流，总人数为 17,905 人。本文使用“低负荷”和“高负荷”描述输入规模，不将 2,187 人强行命名为节假日、周末或工作日场景。")
    add_table(doc, "表 1  两种需求场景的线路构成（人）", ["线路", "基准需求", "每列列车", "叠加列车数", "高负荷总量"], [
        ("2号线", "1,112", "2,400", "2", "5,912"),
        ("7号线", "500", "1,620", "2", "3,740"),
        ("16号线", "84", "1,230", "2", "2,544"),
        ("18号线", "491", "1,650", "2", "3,791"),
        ("磁浮", "0", "959", "2", "1,918"),
        ("合计", "2,187", "—", "—", "17,905"),
    ], [2.6, 3.0, 3.0, 2.6, 3.1])

    add_heading(doc, "3 到达时刻队列感知 A*", 1)
    add_heading(doc, "3.1 共同物理执行层与基线", 2)
    add_p(doc, "Improved A* 作为已有路径规划基线，仅用于比较。Improved A* 与 AA* 使用相同网络、1 s 时间步长、密度相关移动、设施服务容量、下游有限接纳、溢回和完全疏散终止条件。路径算法只决定客流批次的后续路线；路线被接受后，移动、服务和空间接纳由同一物理执行器推进。因此，网络层比较的主要差异来自路径决策和客流分配，而不是两套不同的速度或容量规则。本文不在正文展开 Improved A* 的原始代价公式，相关实现细节保留在代码和补充材料中。")

    add_heading(doc, "3.2 物理队列、承诺事件与规划意图", 2)
    add_p(doc, "对每个资源 r，物理队列 Qr(t) 表示已经选择该资源、尚未获得入口容量且仍在上游的整数人数。已接受承诺事件记录为 (ak,mk)，其中 ak 是已经确定路线的批次预计到达资源 r 的时刻，mk 是该批次人数。未被接受的规划意图不进入物理队列、不写入未来到达事件，也不占用空间容量。逻辑批次还保存 source_group、queue_enter_time 和 current_path；同一批次不能同时属于两个资源队列。")
    add_p(doc, "三类状态在执行过程中严格分离。物理状态由乘客当前所在节点、边或真实资源队列表示；已接受承诺由在途批次和未来到达事件表示；规划意图只存在于路径搜索阶段。已经进入在途执行队列的批次继续执行已承诺边，不在途中即时改路。仍处于明确选择状态的批次才允许重新规划；改路前从旧逻辑队列删除，选定新路径后立即加入新资源队列。搜索用已访问节点集合禁止循环，并通过有向边、出口可达性和有限时间检查避免反向或断裂路径。")

    add_heading(doc, "3.3 到达事件队列推进", 2)
    add_p(doc, "在决策时刻 t0，AA* 从当前物理队列和已经接受的未来到达事件开始。事件按到达时刻升序排列。事件之间，服务过程按照 μr 消化队列；事件发生时，将该批次的整数人数加入队列。若候选批次预计在 τ 时刻到达，则队列推进到 τ 后得到资源等待估计。")
    add_equation(doc, 2, "queue_minus")
    add_latex_note(doc, r"Q_k^-=\max\{Q_{k-1}^+-\mu_r(a_k-a_{k-1}),0\}")
    add_equation(doc, 3, "queue_plus")
    add_latex_note(doc, r"Q_k^+=Q_k^-+m_k")
    add_equation(doc, 4, "queue_target")
    add_latex_note(doc, r"\widehat Q_r(\tau)=\max\{Q_K^+-\mu_r(\tau-a_K),0\}")
    add_p(doc, "图 3b 给出一个数值例子：当前队列为 80 人，服务率 μ=2 人/s；20 s 时到达 30 人，35 s 时到达 20 人，候选批次在 50 s 到达。0–20 s 消化 40 人后剩余 40 人；加入 30 人后为 70 人；20–35 s 再消化 30 人后为 40 人；加入 20 人后为 60 人；35–50 s 消化 30 人，故候选批次到达时的预测队列为 30 人，而不是只读取 t0 时的 80 人或把所有未来客流一次性相加。")
    add_p(doc, "路线一旦被接受，候选批次的整数到达事件立即写回资源事件索引；同一轮内后续批次按确定性顺序处理，因此后续搜索可看到前一批次已经产生的未来竞争。这个写回过程只记录已接受路线，不将未接受的候选路线误计为物理需求。")

    add_heading(doc, "3.4 时间依赖路径代价", 2)
    add_p(doc, "沿候选路径从节点 i 推进至 j 时，后续资源的预计到达时刻包含此前所有移动和等待。对设施资源，代码还加入当前批次的平均服务时间项；普通边资源不加入该项。")
    add_equation(doc, 5, "batch")
    add_latex_note(doc, r"b_r(m)=m/(2\mu_r)\ \text{for facility resources; }0\ \text{otherwise}")
    add_equation(doc, 6, "arrival")
    add_latex_note(doc, r"\tau_j=\tau_i+t_{ij}^{move}+\widehat Q_r(\tau_i)/\mu_r+b_r(m)+s_j^{space}")
    add_p(doc, "其中，tmove 是共同物理模型给出的边移动时间；Q̂r/μr 是候选批次在预计到达时刻面对的资源等待；br(m) 是设施批次服务项；s_space 是下游空间有限接纳导致的等待。各项均以秒计，不在路径搜索中做无依据的跨指标归一化。密度暴露通过等待阶段和移动阶段的密度风险分别累计。")
    add_equation(doc, 7, "risk")
    add_latex_note(doc, r"R_j=R_i+R_{ij}^{wait}+R_{ij}^{move}")
    add_equation(doc, 8, "objective")
    add_latex_note(doc, r"C_j=(\tau_j-\tau_0)+\lambda R_j")
    add_p(doc, "λ 为代码中的安全权重，用于将密度风险时间纳入广义目标。本文把移动、资源等待、批次服务和空间接纳作为时间项，把密度暴露作为风险项；结果部分分别报告完成时间、静止暴露、移动距离和计算时间，以避免将不同物理量合并成单一“效率”。")
    add_figure(doc, "fig2_aa_mechanism_numeric_cn.png", "图 2  AA* 的到达事件和时间依赖代价。图中区分物理状态、已接受承诺和规划意图，给出真实队列、未来到达事件、服务消化以及候选批次到达时的数值推进示例，并将物理移动、资源等待、批次服务、空间接纳和密度暴露分解为逐段代价。")

    add_heading(doc, "3.5 多标签搜索、剪枝与批次级改路", 2)
    add_p(doc, "AA* 的搜索标签保存当前节点 v、累计预计时间 t、累计风险 R、广义目标 C、前驱标签和已访问节点集合。由于同一节点在不同到达时刻面对的未来队列不同，较晚但成本较低的标签不能仅因到达较晚而删除。代码只在到达时间相同的容差范围内，使用风险和目标值进行支配判断；若新标签在相同到达时刻同时不差于旧标签，则丢弃新标签，反之停用旧标签。")
    add_p(doc, "为避免标签数量因循环无限增长，搜索对每个标签保存 visited_nodes，禁止重复进入已经访问的节点；有限节点图上的简单路径数量有限。与此同时，代码使用自由流剩余时间下界、当前最优出口代价和目标截止值进行上界剪枝，并记录生成标签数、扩展标签数、单节点最大活动标签数、开放堆大小和搜索运行时间。该实现不设置未经实验支持的固定全局标签上限，因此论文不宣称一个与图规模无关的多项式最坏复杂度；实际计算代价由运行诊断报告。")
    add_p(doc, "批次级改路只作用于仍处于选择状态的批次。设保留原路线的剩余成本为 Cstay，候选路线成本为 Cswitch，实际收益率为 Rsw=(Cstay−Cswitch)/Cstay。仅当候选路线合法、下一跳确实改变、Cswitch<Cstay 且 Rsw≥gmin 时接受改路。高负荷 Mode 4 使用 gmin=0.20，作为当前场景的防摆动阈值，而不是外部标定的通用参数。代码同时检查同一步 A–B–A 反向循环和重复改路机会。")
    add_code_block(doc, "AA*(G, batch):\n  1. 读取批次当前节点、人数 m、时刻 t0 和资源事件索引。\n  2. 将起始标签 (v0, 0, 0, 0, ∅, {v0}) 放入优先队列。\n  3. 弹出最小估计目标标签；对每条可行边 (u,v)：\n       识别资源 r(u,v)，推进 Qr 到标签预计到达时刻；\n       计算资源等待、设施批次服务、物理移动和空间接纳等待；\n       更新到达时刻、密度风险和广义目标；\n       若循环、不可达、非有限移动时间或上界不可行，则丢弃标签；\n       若同到达时刻存在支配标签，则丢弃，否则保存新标签。\n  4. 从所有可达出口标签中选择最小广义目标路径。\n  5. 接受路径后，将整数人数及预计到达时刻写回资源事件索引。\n  6. 仅对未进入在途队列的批次执行收益阈值改路；已在途批次继续执行承诺路线。")
    add_figure(doc, "fig3_workflow_cn.png", "图 3  网络规划、共同物理执行层和 Pathfinder 连续空间评估流程。网络层负责生成路径和出口分配；共同执行层用于公平比较 Improved A* 与 AA*；Pathfinder 用于执行预分配路线并提供软件原生 Goto Any Exit 参照。")

    add_heading(doc, "4 实验设计与评价指标", 1)
    add_heading(doc, "4.1 网络层比较与评价指标", 2)
    add_p(doc, "两种负荷分别运行 Improved A* 和 AA*，只改变路径规划方法，保持站体、需求、速度、容量、空间接纳、溢回和终止条件一致。站级清空用平均完成时间和 Tp 表示，其中 Tp 是累计完成比例首次达到 p 的时刻；机制指标包括累计静止暴露、总移动距离、出口和关键设施负荷均衡以及墙钟运行时间。")
    add_equation(doc, 9, "tp")
    add_latex_note(doc, r"T_p=\min\{t:E_{out}(t)\ge pN\},\ p\in\{0.50,0.80,0.95,0.99,1.00\}")
    add_equation(doc, 10, "wstat")
    add_latex_note(doc, r"W_{stat}=\sum_t N_{stationary}(t)\Delta t")
    add_equation(doc, 11, "jain")
    add_latex_note(doc, r"J(x)=(\sum_i x_i)^2/[n\sum_i x_i^2]")
    add_p(doc, "网络层 Wstat 是共同执行层中站内未移动人员的累计人·秒，包含资源排队、空间阻塞和其他站内静止状态。Pathfinder 的 congestion time total(s) 是连续空间模型为每位乘客记录的拥堵时间，两者不作同一物理量处理。出口 Jain 指数采用所有最终出口的疏散人数，关键设施 Jain 指数采用代码定义的设施负荷；Jain 越高表示分配更均衡，但均衡本身不等同于更短疏散时间。")

    add_heading(doc, "4.2 单模块消融", 2)
    add_p(doc, "消融在 17,905 人高负荷下进行，以完整 AA* 为参照，每次只关闭一个规划组件：资源队列等待代价、到达时刻队列预测、空间接纳等待、密度暴露代价或多标签搜索。共同物理执行层保持开启。当前输出为每个确定性配置的一次完整运行，因此报告相对变化和机制一致性，不构造不存在的重复样本、显著性检验或置信区间。")

    add_heading(doc, "4.3 Pathfinder 跨模型一致性评估", 2)
    add_p(doc, "Pathfinder 复现使用相同的站体几何文件，网络模型输出的完整批次路线被转换为 Pathfinder 可执行的路径分配，分别形成 P-Improved 和 P-AA；同时运行 Pathfinder 原生 Goto Any Exit 作为自主选出口参照。评价使用所有乘客的经验累计完成曲线、平均完成时间、T95、T100、拥堵时间和移动距离，结果以场景级分布和汇总指标解释。若不同协议的初始位置、房间或朝向字段存在软件层差异，结论不延伸到单个乘客层面的因果比较。")
    add_equation(doc, 12, "ecdf")
    add_latex_note(doc, r"\hat F(t)=N^{-1}\sum_{i=1}^N\mathbf{1}(T_i\le t)")
    add_p(doc, "高、低负荷 Pathfinder CSV 均包含 id、name、exit time(s)、congestion time total(s) 和 distance (m) 字段，人数分别为 17,905 和 2,187。三种协议使用同一 SHA-256 几何文件。Pathfinder 不是网络层的真实值，而是检验路径分配方向能否在连续空间运动中保持的跨模型工具。")

    add_heading(doc, "5 结果", 1)
    add_heading(doc, "5.1 两种负荷下的网络层结果", 2)
    add_p(doc, "低负荷网络运行包含 2,187 人，两种方法均完成疏散，T100 均为 325 s。AA* 将平均完成时间从 138.924 s 降至 135.772 s，T95 从 289 s 降至 258 s（降低 10.7%），累计静止暴露从 108,228 降至 91,853 人·s（降低 15.1%）。与此同时，总移动距离从 213,389.2 m 增至 221,538.7 m。L2 和 L16 的线路清空时间分别提前 57 s 和 44 s，但 L7 仍为最后清空线路，因此等待暴露的降低没有改变最终 T100。")
    add_p(doc, "高负荷网络运行包含 17,905 人，两种方法均完成疏散。AA* 将平均完成时间从 438.850 s 降至 335.537 s，T95 从 1125 s 降至 889 s（降低 21.0%），T99 从 1309 s 降至 1097 s，T100 从 1486 s 降至 1222 s（降低 17.8%）。累计静止暴露从 6,552,858 降至 4,431,142 人·s（降低 32.4%），总移动距离从 1,481,120.5 m 增至 1,710,553.1 m（增加 15.5%）。出口 Jain 指数从 0.626 增至 0.711，关键设施 Jain 指数从 0.162 增至 0.409；墙钟运行时间从 37.7 s 增至 615.3 s，约为基线的 16.3 倍。")
    add_figure(doc, "fig4_network_load_stratified_cn.png", "图 4  两种负荷下的网络层结果。上排为 2,187 人低负荷，下排为 17,905 人高负荷；各面板分别报告平均完成时间、T95、T100、累计静止暴露、总移动距离、墙钟运行时间、出口 Jain 指数和关键设施 Jain 指数。低、高负荷采用分面柱状图，不表示连续负荷曲线。")
    add_table(doc, "表 2  两种负荷下的网络层主要结果", ["负荷", "方法", "平均 (s)", "T95 (s)", "T99 (s)", "T100 (s)", "静止暴露 (人·s)", "距离 (m)"], [
        ("低负荷", "Improved A*", "138.924", "289", "316", "325", "108,228", "213,389.2"),
        ("低负荷", "AA*", "135.772", "258", "315", "325", "91,853", "221,538.7"),
        ("高负荷", "Improved A*", "438.850", "1125", "1309", "1486", "6,552,858", "1,481,120.5"),
        ("高负荷", "AA*", "335.537", "889", "1097", "1222", "4,431,142", "1,710,553.1"),
    ], [4.0, 2.4, 2.2, 1.8, 1.8, 1.8, 3.3, 2.8], font_size=7.6)

    add_heading(doc, "5.2 共享瓶颈的时序机制", 2)
    add_p(doc, "原始 gate_backlog_step_trace.csv 提供了逐秒闸机路由队列、服务率、选入人数和已服务人数，因此可直接检验 AA* 的瓶颈协调，而不依赖抽象示意。高负荷下，AA* 的聚合闸机路由队列峰值和持续时间均低于 Improved A* 的集中拥堵过程；在低负荷下，两条聚合队列曲线整体较低，方法差异主要集中在短时峰值。高负荷差值热图显示，AA* 并不是使所有闸机的队列都降低，而是将客流从部分持续拥堵闸机重新分配到其他可用资源。")
    add_p(doc, "这种机制与线路清空结果一致：高负荷下 L2、L7、L16 和 L18 的清空时间分别由 1483、1486、874 和 1309 s 变为 993、1222、344 和 418 s，磁浮由 590 s 增至 651 s。AA* 不是让每条线路同时改善，而是以部分线路和部分出口的让渡换取全站尾部缩短。")
    add_figure(doc, "fig5_shared_bottleneck_timeseries_cn.png", "图 5  共享闸机瓶颈的真实时序证据。a、b 分别为高、低负荷下各闸机路由队列的聚合时间序列；c 为高负荷 AA*−Improved A* 的闸机队列差值热图，时间按 20 s 分箱，仅使用原始 gate_backlog_step_trace.csv；d 为高负荷各线路清空时间。来源—出口重分配结果在正文中结合出口 Jain 指数和线路清空时间解释，不另设主图。该图不把闸机节点占用人数和资源队列混为同一指标。")

    add_heading(doc, "5.3 来源—出口流量重分配", 2)
    add_p(doc, "高负荷来源—出口矩阵表明，AA* 的主要路径差异集中在少数大流量重分配。18 号线有 1,650 人从 Exit 17 转向 Exit 12；16 号线减少向 Exit 10 的分配并增加向 Exit 11 西侧的分配；磁浮客流在 18–21 号出口间重新分布；2 号线在多个本线出口之间进行数百人的调整。每条来源线路总人数保持不变，因此这些差值代表路径组合重构，而不是需求变化。")

    add_heading(doc, "5.4 单模块消融", 2)
    add_p(doc, "完整 AA* 的高负荷基准为 T95=889 s、T100=1222 s。去除资源队列等待代价后，T95 和 T100 分别增加 196.3% 和 182.7%，累计静止暴露增加 164.6%，关键设施 Jain 指数相对完整 AA* 降低 69.4%，是退化最大的变体。保留资源等待但关闭到达时刻队列预测后，T95 增至 1060 s、T100 增至 1427 s，累计静止暴露增加 45.2%，运行时间减少约 11.7%，说明未来到达事件预测是主要增益来源。")
    add_p(doc, "去除密度暴露项和改为单标签搜索均不改变 T95，但分别将 T100 增至 1254 s 和 1251 s，表明密度风险和多标签搜索的独立作用主要集中在残余客流尾部。关闭规划层空间接纳等待后，汇总结果与完整 AA* 相同；共同物理层的有限接纳和溢回仍然开启，因此该消融只检验规划层空间等待项，而不是关闭物理执行器中的接纳规则。")
    add_figure(doc, "fig6_ablation_cn.png", "图 6  AA* 的高负荷单模块消融。面板 a 显示各变体相对完整 AA* 的 T50、T95、T100、静止暴露、出口 Jain 和关键设施 Jain 变化；面板 b 显示 T100 与墙钟运行时间变化。每个变体只关闭一个规划组件，数值来自单次确定性运行。")
    add_table(doc, "表 3  AA* 高负荷单模块消融结果（相对完整 AA* 的变化，%）", ["变体", "T95", "T100", "平均完成时间", "静止暴露", "关键设施 Jain", "运行时间"], [
        ("去除资源队列等待", "+196.3", "+182.7", "+115.8", "+164.6", "−69.4", "+103.3"),
        ("去除到达时刻预测", "+19.2", "+16.8", "+29.0", "+45.2", "−59.8", "−11.7"),
        ("去除密度暴露", "0.0", "+2.6", "+0.1", "+0.1", "0.0", "+5.3"),
        ("单标签搜索", "0.0", "+2.4", "+0.0", "−0.0", "−4.2", "−9.3"),
        ("去除空间接纳等待", "0.0", "0.0", "0.0", "0.0", "0.0", "+3.5"),
    ], [4.6, 1.6, 1.6, 2.6, 2.4, 2.9, 2.1], font_size=7.7)

    add_heading(doc, "5.5 Pathfinder 高负荷连续空间评估", 2)
    add_p(doc, "高负荷 Pathfinder 运行中，P-AA 的平均完成时间为 396.0 s，低于 P-Improved 的 435.5 s；T95 从 1024.5 s 降至 985.1 s，T100 从 1414.6 s 降至 1298.3 s。P-AA 的平均拥堵时间为 276.6 s，低于 P-Improved 的 311.3 s，而平均移动距离为 131.7 m，与 P-Improved 的 131.8 m 基本相同。网络层 AA* 的改善方向在连续空间执行中得到复现。")
    add_p(doc, "Pathfinder Goto Any Exit 的平均完成时间为 362.5 s，低于两种预分配路径；其 T95=1034.1 s、T100=1458.5 s，均高于 P-AA，说明较低的典型乘客完成时间并不等于更短的全站尾部。其平均移动距离 114.6 m 也是三种协议中最短，但这一几何优势没有转化为 T100 优势。")
    add_table(doc, "表 4  Pathfinder 高负荷场景级结果", ["协议", "平均完成 (s)", "T95 (s)", "T100 (s)", "平均拥堵 (s/人)", "平均距离 (m)"], [
        ("P-Improved", "435.5", "1024.5", "1414.6", "311.3", "131.8"),
        ("P-AA", "396.0", "985.1", "1298.3", "276.6", "131.7"),
        ("Goto Any Exit", "362.5", "1034.1", "1458.5", "253.0", "114.6"),
    ], [4.4, 2.5, 2.1, 2.1, 2.8, 2.7], font_size=8.0)

    add_heading(doc, "5.6 Pathfinder 低负荷连续空间评估", 2)
    add_p(doc, "低负荷 Pathfinder 三种协议均使用 2,187 名乘客和同一几何文件。P-Improved 的平均完成时间、T95 和 T100 分别为 117.6、254.6 和 319.0 s；P-AA 分别为 120.8、254.6 和 319.0 s。P-AA 的平均拥堵时间为 9.2 s/人，略低于 P-Improved 的 9.4 s/人，但平均移动距离由 120.7 m 增至 124.7 m。由此可见，网络层低负荷下的等待暴露降低并未转化为 Pathfinder 中更短的平均完成时间，且 T95、T100 均未改善。")
    add_p(doc, "Goto Any Exit 的平均完成时间为 102.2 s，T95 为 224.8 s，均低于两种预分配路径，但 T100 为 320.0 s，略高于二者的 319.0 s；其平均移动距离为 97.3 m，平均拥堵时间为 14.8 s/人。该结果再次显示典型乘客效率、拥堵暴露和最终尾部清空不是同一个目标。")
    add_figure(doc, "fig7_pathfinder_load_stratified_cn.png", "图 7  Pathfinder 两种负荷下的连续空间评估。第一行低负荷，第二行高负荷；面板分别展示经验累计完成曲线、完成时间分位剖面、平均完成时间—T100 联合比较以及拥堵暴露尾部。低、高负荷使用独立的时间坐标范围，结果按场景级分布解释。")
    add_table(doc, "表 5  Pathfinder 低负荷场景级结果", ["协议", "平均完成 (s)", "T95 (s)", "T100 (s)", "平均拥堵 (s/人)", "平均距离 (m)"], [
        ("P-Improved", "117.6", "254.6", "319.0", "9.4", "120.7"),
        ("P-AA", "120.8", "254.6", "319.0", "9.2", "124.7"),
        ("Goto Any Exit", "102.2", "224.8", "320.0", "14.8", "97.3"),
    ], [4.4, 2.5, 2.1, 2.1, 2.8, 2.7], font_size=8.0)

    add_heading(doc, "6 讨论", 1)
    add_heading(doc, "6.1 高负荷下共享瓶颈竞争放大了 AA* 的收益", 2)
    add_p(doc, "高负荷下，列车到站客流使多个来源批次在更短时间内进入同一组竖向设施和闸机。只读取当前状态的路径代价无法充分表示已经接受的未来竞争，而 AA* 在每个候选批次的预计到达时刻推进队列，因此能够用一部分额外移动距离换取更少的资源等待。高负荷结果中总移动距离增加 15.5%，但累计静止暴露降低 32.4%，平均完成时间降低 23.5%，说明主要收益来自等待压缩而非几何路径缩短。")
    add_heading(doc, "6.2 低负荷下等待改善不一定改变最终尾部", 2)
    add_p(doc, "低负荷下，AA* 仍将 T95 降低 10.7%，累计静止暴露降低 15.1%，但两种方法的 T100 均为 325 s。L7 仍然是最后清空线路，表明最终尾部由一个持续时间更长的线路级过程控制，局部队列协调不足以改变最后完成事件。Pathfinder 结果进一步显示，P-AA 与 P-Improved 的 T95 和 T100 相同，P-AA 平均完成时间略高。这不是网络层结果错误，而是网络批次分配与连续空间运动规则之间存在尺度差异：网络层减少的资源等待可能被连续空间中的初始位置、局部绕行和微观冲突重新分配。")
    add_heading(doc, "6.3 平均效率、尾部清空、绕行和计算代价需要同时报告", 2)
    add_p(doc, "Goto Any Exit 在两种负荷下都给出较低的平均完成时间或 T95，但 T100 并未优于 P-AA。该参照说明“让大多数乘客尽早完成”和“压缩最后一批乘客的清空尾部”是不同目标。AA* 的额外计算代价也不能忽略：高负荷墙钟时间约为 Improved A* 的 16.3 倍。因而 AA* 更适合以列车到站、站台或批次级引导的方式实施，而不是要求每名乘客高频个体重规划。")
    add_heading(doc, "6.4 适用范围", 2)
    add_p(doc, "本文结果支持的工程含义是：当多批次客流同时竞争共享服务设施时，预测到达事件具有明确价值；当负荷较低或最终尾部由单一线路控制时，AA* 更可能减少等待暴露而非改变 T100。该结论来自龙阳路站的两种确定性需求场景，不将收益幅度外推为所有地铁站或所有灾害情景下的通用规律。")

    add_heading(doc, "7 结论", 1)
    add_p(doc, "本文将多线换乘站疏散路径规划表述为共享服务能力上的到达时刻协调问题，并提出与代码实现一致的 AA*。该方法把楼梯、自动扶梯和闸机作为设施资源，把普通移动边作为边资源；将当前物理队列、已接受未来到达事件和服务过程推进到候选批次的预计到达时刻；在时变多标签搜索中累计移动、资源等待、批次服务、空间接纳和密度暴露，并通过批次级改路和事件回写维持物理队列守恒。")
    add_p(doc, "两种负荷的网络结果表明，AA* 的收益受负荷条件调节。高负荷下，AA* 将平均完成时间从 438.9 s 降至 335.5 s，T95 从 1125 s 降至 889 s，T100 从 1486 s 降至 1222 s，累计静止暴露降低 32.4%，但总移动距离增加 15.5%，墙钟运行时间增至 615.3 s。低负荷下，AA* 将 T95 从 289 s 降至 258 s，静止暴露降低 15.1%，但 T100 保持 325 s。消融结果把资源等待识别为共享瓶颈识别的基础，把到达时刻预测识别为主要增益来源，并显示多标签和密度暴露主要影响残余尾部。")
    add_p(doc, "Pathfinder 跨模型评估显示，高负荷下 P-AA 的平均完成时间和 T100 改善方向得到连续空间复现；低负荷下 P-AA 与 P-Improved 的 T95、T100 相同，平均完成时间略高，仅平均拥堵时间略低。Goto Any Exit 在两种负荷下都使较大比例乘客更早完成，但没有同步改善最终清空尾部。AA* 因此不应被表述为普遍最优算法，而应被理解为一种面向共享瓶颈竞争、收益受负荷条件调节的批次级路径协调方法。")

    add_heading(doc, "数据可用性声明", 1)
    add_p(doc, "本研究使用项目维护的龙阳路站 CAD 几何、Pathfinder 模型、路径规划代码和仿真输出。站体工程数据和模型文件的公开范围需依据数据持有方许可确定；在许可范围内，可提供用于复核本文汇总结果的派生统计表、指标定义和运行配置。", indent=False, size=9.5)
    add_heading(doc, "生成式人工智能使用声明（投稿前按期刊要求处理）", 1)
    add_p(doc, "本文的研究设计、代码、Pathfinder 模型、仿真运行和数值结果由作者提供并核对；语言组织、图表排版和文档格式化阶段可使用生成式工具辅助，但所有方法描述、结果数字和结论均须由作者根据源代码和原始输出最终确认。", indent=False, size=9.5)

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Shen Y, Yang H, Ren G, Ran B. Model cascading overload failure and dynamic vulnerability analysis of facility network of metro station. Reliability Engineering & System Safety 242 (2024) 109711. https://doi.org/10.1016/j.ress.2023.109711.",
        "[2] Wen X, Si B, Xu M, Zhao F, Jiang R. A passenger flow spatial-temporal distribution model for a passenger transit hub considering node queuing. Transportation Research Part C 163 (2024) 104640. https://doi.org/10.1016/j.trc.2024.104640.",
        "[3] Guo K, Zhang L. Simulation-based passenger evacuation optimization in metro stations considering multi-objectives. Automation in Construction 133 (2022) 104010. https://doi.org/10.1016/j.autcon.2021.104010.",
        "[4] Yang X, Yang Y, Li Y, Yang X. Path planning for guided passengers during evacuation in subway station based on multi-objective optimization. Applied Mathematical Modelling 111 (2022) 777–801. https://doi.org/10.1016/j.apm.2022.07.024.",
        "[5] Yang X, Dai W, Li Y, Yang X. An efficient evacuation path optimization for passengers in subway stations under floods. Tunnelling and Underground Space Technology 143 (2024) 105473. https://doi.org/10.1016/j.tust.2023.105473.",
        "[6] Yang X, Wan J, Zhu H, Xie C-Z, Zhang B. Optimization of passenger evacuation path in flood scenarios considering companion behaviors. Simulation Modelling Practice and Theory 145 (2025) 103212. https://doi.org/10.1016/j.simpat.2025.103212.",
        "[7] Hua Y, Zhao J, Li H-T, Duan L. Shortest or locally quickest? A prediction-based approach for evacuation choice simulation between multiple staircases. Journal of Safety Science and Resilience 5 (2024) 281–294. https://doi.org/10.1016/j.jnlssr.2024.04.001.",
        "[8] Thunderhead Engineering. Pathfinder Technical Reference Manual: Path Planning. https://www.thunderheadeng.com/docs/2024-2/pathfinder/technical-reference-manual/.",
        "[9] Thunderhead Engineering. Pathfinder User Manual: Behaviors—Goto Any Exit. https://www.thunderheadeng.com/docs/2026-1/pathfinder/behaviors/.",
        "[10] 蒙盾, 胡志强, 张洪雨. 基于改进 A* 算法的多层邮轮疏散系统仿真. 系统仿真学报 34(6) (2022) 1375–1382. https://doi.org/10.16182/j.issn1004731x.joss.21-0075.",
        "[11] Yu L, Liu H, Fang Z, Ye R, Huang Z, You Y. A new approach on passenger flow assignment with multi-connected agents. Physica A 628 (2023) 129175. https://doi.org/10.1016/j.physa.2023.129175.",
    ]
    for ref in refs:
        p = add_p(doc, ref, indent=False, size=8.7, after=3)
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.65)
        p.paragraph_format.line_spacing = 1.12

    doc.core_properties.title = "面向多线换乘站共享瓶颈协调的到达时刻队列感知疏散路径规划"
    doc.core_properties.subject = "AA*；共享瓶颈；Pathfinder；低负荷与高负荷"
    doc.core_properties.keywords = "metro evacuation; arrival-time queue; shared bottleneck; Pathfinder; AA*"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
