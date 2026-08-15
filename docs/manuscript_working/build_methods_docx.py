from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
TEX_PATH = ROOT / "docs" / "manuscript_working" / "31_methods_experiments_cn_formal_v12_networkflow_aa_revision.tex"
DOCX_PATH = ROOT / "docs" / "manuscript_working" / "31_methods_experiments_cn_formal_v12_networkflow_aa_revision.docx"


def set_run_font(run, latin="Calibri", east_asia="宋体", size=11, bold=None, color=None):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    total = sum(widths)
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    for child in list(old_grid):
        old_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        old_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), "B7C4D2")


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run._r.addnext(fld)


def parse_braced(text, start):
    open_pos = text.find("{", start)
    if open_pos < 0:
        return "", len(text)
    depth = 0
    for idx in range(open_pos, len(text)):
        if text[idx] == "{":
            depth += 1
        elif text[idx] == "}":
            depth -= 1
            if depth == 0:
                return text[open_pos + 1:idx], idx + 1
    return text[open_pos + 1:], len(text)


def replace_command_arg(text, command, replacement):
    token = "\\" + command
    while token in text:
        start = text.find(token)
        arg, end = parse_braced(text, start + len(token))
        if end <= start:
            break
        text = text[:start] + replacement(arg) + text[end:]
    return text


def math_plain(raw):
    text = raw.strip()
    if "\\begin{cases}" in text:
        before = text.split("\\begin{cases}", 1)[0].strip()
        body = text.split("\\begin{cases}", 1)[1].split("\\end{cases}", 1)[0]
        rows = []
        for row in body.split("\\\\"):
            row = row.strip().rstrip(",")
            if row:
                row = row.replace("&", " ")
                rows.append(math_plain(row))
        return before + " " + "; ".join(rows)

    while "\\frac" in text:
        start = text.find("\\frac")
        first, first_end = parse_braced(text, start + len("\\frac"))
        second, second_end = parse_braced(text, first_end)
        if first_end <= start or second_end <= first_end:
            break
        text = text[:start] + "(" + first + ")/(" + second + ")" + text[second_end:]
    text = re.sub(r"\\sum_\{[^{}]*\}\^\{[^{}]*\}", "Σ", text)
    text = re.sub(r"\\int_\{[^{}]*\}\^\{[^{}]*\}", "∫", text)
    text = text.replace("\\bar N_v", "N̄_v")
    text = text.replace("\\widehat Q_r", "Q̂_r")
    text = replace_command_arg(text, "text", lambda a: a)
    text = replace_command_arg(text, "mathrm", lambda a: a)
    text = replace_command_arg(text, "mathbf", lambda a: a)
    text = text.replace("\\left\\lfloor", "⌊").replace("\\right\\rfloor", "⌋")
    replacements = {
        "\\rightarrow": "→",
        "\\leftrightarrow": "↔",
        "\\leftarrow": "←",
        "\\left": "",
        "\\right": "",
        "\\widehat": "",
        "\\overline": "",
        "\\bar": "",
        "\\mathcal": "",
        "\\mathrm": "",
        "\\displaystyle": "",
        "\\alpha": "α",
        "\\beta": "β",
        "\\gamma": "γ",
        "\\lambda": "λ",
        "\\Delta": "Δ",
        "\\delta": "δ",
        "\\mu": "μ",
        "\\xi": "ξ",
        "\\rho": "ρ",
        "\\eta": "η",
        "\\tau": "τ",
        "\\pi": "π",
        "\\psi": "ψ",
        "\\ell": "ℓ",
        "\\Omega": "Ω",
        "\\Phi": "Φ",
        "\\max": "max",
        "\\min": "min",
        "\\sum": "Σ",
        "\\int": "∫",
        "\\in": "∈",
        "\\leq": "≤",
        "\\geq": "≥",
        "\\neq": "≠",
        "\\cup": "∪",
        "\\varnothing": "∅",
        "\\emptyset": "∅",
        "\\infty": "∞",
        "\\rightarrow": "→",
        "\\mid": "|",
        "\\cdot": "·",
        "\\ldots": "...",
        "\\mathrm{d}": "d",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = text.replace("\\,", " ")
    text = text.replace("\\;", " ")
    text = text.replace("\\!", "")
    text = text.replace("\\,", " ")
    text = text.replace("\\", "")
    text = text.replace("{", "").replace("}", "")
    text = text.replace("^2", "²").replace("^3", "³")
    text = re.sub(r"\s+", " ", text).strip()
    return text


def clean_text(text):
    text = text.strip()
    text = text.replace("\\noindent", "").strip()
    text = text.replace("~", " ")
    text = text.replace("--", "—")
    text = re.sub(r"\\textbf\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\textit\{([^{}]*)\}", r"\1", text)
    text = re.sub(r"\\\((.*?)\\\)", lambda m: math_plain(m.group(1)), text)
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def add_plain_paragraph(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.3)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, bold=True)
        add_math_runs(p, text[len(bold_prefix):])
    else:
        add_math_runs(p, text)
    return p


def add_math_runs(paragraph, text, size=11, bold=None):
    def flush(buf):
        if not buf:
            return
        run = paragraph.add_run("".join(buf))
        set_run_font(run, latin="Cambria Math", east_asia="宋体", size=size, bold=bold)

    delimiters = set(" =+-*/(),.;:|[]{}<>≤≥≠∈∪∩→↔·")
    buf = []
    i = 0
    while i < len(text):
        marker = text[i]
        if marker not in ("^", "_"):
            buf.append(marker)
            i += 1
            continue

        flush(buf)
        buf = []
        i += 1
        if i >= len(text):
            buf.append(marker)
            break

        start = i
        if text[i] == "{":
            depth = 1
            i += 1
            start = i
            while i < len(text) and depth:
                if text[i] == "{":
                    depth += 1
                elif text[i] == "}":
                    depth -= 1
                    if depth == 0:
                        break
                i += 1
            token = text[start:i]
            if i < len(text) and text[i] == "}":
                i += 1
        else:
            while i < len(text) and text[i] not in delimiters:
                i += 1
            token = text[start:i] or marker

        run = paragraph.add_run(token)
        set_run_font(run, latin="Cambria Math", east_asia="宋体", size=size, bold=bold)
        if marker == "^":
            run.font.superscript = True
        else:
            run.font.subscript = True

    flush(buf)


def add_equation(doc, raw):
    p = doc.add_paragraph(style="Equation")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_together = True
    add_math_runs(p, math_plain(raw), size=11)
    return p


def extract_table(block):
    caption_match = re.search(r"\\caption\{(.*?)\}", block, re.S)
    caption = clean_text(caption_match.group(1)) if caption_match else ""
    tabular_match = re.search(r"\\begin\{tabular\}.*?\n(.*?)\\end\{tabular\}", block, re.S)
    rows = []
    if tabular_match:
        for raw in tabular_match.group(1).splitlines():
            line = raw.strip()
            if not line or line.startswith("\\toprule") or line.startswith("\\midrule") or line.startswith("\\bottomrule"):
                continue
            line = line.replace("\\small", "").strip()
            if "\\\\" in line:
                line = line.split("\\\\", 1)[0].strip()
            cells = [clean_text(cell) for cell in line.split("&")]
            if cells and any(cells):
                rows.append(cells)
    return caption, rows


def add_table_from_rows(doc, caption, rows):
    if not rows:
        return
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(caption)
        set_run_font(r, bold=True, size=10)

    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    if col_count == 3:
        widths = [2100, 3200, 4060]
    else:
        widths = [780, 1450, 620, 760, 760, 760, 1000, 1100, 2130]
        widths = widths[:col_count]
        widths[-1] += 9360 - sum(widths)
    set_table_geometry(table, widths)

    for row_idx, values in enumerate(rows):
        tr_pr = table.rows[row_idx]._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for col_idx in range(col_count):
            cell = table.cell(row_idx, col_idx)
            value = values[col_idx] if col_idx < len(values) else ""
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 or col_idx >= 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            add_math_runs(p, value, size=8.5 if col_count > 3 else 9, bold=row_idx == 0)
            if row_idx == 0:
                set_cell_shading(cell, "F4F6F9")
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in (
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    eq = doc.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    eq.font.name = "Cambria Math"
    eq._element.rPr.rFonts.set(qn("w:ascii"), "Cambria Math")
    eq._element.rPr.rFonts.set(qn("w:hAnsi"), "Cambria Math")
    eq._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    eq.font.size = Pt(11)
    eq.paragraph_format.space_before = Pt(4)
    eq.paragraph_format.space_after = Pt(8)

    for style_name in ("List Number", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    r = header.add_run("方法与实验部分")
    set_run_font(r, size=9, color="6B7280")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    r = footer.add_run("第 ")
    set_run_font(r, size=9, color="6B7280")
    add_page_field(footer)
    r = footer.add_run(" 页")
    set_run_font(r, size=9, color="6B7280")


def add_title(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("方法与实验部分")
    set_run_font(r, east_asia="黑体", size=20, bold=True, color="0B2545")
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(18)
    r2 = p2.add_run("大型多线换乘站整体应急疏散")
    set_run_font(r2, east_asia="宋体", size=10, color="6B7280")


def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.15
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)
    border = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "16")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "2E74B5")
    border.append(left)
    p_pr.append(border)
    r = p.add_run("说明：")
    set_run_font(r, bold=True, color="1F4D78")
    r2 = p.add_run(text)
    set_run_font(r2, size=10)


def build():
    tex = TEX_PATH.read_text(encoding="utf-8")
    lines = tex.splitlines()
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title(doc)

    in_document = False
    paragraph_lines = []
    list_mode = None
    i = 0

    def flush_paragraph():
        nonlocal paragraph_lines
        if paragraph_lines:
            text = clean_text(" ".join(x.strip() for x in paragraph_lines))
            if text:
                if text.startswith("说明："):
                    add_note(doc, text[len("说明："):].strip())
                else:
                    add_plain_paragraph(doc, text)
        paragraph_lines = []

    while i < len(lines):
        line = lines[i].strip()
        if line == r"\begin{document}":
            in_document = True
            i += 1
            continue
        if not in_document:
            i += 1
            continue
        if line in (r"\maketitle", r"\end{document}") or line.startswith(r"\setcounter"):
            i += 1
            continue
        if line == "":
            flush_paragraph()
            i += 1
            continue
        if line.startswith(r"\["):
            flush_paragraph()
            math_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != r"\]":
                math_lines.append(lines[i].strip())
                i += 1
            add_equation(doc, " ".join(math_lines))
            i += 1
            continue
        if line.startswith(r"\begin{table}"):
            flush_paragraph()
            block = [line]
            i += 1
            while i < len(lines):
                block.append(lines[i])
                if lines[i].strip() == r"\end{table}":
                    break
                i += 1
            caption, rows = extract_table("\n".join(block))
            add_table_from_rows(doc, caption, rows)
            i += 1
            continue
        if line in (r"\begin{enumerate}", r"\begin{itemize}"):
            flush_paragraph()
            list_mode = "List Number" if "enumerate" in line else "List Bullet"
            i += 1
            continue
        if line in (r"\end{enumerate}", r"\end{itemize}"):
            flush_paragraph()
            list_mode = None
            i += 1
            continue
        if line.startswith(r"\item"):
            flush_paragraph()
            text = clean_text(line[len(r"\item"):].strip())
            p = doc.add_paragraph(style=list_mode or "List Bullet")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(text)
            set_run_font(r)
            i += 1
            continue
        for command, level in ((r"\section*", 1), (r"\section", 1), (r"\subsection", 2), (r"\subsubsection", 3)):
            if line.startswith(command):
                flush_paragraph()
                title, _ = parse_braced(line, len(command))
                title = clean_text(title)
                p = doc.add_paragraph(style=f"Heading {level}")
                p.paragraph_format.keep_with_next = True
                r = p.add_run(title)
                set_run_font(r, east_asia="黑体", size={1: 16, 2: 13, 3: 12}[level], bold=True,
                             color={1: "2E74B5", 2: "2E74B5", 3: "1F4D78"}[level])
                i += 1
                break
        else:
            paragraph_lines.append(line)
            i += 1
            continue
        continue

    flush_paragraph()
    doc.core_properties.title = "方法与实验部分"
    doc.core_properties.subject = "大型多线换乘站整体应急疏散"
    doc.core_properties.author = ""
    doc.save(DOCX_PATH)
    print(DOCX_PATH)


if __name__ == "__main__":
    build()
