from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "manuscript_working"
FIG = OUT / "figures_cn"
DOCX = OUT / "35_tust_manuscript_working_draft_cn.docx"
MD = OUT / "35_tust_manuscript_working_draft_cn.md"
ABLATION_ROOT = ROOT / "outputs" / "ablation"
ABLATION_ORDER = [
    "Full AA*", "No arrival-time queue prediction", "No resource-queue waiting cost",
    "No spatial receiving wait", "Single-label search", "No density-risk penalty",
]
ABLATION_LABELS = {
    "Full AA*": "完整 AA*",
    "No arrival-time queue prediction": "无到达时刻队列预测",
    "No resource-queue waiting cost": "无资源队列等待代价",
    "No spatial receiving wait": "无空间接收等待",
    "Single-label search": "单标签搜索",
    "No density-risk penalty": "无密度风险惩罚",
}

BLUE = RGBColor(23, 105, 170)
DARK = RGBColor(35, 38, 41)
GREY = RGBColor(90, 96, 101)


def font(run, size=10.5, bold=None, color=DARK):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    run.font.color.rgb = color


def configure(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.2); sec.bottom_margin = Cm(2.1)
    sec.left_margin = Cm(2.25); sec.right_margin = Cm(2.25)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    for level, size in [(1, 15), (2, 12.5), (3, 11.2)]:
        s = doc.styles[f"Heading {level}"]
        s.font.name = "Microsoft YaHei"; s._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        s.font.size = Pt(size); s.font.bold = True; s.font.color.rgb = BLUE
        s.paragraph_format.space_before = Pt(10 if level == 1 else 7)
        s.paragraph_format.space_after = Pt(4)


def h(doc, text, level=1):
    return doc.add_heading(text, level=level)


def p(doc, text, indent=True):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.line_spacing = 1.35
    para.paragraph_format.space_after = Pt(5)
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)
    r = para.add_run(text); font(r)
    return para


def equation(doc, latex, number):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3); para.paragraph_format.space_after = Pt(3)
    r = para.add_run(f"{latex}    ({number})"); font(r, size=10.2)


def caption(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.keep_with_next = False
    r = para.add_run(text); font(r, size=8.7, color=GREY)


def figure(doc, filename, caption_text, width=6.55):
    path = FIG / filename
    if not path.exists():
        return
    para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    shape = para.add_run().add_picture(str(path), width=Inches(width))
    doc_pr = shape._inline.docPr
    doc_pr.set("title", filename)
    doc_pr.set("descr", caption_text)
    caption(doc, caption_text)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tc_pr.append(shd)


def table(doc, headers, rows, widths, caption_text):
    cap = doc.add_paragraph(); cap.paragraph_format.keep_with_next = True
    r = cap.add_run(caption_text); font(r, size=9, bold=True, color=BLUE)
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i, (head, width) in enumerate(zip(headers, widths)):
        t.columns[i].width = Inches(width)
        c = t.rows[0].cells[i]; c.text = head; shade(c, "DCE9F5")
        for rr in c.paragraphs[0].runs: font(rr, size=8.3, bold=True)
    tr_pr = t.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            for rr in cells[i].paragraphs[0].runs: font(rr, size=8.1)
    doc.add_paragraph()


def ablation_rows():
    candidates = sorted(
        ABLATION_ROOT.glob("mode4_*/ablation_results.csv"),
        key=lambda path: path.stat().st_mtime,
        reverse=True,
    )
    latest_by_variant = {}
    for path in candidates:
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            for row in csv.DictReader(handle):
                latest_by_variant.setdefault(row["variant"], row)
    return [latest_by_variant[name] for name in ABLATION_ORDER if name in latest_by_variant]


def build_docx():
    doc = Document(); configure(doc)
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("面向多线换乘站共享瓶颈协调的到达时刻队列感知疏散路径规划"); font(r, size=18, bold=True, color=BLUE)
    sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Tunnelling and Underground Space Technology 中文工作稿"); font(r, size=10, color=GREY)

    h(doc, "摘要", 1)
    p(doc, "大型地下换乘站的疏散效率不仅由路径长度决定，还取决于不同客流批次到达共享楼扶梯、通道、闸机区和出口的时序。决策时看似通畅的设施，在乘客实际到达前可能已接收其他已分配客流。针对这种状态错位，本文提出到达时刻队列感知 A*（arrival-time queue-aware A*，AA*）：将共享设施表示为具有统一服务率的资源，根据当前队列、已确认到达事件及到达前可完成的服务量预测预计等待，并在时变多标签搜索中累计行走、资源等待、空间接收等待与密度暴露代价。研究以龙阳路多线换乘站为案例，站体几何来自项目 CAD，设置 2,187 人基准需求和 17,905 人列车到达叠加需求；Improved A* 与 AA* 共用同一物理加载层，高负荷路径分配进一步在 Pathfinder 中复现。高负荷网络仿真中，AA* 将 T95 和 T100 分别降低 21.0% 和 17.8%，累计停滞暴露降低 32.4%，同时增加了行走距离与计算时间。Pathfinder 中，P-AA 相对 P-Improved 将平均完成时间降低 9.1%、T100 降低 8.2%。Pathfinder 原生局部最快 Any Exit 的平均完成时间最短，但 T100 最长，显示平均效率与尾部清空并非同一目标。单模块消融表明，显式资源等待代价构成基础机制，到达时刻队列预测提供主要增益；多标签与密度风险项主要影响最末端尾部。")
    p(doc, "关键词：地铁站疏散；动态路径规划；到达时刻队列；共享瓶颈；Pathfinder；换乘站", indent=False)

    h(doc, "1 引言", 1)
    p(doc, "多线换乘站在有限的地下空间内连接多个站台、站厅和换乘方向。全站疏散时，站台候车人员、列车到达人员、站厅人员和换乘人员并不沿相互独立的通道运动，而是依次共享楼梯、自动扶梯、通道、闸机区和出口。局部设施的服务能力不足会沿客流链向上游传播，并延迟远离可见瓶颈的客流批次。将站内功能空间与客流组织为有向设施网络，有助于刻画这种跨线路耦合；进一步把节点排队与动态加载结合，则能够描述容量受限设施中的时空积聚[1,2]。")
    p(doc, "既有疏散路径研究已考虑距离、当前密度、拥堵、风险和多目标优化，并广泛采用微观仿真检验引导方案[3-6]。然而，对共享瓶颈而言，真正影响乘客的是实际到达时的设施状态，而不是路径决策瞬间的状态。若干客流批次若在此前已被分配至同一设施，当前状态代价会低估未来等待，并可能持续把新需求导向已被未来到达量占用的瓶颈。")
    p(doc, "本文把这一问题表述为多源客流的到达时刻协调。对任一候选路径，AA* 根据当前排队、已确认在途到达与到达前可完成的服务量，预测客流批次到达每个共享设施时面对的队列，并将该等待代价嵌入时变多标签 A*。研究不预设某种方法在所有指标上全面占优，而检验到达时刻队列感知是否改善等待暴露、上尾清空与设施负荷分配，以及这种改善对应的绕行与计算代价。")
    p(doc, "本文的研究问题为：在无火灾、建模设施均可用的多线换乘站应急疏散中，到达时刻队列感知能否通过协调共享瓶颈，改善 T95-T100 尾部清空和累计等待暴露；这种改善由哪些规划机制驱动，并能否在独立的 Pathfinder 微观运动模型中复现？围绕该问题，论文设置三个子问题：第一，AA* 相对 Improved A* 是否改变站级与线路级清空过程；第二，改善是否以更长移动距离和更高计算成本为交换；第三，到达时刻队列预测、资源等待、空间接收等待、多标签搜索和密度风险惩罚分别发挥何种作用。")
    figure(doc, "fig1_study_design_cn.png", "图1 研究问题与证据链。网络仿真用于控制物理层并检验路径机制；Pathfinder 用于在不同运动模型中复现高负荷路径分配结果；消融实验用于分解 AA* 的规划机制。")

    h(doc, "2 文献基础与研究定位", 1)
    h(doc, "2.1 设施网络、动态加载与路径优化", 2)
    p(doc, "Shen 等[1]将地铁站设施和有向客流组织为层级网络，分析设施过载后的流量重分配；Wen 等[2]在综合客运枢纽中耦合动态路径选择、节点排队和客流加载。这类研究说明，换乘站的疏散分析应把站内设施作为相互依赖的服务链，而非仅以几何最短路表示。Yang 等[3]在 TUST 的研究中构建了从站体结构、疏散网络、节点时间预测、多目标优化到 Pathfinder 比较的完整流程；相关工作还将引导人员布设、路径分配、拥堵与风险目标纳入地铁站疏散优化[4-6]。本文沿用“真实站体—计算网络—微观复现”的完整证据结构，但研究对象限定为无火灾、设施可用条件下的共享瓶颈协调。")
    h(doc, "2.2 Pathfinder Any Exit 的研究角色", 2)
    p(doc, "最短路径与局部最快路径并不等价。Hua 等[7]通过全尺寸实验研究多个楼梯之间的最短与局部最快选择，并使用 Pathfinder 作为模型比较的一部分。Pathfinder 的路径规划以局部最快门选择为基础，Goto Any Exit 允许乘客选择任一可达出口[8,9]。因此，本文保留 PF-LQ 作为原生局部最快行为参照，用于揭示平均完成与尾部清空之间的取舍；它不是 AA* 的第三种实现，也不被当作真值。")

    h(doc, "3 方法", 1)
    h(doc, "3.1 站内设施—通道有向网络", 2)
    p(doc, "研究范围为无火灾、无烟气、所有建模疏散设施可用的全站应急疏散。公共交通空间表示为有向图 G=(V,E)。节点包括站台、候车分区、站厅与换乘连接、楼梯、自动扶梯、闸机排队区、出口前区和安全出口；有向边表示相邻空间之间的可行移动。由 CAD 派生的当前模型包含 572 个节点和 1,414 条有向边，其中包括 16 个出口、32 个楼梯和 63 个自动扶梯节点。")
    figure(doc, "fig_station_network_cn.png", "图2 CAD 派生的设施—通道网络。（a）站内主要物理设施节点及其连接；（b）模型规模与设施类型。图中坐标来自当前计算模型，不是通用地铁站示意图。")
    p(doc, "同一物理设施可能在网络中连接多条边。本文定义资源集合 R 与映射 ρ:E→R∪{∅}，所有映射到同一资源 r 的入流共同竞争服务率 μ_r，从而避免把一个闸机组、楼梯或自动扶梯的容量按连接边数重复计算。")
    equation(doc, r"\rho:E\rightarrow\mathcal{R}\cup\{\varnothing\}", 1)

    h(doc, "3.2 共同物理加载层", 2)
    p(doc, "Improved A* 与 AA* 共用 1 s 时间步长、密度相关行走速度、设施服务队列、下游有限接收能力和溢回规则。Improved A* 的路径代价遵循改进 A* 参考实现中的长度—行走时间结构，其参数 α=0.15、β=0.85 和启发项 γ=0.10 来自参考模型适配[10]；这些参数作为模型输入使用，不表述为龙阳路现场标定值。")
    equation(doc, r"v(k)=\begin{cases}1.427,&k\le 0.2\\\max(1.427-0.3549k,0),&0.2<k\le4.0\\0,&k>4.0\end{cases}", 2)
    equation(doc, r"c_e^{Imp}(t)=\alpha l_e+\beta t_e^{move}(t),\quad h(n)=\gamma d(n,\mathcal{X})", 3)

    h(doc, "3.3 到达时刻队列预测", 2)
    p(doc, "设 t 为当前决策时刻，Q_r(t) 为资源 r 的当前队列，A_r(t,τ) 为已经确认会在 τ 之前到达该资源的在途客流，μ_r 为资源服务率。AA* 将队列推进到候选批次的预计到达时刻：")
    equation(doc, r"\widehat{Q}_r(\tau)=\max\{0,Q_r(t)+A_r(t,\tau)-\mu_r(\tau-t)\}", 4)
    equation(doc, r"w_r(\tau)=\widehat{Q}_r(\tau)/\mu_r", 5)
    p(doc, "实际实现按已确认到达事件的时间顺序推进服务，避免把所有到达量一次性叠加。对路径 P，AA* 累计物理行走时间 t、资源队列等待 wq、空间接收等待 ws 和密度暴露风险 R(P)：")
    equation(doc, r"C(P)=\sum_{(i,j)\in P}[t_{ij}+w^q_{ij}+w^s_{ij}]+\lambda R(P)", 6)

    h(doc, "3.4 时变多标签搜索", 2)
    p(doc, "当下游队列随到达时刻变化时，同一节点上代价较低但到达较晚的标签不一定支配到达较早的标签。AA* 因而分别保存节点、预计到达时间、累计时间和密度风险，只在安全的同到达时刻条件下执行支配剪枝，并以自由流剩余时间作为下界。已接受的移动会登记为后续决策可见的到达事件，使路径搜索与共同物理执行器形成闭环。")
    figure(doc, "fig3_aa_method_cn.png", "图3 AA* 的到达时刻代价与模块结构。（a）沿候选路径推进预计到达时刻；（b）规划目标；（c）同一节点保留多个未来到达标签；（d）本文逐项检验的五个规划模块。")

    h(doc, "4 案例与实验设计", 1)
    h(doc, "4.1 需求场景", 2)
    p(doc, "2,187 人基准需求来自项目人数设计记录，并按线路分配至站台候车、站厅与换乘空间。本文把它称为基准需求场景，不以节假日、周末或工作日标签强化场景含义。列车到达叠加需求在同一基准分布上加入代码中冻结的各列车载客量：2号线 2×2,400 人，7号线 2×1,620 人，16号线 2×1,230 人，18号线 2×1,650 人，磁浮 2×959 人，总计 15,718 人；场景总人数为 17,905 人。")
    figure(doc, "fig2_demand_composition_cn.png", "图4 需求构成。（a）17,905 人场景的线路与人口来源；（b）基准需求和列车到达叠加需求。")
    table(doc, ["线路", "站内基准客流", "列车到达", "合计"], [
        ("2号线",1112,4800,5912),("7号线",500,3240,3740),("16号线",84,2460,2544),
        ("18号线",491,3300,3791),("磁浮",0,1918,1918),("合计",2187,15718,17905)
    ], [1.2,1.4,1.4,1.2], "表1 需求场景的可追溯人口构成（人）")

    h(doc, "4.2 网络仿真比较", 2)
    p(doc, "网络实验只改变路径方法。几何、人口、速度—密度关系、设施服务率、下游接收限制、溢回和终止标准保持相同。主指标为 T95、T100、累计停滞人·秒和人均完成时间；机制指标包括线路清空、出口与关键设施 Jain 均衡指数、总移动距离和墙钟运行时间。所有高负荷运行均要求 17,905 人完全疏散并满足人口守恒。")
    equation(doc, r"T_p=\min\{t:E^{out}(t)\ge pN\},\ p\in\{0.50,0.80,0.90,0.95,0.99,1.00\}", 7)
    equation(doc, r"W^{stat}=\sum_tN^{stationary}(t)\Delta t", 8)

    h(doc, "4.3 Pathfinder 外部复现", 2)
    p(doc, "高负荷 Pathfinder 实验在 2023.3.1206 Steering 模式下进行，每个协议均含 17,905 名乘客。P-Improved 和 P-AA 分别执行由网络模型导出的 Improved A* 与 AA* 路径/出口分配；PF-LQ 使用 Pathfinder 原生 Goto Any Exit。三组 .geom 文件经 SHA-256 核验完全一致。P-Improved 与 P-AA 的乘客姓名集合一致，支持按姓名进行描述性配对；PF-LQ 为独立生成的原生行为场景，只进行总体分布比较。")
    p(doc, "外部复现的完整流程为：固定同一站体几何与人口规模；将网络模型的已执行完整路线合并为 Pathfinder 可读的行为分配；分别运行 P-Improved、P-AA 与 PF-LQ；从全部乘客记录计算经验累计完成曲线、T50-T100、拥堵时间和移动距离；最后将 P-AA 与 P-Improved 做按姓名配对，并将 PF-LQ 解释为局部最快情境参照。该流程完整保留了路径分配来源、微观运动执行和结果统计三个层次。")

    h(doc, "4.4 单模块消融", 2)
    p(doc, "消融实验采用 leave-one-component-out 设计。完整 AA* 之外，每次只关闭一个规划机制：到达时刻资源队列预测、资源队列等待代价、空间接收等待、多标签搜索或密度风险惩罚。站体、人口和共同物理执行器均保持不变。Improved A* 作为外部算法参照，不参与 AA* 模块贡献计算。每个变体报告 T95、T100、累计停滞、人均完成时间、总距离、负荷均衡和运行时间。")

    h(doc, "5 结果", 1)
    h(doc, "5.1 高负荷网络仿真", 2)
    p(doc, "两种方法均完成 17,905 人疏散。与 Improved A* 相比，AA* 将 T95 从 1,125 s 降至 889 s，将 T100 从 1,486 s 降至 1,222 s；累计停滞由 6,552,858 降至 4,431,142 人·秒，人均完成时间由 438.9 降至 335.5 s。与此同时，总移动距离由 1,481,121 m 增至 1,710,553 m，墙钟时间由 37.65 s 增至 615.28 s。结果呈现清晰的“绕行换等待”机制：AA* 允许部分客流走更长路径，以减少共享瓶颈前的持续停滞和最终尾部。")
    table(doc,["指标","Improved A*","AA*","变化"],[
        ("T95 (s)",1125,889,"-21.0%"),("T100 (s)",1486,1222,"-17.8%"),
        ("累计停滞 (人·s)",6552858,4431142,"-32.4%"),("人均完成时间 (s)",438.85,335.54,"-23.5%"),
        ("总移动距离 (m)",1481121,1710553,"+15.5%"),("出口 Jain",0.626,0.711,"+13.6%"),
        ("关键设施 Jain",0.162,0.409,"+152.4%"),("运行时间 (s)",37.65,615.28,"16.3倍")
    ],[2.2,1.3,1.3,1.1],"表2 高负荷网络仿真结果")
    p(doc, "线路级结果显示，AA* 缩短了 2、7、16 和 18 号线的清空时间，其中 18 号线由 1,309 s 降至 418 s；磁浮由 590 s 增至 651 s。7号线在两种方法下均为最后清空线路。站级改善由共享能力在不同线路间重新分配产生，而不是要求每条线路同时改善。")
    figure(doc, "fig4_network_high_load_cn.png", "图5 高负荷网络结果。（a-c）站级 T95、T100 和人均完成时间；（d）线路清空时间；（e）出口与关键设施的 Jain 均衡指数。")
    p(doc, "来源线路—出口分配矩阵进一步显示，站级改善并非各出口负荷的等比例缩放，而是跨线路出口选择的重新组织。AA* 减少部分过度集中的流向，同时增加原本利用不足的出口与通道组合，从而使出口和关键设施 Jain 指数分别由 0.626、0.162 提高至 0.711、0.409。")
    figure(doc, "fig6_source_exit_redistribution_cn.png", "图6 来源线路—出口分配重构。（a-b）Improved A* 与 AA* 的来源线路—出口人数矩阵；（c）AA* 相对 Improved A* 的人数差。该图使用全部 17,905 人的网络仿真出口记录。")

    h(doc, "5.2 Pathfinder 高负荷外部复现", 2)
    p(doc, "P-AA 相对 P-Improved 将平均完成时间从 435.5 s 降至 396.0 s，将 T100 从 1,414.6 s 降至 1,298.3 s，平均拥堵时间从 311.3 s 降至 276.6 s。按姓名配对的 17,905 名乘客中，10,073 人（56.3%）在 P-AA 中更快，7,825 人更慢，7 人相同；平均和中位节省分别为 39.5 s 和 12.2 s。")
    p(doc, "PF-LQ 的平均完成时间、T50 和 T90 最短，但其曲线在上尾与 P-AA 交叉：PF-LQ 的 T95 和 T100 分别为 1,034.1 s 与 1,458.5 s，而 P-AA 为 985.1 s 与 1,298.3 s。PF-LQ 更有利于典型乘客的快速完成，P-AA 则在 T95-T100 残余尾部形成更早的清空。")
    table(doc,["协议","平均","T50","T95","T99","T100","拥堵","距离(m)"],[
        ("P-Improved","435.5","346.6","1024.5","1296.6","1414.6","311.3","131.8"),
        ("P-AA","396.0","315.1","985.1","1189.2","1298.3","276.6","131.7"),
        ("PF-LQ","362.5","292.8","1034.1","1344.1","1458.5","253.0","114.6")
    ],[1.05,.65,.65,.65,.65,.7,.7,.7],"表3 Pathfinder 高负荷完成、拥堵与距离结果（除距离外单位为 s）")
    figure(doc, "fig5_pathfinder_validation_cn.png", "图7 Pathfinder 高负荷外部复现。（a）全部 17,905 名乘客的累计完成曲线；（b）T50-T100；（c-d）平均与最大完成时间；（e）P-AA 与 P-Improved 的按姓名配对结果。")
    p(doc, "拥堵构成和移动距离揭示了三种协议的不同作用方式。P-AA 相对 P-Improved 同时降低水平区域和楼梯拥堵，而人均移动距离基本不变；PF-LQ 通过更短路径取得最低平均完成时间，却未同步缩短 T95-T100。因而三者不是单一指标上的线性排序，而是平均效率、尾部清空与路径距离之间的不同运行点。")
    figure(doc, "fig8_pathfinder_tradeoff_cn.png", "图8 Pathfinder 高负荷效率—尾部—距离取舍。（a）人均水平区域与楼梯拥堵时间；（b）人均移动距离；（c）平均完成时间与 T100 的联合位置；（d）P-AA、PF-LQ 相对 P-Improved 的指标变化。")

    h(doc, "5.3 单模块消融结果", 2)
    rows = ablation_rows()
    if len(rows) >= 2:
        full, no_pred = rows[0], rows[1]
        p(doc, f"同一批次重跑的完整 AA* 得到 T95={float(full['T95_s']):.0f} s、T100={float(full['T100_s']):.0f} s 和累计停滞 {float(full['stationary_person_s']):,.0f} 人·秒，与原正式高负荷结果一致。关闭到达时刻队列预测后，T95 增至 {float(no_pred['T95_s']):.0f} s，T100 增至 {float(no_pred['T100_s']):.0f} s，累计停滞增至 {float(no_pred['stationary_person_s']):,.0f} 人·秒；总移动距离则由 {float(full['total_movement_distance_m']):,.0f} m 降至 {float(no_pred['total_movement_distance_m']):,.0f} m。该变体走得更短但等得更久，直接支持“未来瓶颈预测通过适度绕行减少等待”的机制解释。")
        table(doc,["AA* 配置","T95(s)","T100(s)","累计停滞(人·s)","总距离(m)","运行时间(s)"],[(ABLATION_LABELS.get(r['variant'], r['variant']),f"{float(r['T95_s']):.0f}",f"{float(r['T100_s']):.0f}",f"{float(r['stationary_person_s']):,.0f}",f"{float(r['total_movement_distance_m']):,.0f}",f"{float(r['wall_clock_s']):.1f}") for r in rows],[2.15,.72,.72,1.25,1.05,.9],"表4 AA* 高负荷单模块消融")
        if len(rows) >= 4:
            no_wait, no_space = rows[2], rows[3]
            p(doc, f"关闭资源队列等待代价产生最强退化：T95={float(no_wait['T95_s']):.0f} s、T100={float(no_wait['T100_s']):.0f} s，累计停滞达到 {float(no_wait['stationary_person_s'])/1e6:.2f} 百万人·秒，说明显式资源服务等待是 AA* 的基础机制。相反，关闭规划层空间接收等待后，T95、T100、累计停滞与总距离均与完整 AA* 相同。该结果表明，有限接收和溢回仍是共同物理层的重要约束，但额外空间等待预测在当前高负荷场景中没有独立贡献。")
        if len(rows) >= 6:
            single_label, no_density = rows[4], rows[5]
            p(doc, f"单标签搜索与去除密度风险惩罚均未改变 T95，但分别将 T100 增至 {float(single_label['T100_s']):.0f} s 和 {float(no_density['T100_s']):.0f} s；人均完成时间、累计停滞和总距离与完整 AA* 接近。这说明两项机制的独立作用集中在最后少量残余客流，而非整体完成过程。综合六组结果，资源等待代价决定方法能否识别服务瓶颈，到达时刻预测进一步协调未来竞争，多标签与密度风险项用于收紧最末端尾部。")
        if (FIG / "fig7_ablation_cn.png").exists():
            figure(doc, "fig7_ablation_cn.png", "图9 AA* 高负荷单模块消融。各变体每次仅关闭一个规划机制；物理执行器、站体、人口和设施能力保持一致。")
    else:
        p(doc, "[待高负荷消融运行完成后自动插入。]")

    h(doc, "5.4 低负荷结果插入位置", 2)
    p(doc, "本节将在 2,187 人 Pathfinder 三协议运行冻结后，使用与高负荷完全一致的累计完成曲线、分位完成时间、拥堵时间与配对统计。正文只写实际结果，不预设低负荷与高负荷具有相同的效应方向或幅度。")

    h(doc, "6 讨论", 1)
    h(doc, "6.1 贡献在于到达时刻协调，而非几何最短化", 2)
    p(doc, "网络结果中 AA* 的移动距离增加而停滞显著下降；关闭到达时刻预测后，路径距离缩短但等待和尾部均恶化。这两组证据共同说明，核心作用不是寻找更短几何路径，而是识别未来服务能力已被在途客流占用的设施，并在客流真正到达前完成分流。")
    h(doc, "6.2 平均完成与尾部清空", 2)
    p(doc, "PF-LQ 说明平均完成时间与最终清空并非同一评价维度。局部最快行为可使多数乘客迅速完成，却可能留下少量长路径或持续拥堵的残余客流。AA* 的优势集中在 T95-T100，因而更适合被表述为共享瓶颈协调和尾部抑制方法，而不是所有乘客或所有指标上的全面最优方法。")
    h(doc, "6.3 运行代价与应用形式", 2)
    p(doc, "高负荷 AA* 的运行时间约为 Improved A* 的 16.3 倍，主要来自事件索引的未来队列查询、旧路径状态刷新和多标签搜索。路径结果更适合作为客流批次级引导策略，通过动态标志、现场人员和线路/站台分配实施，而不是要求每名乘客执行实时个体最优。")

    h(doc, "7 结论", 1)
    p(doc, "本文将复杂换乘站疏散表述为共享设施容量上的到达时刻协调问题，并提出将预计资源队列、空间接收等待和密度暴露嵌入时变多标签搜索的 AA*。17,905 人高负荷网络仿真表明，AA* 缩短 T95 和 T100、减少累计停滞并改善设施负荷分配，同时付出更长移动距离与更高计算成本；Pathfinder 外部复现保留了 P-AA 相对 P-Improved 的平均和尾部优势，而 PF-LQ 呈现平均更快、尾部更慢的另一运行点。消融进一步识别了分层机制：显式资源等待是基础，到达时刻预测提供主要性能增益，多标签与密度风险项收紧极端尾部，额外空间等待预测在本高负荷工况下没有独立贡献。低负荷结果完成后，将进一步检验共享设施竞争减弱时各机制的效应幅度。")

    h(doc, "参考文献", 1)
    refs = [
        "[1] Shen Y, Yang H, Ren G, Ran B. Model cascading overload failure and dynamic vulnerability analysis of facility network of metro station. Reliability Engineering & System Safety 242 (2024) 109711. https://doi.org/10.1016/j.ress.2023.109711.",
        "[2] Wen X, Si B, Xu M, Zhao F, Jiang R. A passenger flow spatial-temporal distribution model for a passenger transit hub considering node queuing. Transportation Research Part C 163 (2024) 104640. https://doi.org/10.1016/j.trc.2024.104640.",
        "[3] Yang X, Dai W, Li Y, Yang X. An efficient evacuation path optimization for passengers in subway stations under floods. Tunnelling and Underground Space Technology 144 (2024) 105473. https://doi.org/10.1016/j.tust.2023.105473.",
        "[4] Yang X, Yang Y, Li Y, Yang X. Path planning for guided passengers during evacuation in subway station based on multi-objective optimization. Applied Mathematical Modelling 111 (2022) 777-801. https://doi.org/10.1016/j.apm.2022.07.024.",
        "[5] Guo K, Zhang L. Simulation-based passenger evacuation optimization in metro stations considering multi-objectives. Automation in Construction 133 (2022) 104010. https://doi.org/10.1016/j.autcon.2021.104010.",
        "[6] Xu H, Wei Y, Tan Y. Optimization of emergency evacuation in complex rail transit station. Journal of Building Engineering 98 (2024) 110321. https://doi.org/10.1016/j.jobe.2024.110321.",
        "[7] Hua Y, Zhao J, Li H-T, Duan L. Shortest or locally quickest? A prediction-based approach for evacuation choice simulation between multiple staircases. Journal of Safety Science and Resilience 5 (2024) 281-294. https://doi.org/10.1016/j.jnlssr.2024.04.001.",
        "[8] Thunderhead Engineering. Pathfinder Technical Reference Manual: path planning and locally quickest door choice. 2021.",
        "[9] Thunderhead Engineering. Pathfinder User Manual: Behaviors and Goto Any Exit. 2022.",
        "[10] 蒙盾, 胡志强, 张洪雨. 基于改进A*算法的多层邮轮疏散系统仿真. 系统仿真学报 34(6) (2022) 1375-1382. https://doi.org/10.16182/j.issn1004731x.joss.21-0075.",
    ]
    for ref in refs:
        para = doc.add_paragraph(); para.paragraph_format.left_indent = Cm(.7); para.paragraph_format.first_line_indent = Cm(-.7)
        rr = para.add_run(ref); font(rr, size=8.8)
    doc.save(DOCX)
    return doc


def build_md():
    MD.write_text("""# 中文稿说明\n\n本文件的正式内容见同目录 DOCX。当前版本已按研究问题—方法—实验—结果—消融重排，并嵌入 9 张已完成主图。低/高负荷综合对比图在低负荷 Pathfinder 数据冻结后补入。\n\n数学公式的 LaTeX 源式已在 DOCX 正文中保留。\n""", encoding="utf-8")


if __name__ == "__main__":
    build_docx(); build_md(); print(DOCX)
