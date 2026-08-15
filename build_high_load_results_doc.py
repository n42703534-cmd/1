from __future__ import annotations

import csv
from collections import defaultdict
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "outputs" / "algorithm_compare" / "mode4_20260705_154656"
OUT = ROOT / "outputs" / "龙阳路高负荷结果对比表_20260705.docx"

# compact_reference_guide + named landscape_results_tables override
PAGE_W = 11.0
PAGE_H = 8.5
MARGIN = 0.55
CONTENT_DXA = 14256  # 9.9 in
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 90, "bottom": 90, "start": 120, "end": 120}

BLUE = "D9EAF7"
BLUE_DARK = "1F4E79"
BLUE_TEXT = RGBColor(31, 78, 121)
GREEN = "E2F0D9"
RED = "FCE4D6"
GRAY = "F2F2F2"
WHITE = "FFFFFF"
BLACK = RGBColor(0, 0, 0)
MUTED = RGBColor(90, 90, 90)


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def set_run_font(run, name="宋体", size=10.5, bold=False, color=BLACK):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in CELL_MARGINS.items():
        tag = "start" if edge == "start" else "end" if edge == "end" else edge
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="666666", size="5"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = borders.find(qn(f"w:{edge}"))
        if elem is None:
            elem = OxmlElement(f"w:{edge}")
            borders.append(elem)
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), size)
        elem.set(qn("w:space"), "0")
        elem.set(qn("w:color"), color)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: Iterable[int]):
    widths = list(widths)
    assert sum(widths) == CONTENT_DXA, (widths, sum(widths))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[min(idx, len(widths) - 1)]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table)


def format_cell(cell, text: str, *, bold=False, size=10.0, align=WD_ALIGN_PARAGRAPH.CENTER,
                fill=None, color=BLACK):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(str(text))
    set_run_font(run, size=size, bold=bold, color=color)
    if fill:
        set_cell_shading(cell, fill)


def add_heading(doc, text: str, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    # style supplies font; explicit East Asian font avoids substitution.
    set_run_font(r, name="微软雅黑", size=15 if level == 1 else 12, bold=True, color=BLUE_TEXT)
    return p


def add_note(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_run_font(r, size=8.5, color=MUTED)
    return p


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def keep_rows_together(rows):
    """Keep a short logical table/group on one page where Word can do so."""
    rows = list(rows)
    for i, row in enumerate(rows):
        set_row_cant_split(row)
        if i < len(rows) - 1:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True


def make_table(doc, headers: list[str], rows: list[list[str]], widths: list[int],
               change_col: int | None = None, change_signs: list[int] | None = None,
               font_size=9.6, keep_together=True):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for j, h in enumerate(headers):
        format_cell(table.rows[0].cells[j], h, bold=True, size=10.0, fill=BLUE)
    for i, values in enumerate(rows):
        cells = table.add_row().cells
        for j, value in enumerate(values):
            fill = None
            if change_col is not None and j == change_col and change_signs is not None:
                fill = GREEN if change_signs[i] > 0 else RED if change_signs[i] < 0 else GRAY
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 and len(value) > 22 else WD_ALIGN_PARAGRAPH.CENTER
            format_cell(cells[j], value, size=font_size, fill=fill, align=align)
    set_table_geometry(table, widths)
    if keep_together:
        keep_rows_together(table.rows)
    else:
        for row in table.rows:
            set_row_cant_split(row)
    return table


def pct_change(base: float, ours: float) -> str:
    if base == 0:
        return "—"
    d = ours - base
    return f"{d:+.0f} 人（{d/base:+.1%}）"


def friendly_source(group: str, people: int) -> str:
    if "_platform::Platform_" in group:
        node = group.split("::Platform_")[-1]
        return f"{node}（{people}人）"
    if group.endswith("_transfer"):
        core = group.removesuffix("_transfer")
        parts = core.split("_")
        if len(parts) >= 2:
            return f"{parts[0]}→{parts[1]}换乘（{people}人）"
    if group.endswith("_hall"):
        return f"{group.split('_')[0]}站厅（{people}人）"
    return f"{group}（{people}人）"


def build_document():
    summary = {r["method"]: r for r in read_csv(DATA_DIR / "summary_metrics.csv")}
    route_rows = read_csv(DATA_DIR / "route_chain.csv")
    pf_metrics = {r["metric"]: r for r in read_csv(ROOT / "outputs" / "pathfinder_high_load_quantification_20260704.csv")}

    # route_chain is based on destination-node arrivals and therefore counts
    # each gate traversal once. The historical facility_throughput.csv summed
    # both ends of every incident edge and doubled internal-node throughput.
    gate_arrivals: dict[str, dict[str, float]] = {}
    for row in route_rows:
        node = row["node"]
        if row["chain_type"] != "facility" or not node.startswith("Gate_"):
            continue
        method = row["method"]
        gate_arrivals.setdefault(node, {})
        gate_arrivals[node][method] = gate_arrivals[node].get(method, 0.0) + float(row["people"])

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Inches(PAGE_W)
    sec.page_height = Inches(PAGE_H)
    sec.top_margin = Inches(MARGIN)
    sec.bottom_margin = Inches(MARGIN)
    sec.left_margin = Inches(MARGIN)
    sec.right_margin = Inches(MARGIN)
    sec.header_distance = Inches(0.28)
    sec.footer_distance = Inches(0.28)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for level, size, before, after in ((1, 16, 14, 7), (2, 13, 10, 5), (3, 11, 8, 4)):
        st = styles[f"Heading {level}"]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = BLUE_TEXT
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    # Running header/footer (memo_masthead pattern, quiet implementation).
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hr = hp.add_run("龙阳路高负荷疏散优化 | 结果对比")
    set_run_font(hr, name="微软雅黑", size=8.5, bold=True, color=MUTED)
    fp = sec.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = fp.add_run("第 ")
    set_run_font(fr, size=8.5, color=MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    fp._p.append(fld)
    fr2 = fp.add_run(" 页")
    set_run_font(fr2, size=8.5, color=MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(2)
    title.paragraph_format.space_after = Pt(3)
    tr = title.add_run("龙阳路高负荷场景算法结果对比")
    set_run_font(tr, name="微软雅黑", size=20, bold=True, color=BLACK)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.paragraph_format.space_after = Pt(10)
    sr = sub.add_run("ImprovedAStar vs AdaptiveQueueAwareAStar（AA） | 17,905人")
    set_run_font(sr, name="微软雅黑", size=10.5, bold=True, color=MUTED)

    add_heading(doc, "1. Pathfinder总体疏散结果", 1)
    pf_defs = [
        ("T50", "T50", "s"), ("T80", "T80", "s"), ("T95", "T95", "s"),
        ("T99", "T99", "s"), ("T100", "T100", "s"),
        ("平均完成时间", "mean_completion_time", "s/人"),
        ("总低速拥堵时间", "total_congestion_time", "人·s"),
        ("楼梯低速拥堵时间", "stair_congestion_time", "人·s"),
        ("平均行走距离", "mean_distance", "m/人"),
    ]
    rows = []
    for label, key, unit in pf_defs:
        r = pf_metrics[key]
        aa = float(r["AA"])
        im = float(r["Improved_1"])
        improvement = float(r["AA_improvement_pct"])
        fmt = ",.3f" if max(abs(aa), abs(im)) < 10000 else ",.1f"
        rows.append([label, format(im, fmt), format(aa, fmt), f"{improvement:.3f}%", unit])
    make_table(doc, ["指标", "Improved（1）", "AA", "AA改善幅度", "单位"], rows,
               [3100, 2500, 2500, 2500, 3656], font_size=9.4)
    add_note(doc, "结论：AA的T100由1,586.200 s降至1,468.525 s，缩短117.675 s（7.419%）；总低速拥堵时间降低5.398%。")
    add_note(doc, "Pathfinder拥堵口径：过去10 s平均速度低于0.25 m/s即累计为低速拥堵时间；总值为全部17,905名乘客的累计人·秒。")

    add_heading(doc, "2. 介观模型关键疏散指标改善", 1)
    base = summary["ImprovedAStar"]
    aa = summary["AdaptiveQueueAwareAStar"]
    metric_defs = [
        ("T50", "T50", "s"), ("R_area", "R_area", "s/人"),
        ("累计排队时间", "queueing_time", "人·s"),
        ("中度拥挤暴露", "congestion_exposure", "人·s"),
        ("重度拥挤暴露", "severe_congestion", "人·s"),
        ("峰值溢出排队", "peak_overflow_queue", "人"),
        ("出口Gini", "exit_gini", "—"), ("运行时间", "wall_clock_s", "s"),
    ]
    rows = []
    for label, key, unit in metric_defs:
        b = float(base[key]); a = float(aa[key]); imp = (b - a) / b * 100 if b else 0
        if key in {"exit_gini"}:
            bv, av = f"{b:.6f}", f"{a:.6f}"
        elif max(abs(b), abs(a)) >= 10000:
            bv, av = f"{b:,.1f}", f"{a:,.1f}"
        else:
            bv, av = f"{b:,.2f}", f"{a:,.2f}"
        rows.append([label, bv, av, f"{imp:.2f}%", unit])
    make_table(doc, ["指标", "ImprovedAStar", "AA", "AA改善幅度", "单位"], rows,
               [3100, 2500, 2500, 2500, 3656], font_size=9.4)
    add_note(doc, "结论：AA的累计排队时间降低37.21%，重度拥挤暴露降低43.41%，出口Gini降低9.89%，体现出更强的削峰与负荷均衡能力。")

    add_heading(doc, "3. Gini降低9.89%的具体体现", 1)
    balance_rows = [
        ["全站出口Gini", "0.370277", "0.333660", "下降0.036617\n（相对降低9.89%）", "总体分配不均衡程度下降"],
        ["最大单出口占比", "12.91%", "10.76%", "下降2.15个百分点", "最高负荷出口占比降低"],
        ["L2四出口分配", "2,054 / 1,301 / 968 / 1,482人\n最大差：1,086人", "1,389 / 1,505 / 1,453 / 1,285人\n最大差：220人", "最大差缩小79.74%", "四个出口人数更接近"],
        ["L7两出口分配", "2,025 / 1,763人\n人数差：262人", "1,898 / 1,890人\n人数差：8人", "人数差缩小96.95%", "两个出口接近均分"],
        ["L18两主要出口分配", "2,312 / 1,439人\n人数差：873人", "1,840 / 1,926人\n人数差：86人", "人数差缩小90.15%", "主要出口接近均分"],
        ["L18四闸机分配", "1,979 / 333 / 190 / 1,249人\nGini：0.419", "1,125 / 715 / 861 / 1,065人\nGini：0.095", "闸机Gini降低77.27%", "局部闸机集中得到修正"],
    ]
    make_table(doc, ["佐证指标", "ImprovedAStar", "AA", "均衡改善", "对Gini的体现"], balance_rows,
               [2200, 3300, 3300, 2600, 2856], change_col=3,
               change_signs=[1, 1, 1, 1, 1, 1], font_size=8.7)
    add_note(doc, "注：Gini降低9.89%是相对降幅；计算对象为同一组全站出口（含零使用出口）。出口人数差缩小是Gini下降的直接表征。")

    # Gate arrival tables grouped by line.
    add_heading(doc, "4. 闸机单次通过人数及客流转移", 1)
    add_note(doc, "统计口径：按到达闸机节点的人数计数，每名乘客通过一个闸机只计1次。绿色表示人数增加，红色表示减少；增减本身不代表性能好坏，也不作为出口Gini改善的证据。")
    gate_groups = [
        ("L7闸机", "Gate_L7_"), ("L2闸机", "Gate_L2_"),
        ("L18闸机", "Gate_L18_"), ("L16闸机", "Gate_L16_"),
        ("磁浮线闸机", "Gate_Maglev_"),
    ]
    for title_text, prefix in gate_groups:
        add_heading(doc, title_text, 2)
        selected = sorted(name for name in gate_arrivals if name.startswith(prefix))
        rows = []
        signs = []
        for gate_name in selected:
            b = gate_arrivals[gate_name].get("ImprovedAStar", 0.0)
            a = gate_arrivals[gate_name].get("AdaptiveQueueAwareAStar", 0.0)
            d = a - b
            rows.append([gate_name, f"{b:,.0f}", f"{a:,.0f}", pct_change(b, a)])
            signs.append(1 if d > 0 else -1 if d < 0 else 0)
        make_table(doc, ["闸机", "ImprovedAStar", "AA", "变化"], rows,
                   [5000, 2800, 2800, 3656], change_col=3, change_signs=signs, font_size=9.5)

    add_heading(doc, "5. AA优势及边界的直接证据", 1)
    evidence_rows = [
        ["Pathfinder T100", "1,586.200 s", "1,468.525 s", "缩短7.419%", "微观模型复核"],
        ["Pathfinder总低速拥堵", "6,116,092.7人·s", "5,785,915.2人·s", "降低5.398%", "微观模型复核"],
        ["介观T50", "395.0 s", "328.5 s", "缩短16.84%", "前半程效率"],
        ["介观T80", "634.5 s", "582.5 s", "缩短8.20%", "中段疏散效率"],
        ["累计排队时间", "630,269.0人·s", "395,718.0人·s", "降低37.21%", "排队压力"],
        ["重度拥挤暴露", "131,549.2人·s", "74,445.2人·s", "降低43.41%", "高风险拥挤"],
        ["峰值溢出排队", "135.44人", "102.38人", "降低24.41%", "峰值削减"],
        ["全站出口Gini", "0.370277", "0.333660", "降低9.89%", "出口负荷均衡"],
        ["L18闸机Gini", "0.419", "0.095", "降低77.27%", "局部闸机均衡"],
        ["介观T95", "837.5 s", "894.0 s", "延长6.75%", "尾部边界"],
        ["介观T100", "1,424.5 s", "1,451.0 s", "延长1.86%", "尾部边界"],
    ]
    make_table(doc, ["证据指标", "ImprovedAStar", "AA", "变化", "证据性质"], evidence_rows,
               [2600, 2700, 2700, 2700, 3556], change_col=3,
               change_signs=[1, 1, 1, 1, 1, 1, 1, 1, 1, -1, -1], font_size=9.0,
               keep_together=False)
    add_note(doc, "结论边界：AA的优势集中在前中期疏散、排队削减、重度拥挤控制和负荷均衡；介观模型T95与T100仍分别延长6.75%和1.86%，不得表述为全时段均优。")

    add_heading(doc, "6. 汇报用结论", 1)
    conclusions = [
        "Pathfinder中，AA将T100缩短117.675 s（7.419%），且T50至T100的改善幅度逐步扩大。",
        "介观模型中，AA将T50和T80分别缩短16.84%和8.20%，R_area降低10.41%。",
        "AA将累计排队时间、重度拥挤暴露和峰值溢出排队分别降低37.21%、43.41%和24.41%。",
        "全站出口Gini由0.370277降至0.333660；L18闸机Gini由0.419降至0.095。",
        "边界：介观T95和T100分别延长6.75%和1.86%，因此结论应限定为前中期效率、拥挤控制与负荷均衡占优。",
    ]
    for text in conclusions:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(text)
        set_run_font(r, size=10.5)

    add_note(doc, "数据来源：outputs/algorithm_compare/mode4_20260705_154656 与 outputs/pathfinder_high_load_quantification_20260704.csv。")

    doc.core_properties.title = "龙阳路高负荷场景算法结果对比"
    doc.core_properties.subject = "ImprovedAStar与AdaptiveQueueAwareAStar结果表"
    doc.core_properties.author = ""
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
