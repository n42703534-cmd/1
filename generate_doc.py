"""Generate discussion document as Word file."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# ── Helper ──────────────────────────────────────────────
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def para(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=len(rows)+1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r+1].cells[c].text = str(val)
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('龙阳路地铁站疏散路径规划研究')
run.bold = True
run.font.size = Pt(18)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run('队列感知A*算法的设计、消融分析与Pathfinder交叉验证').font.size = Pt(13)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('与导师讨论文档  |  2026年6月24日').font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 1. 做了什么
# ══════════════════════════════════════════════════════════
heading('一、整体工作概述', level=1)

heading('1.1 研究框架', level=2)
para('本研究构建了一个双层验证框架：Python网络流模型用于路径规划算法的设计与评估，Pathfinder物理仿真用于验证算法产出的出口分配在实际行人动力学中的效果。')

para('研究对象为上海龙阳路地铁换乘站——5条线路（L2/L7/L16/L18/Maglev）交汇、地下三层结构、17个出口。使用网络流模型（Network Flow Model）对车站进行了拓扑建模。')

heading('1.2 基线算法：Improved A*（Meng et al., 2022）', level=2)

para('该算法被广泛引用于疏散路径规划领域（Journal of System Simulation, 2022），其核心公式为：')

para('  f(n) = g(n) + γ·h(n)')
para('  g_edge = α·l + β·(l/v)    —— 混合量纲：长度 + 时间')
para('  h(n) = γ·d_min              —— 最短距离启发式')
para('  密度 > 3.0 p/m² → 路径阻断（二元截断）')
para('  速度模型：Fruin密度-速度曲线  v(ρ) = v_free - k·ρ')

doc.add_paragraph()
para('该算法的局限性：', bold=True)
para('  (a) g(n)的混合量纲（α 和 β 需手动调参，物理意义不统一）')
para('  (b) 3.0 p/m²的二元截断过于粗糙，忽略3.0-4.0区间仍可缓慢通行')
para('  (c) 未考虑瓶颈节点（闸机/楼梯/扶梯）的排队等待时间')
para('  (d) h(n) = γ·d 仅基于距离，γ=0.10 可能导致低估剩余成本')

heading('1.3 提出算法：Adaptive Queue-Aware A-Star（AQA-AStar）', level=2)

para('在Improved A*基础上进行四项改进：', bold=True)

doc.add_paragraph()
para('改进一：量纲统一为纯时间', bold=True)
para('  原方法：g_edge = α·l + β·(l/v)，需要手动设定α=0.15、β=0.85两个权重')
para('  新方法：g_edge = t_travel + t_wait   （所有项均为秒）')
para('  t_travel = l / v_actual     v_actual = min( v(ρ) , v_cap )')
para('  优势：消除无物理意义的权重参数，所有成本项在统一的时间量纲下可加')

doc.add_paragraph()
para('改进二：瓶颈节点排队感知', bold=True)
para('  新引入：t_wait = N(t) / μ   （仅对闸机/楼梯/扶梯等瓶颈节点）')
para('  N(t)：当前时刻排队人数（含已预约入流）')
para('  μ：节点服务率（GB/T 33668标准通行能力），p/s')
para('  理论依据：Jin (2014) 点排队理论（Transportation Research Part B）')
para('  物理含义：若闸机前有N人排队，通过率为μ p/s，当前排队者需等待N/μ秒')
para('  优势：算法在选择路径时会自动规避排长队的瓶颈，引导人流分散到排队短的闸机')

doc.add_paragraph()
para('改进三：连续Fruin密度-速度模型', bold=True)
para('  原方法：密度 > 3.0 p/m² → 路径阻断（二元），密度 > 4.0 p/m² → 速度归零')
para('  新方法：取消二元截断，使用连续Fruin曲线，速度下限 0.1 m/s')
para('  优势：3.0-4.0 p/m²区间仍可缓慢通行，更符合实际物理行为')

doc.add_paragraph()
para('改进四：启发函数可采纳化', bold=True)
para('  原方法：h(n) = γ·d_min，γ=0.10，可能低估剩余成本（不可采纳）')
para('  新方法：h(n) = d_min / v_free，v_free = 1.427 m/s')
para('  优势：自由流速度下界保证 h(n) ≤ 实际剩余成本（可采纳性），保证A*最优性')

doc.add_paragraph()
para('改进五：边流量溢出重分配', bold=True)
para('  当A*偏好的边已达到容量上限时，将溢出的流量按比例重分配到其他未饱和的替代边上，防止"流量蒸发"人为减速疏散')

doc.add_paragraph()
para('改进六：每步边权只算一次', bold=True)
para('  使用步计数器守卫机制，每个仿真时间步中所有活跃节点的边权统一计算一次后复用，减少冗余计算')

heading('1.4 算法成本函数对比总结', level=2)

heading('1.5 消融实验设计', level=2)
para('为隔离各改进分量的独立贡献，设计了消融变体：')
para('  Full model (AQA-AStar)：t_wait = N/μ，包含所有改进')
para('  NoWaitingTime：t_wait = 0，仅保留密度-速度效应和改进246')
para('  t_wait的贡献 = Full - NoWaitingTime')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 2. 实验设置
# ══════════════════════════════════════════════════════════
heading('二、实验设置', level=1)

heading('2.1 场景设计', level=2)
headers = ['场景', '代号', '列车乘客', '站台+站厅+换乘', '总人数']
rows = [
    ['常规突发', 'Mode 1', '0', '2,187', '2,187'],
    ['双向满载', 'Mode 4', '15,718', '2,187', '17,905'],
]
add_table(headers, rows)

para('Mode 1模拟日常运营中突发紧急情况，乘客为站台候车者+站厅等候者+换乘通道中行人。')
para('Mode 4模拟早晚高峰双向列车同时到站的最不利工况，10列满载列车同时卸载（L7: 2×1620, L2: 2×2400, L16: 2×1230, L18: 2×1650, Maglev: 2×959），叠加站台/站厅/换乘人群。')

heading('2.2 Python仿真配置', level=2)
para('仿真步长 Δt = 0.5s')
para('速度模型：Fruin密度-速度曲线，v_free = 1.427 m/s, ρ_jam = 4.0 p/m², k = 0.3549')
para('密度阈值：中度 > 3.0 p/m², 严重 > 5.0 p/m²')
para('瓶颈容量：按GB/T 33668-2017标准计算（闸机、楼梯、扶梯）')

heading('2.3 Pathfinder验证配置', level=2)
para('软件版本：Pathfinder 2023.3.1206')
para('仿真模式：Steering（Agent-based）')
para('速度模型：Fruin speed profile bundle ("Average All", 正态分布均值~1.2 m/s)')
para('注入方法：Python算法输出每个source_group的出口分配 → 修改PTH文件中对应occupant的behavior字段 → Pathfinder用自己的Locally Quickest算法导航到指定出口')
para('控制粒度：出口级（Exit-level assignment），中间门/楼梯由Pathfinder内部动态选择')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 3. Python仿真结果
# ══════════════════════════════════════════════════════════
heading('三、Python仿真结果', level=1)

heading('3.1 Mode 1 — 常规突发（2,187人）', level=2)
headers = ['指标', 'Improved A*', 'AQA-AStar', '变化']
rows = [
    ['T50 (s)', '95.5', '94.5', '-1.0%'],
    ['T100 (s)', '321.0', '321.0', '0%'],
    ['排队时间 (万人·秒)', '3.9', '3.6', '-7.7%'],
    ['拥挤暴露 (万人·秒)', '0.97', '0.68', '-29.7%'],
    ['重度拥挤 (万人·秒)', '0', '0', '—'],
    ['出口Gini', '0.283', '0.249', '-12.0%'],
    ['墙钟时间 (s)', '7.99', '26.96', '—'],
]
add_table(headers, rows)
para('QA在非饱和态下全面优于基线：拥挤暴露-29.7%，出口更均衡（Gini-12%），总时间持平。排队感知机制在低负荷下有效引导分流。')

heading('3.2 Mode 4 — 双向满载（17,905人）', level=2)
headers = ['指标', 'Improved A*', 'AQA-AStar', '变化']
rows = [
    ['T50 (s)', '350.0', '284.0', '-18.9%'],
    ['T100 (s)', '791.5', '853.5', '+7.8%'],
    ['排队时间 (万人·秒)', '178.8', '118.5', '-33.7%'],
    ['拥挤暴露 (万人·秒)', '190.7', '136.6', '-28.4%'],
    ['重度拥挤 (万人·秒)', '113.1', '65.3', '-42.2%'],
    ['出口Gini', '0.317', '0.364', '+14.9%'],
    ['墙钟时间 (s)', '520.7', '393.8', '-24.4%'],
]
add_table(headers, rows)
para('核心发现：QA在拥挤指标上大幅改善（重度拥挤-42.2%），但T100恶化（+7.8%）且出口Gini恶化（+14.9%）。出现了拥挤安全与出口均衡之间的trade-off。')
para('Gini恶化原因：QA阻止了IA中的"跨线转移"——IA将L16/L18/Maglev的部分人流经由L2出口疏散（通过拥挤的换乘通道），QA将这些人留在本线出口——减少了通道拥挤但导致出口分配更集中。')

heading('3.3 消融实验 — Mode 4', level=2)
headers = ['指标', 'IA基线', 'Full QA', 'NoWaitingTime', 't_wait贡献']
rows = [
    ['T50 (s)', '350.5', '282.0', '301.0', '-19.0s'],
    ['T100 (s)', '791.5', '895.5', '884.0', '+11.5s'],
    ['排队 (万人·秒)', '180.3', '120.2', '167.5', '-47.3万 (-28.2%)'],
    ['拥挤暴露 (万人·秒)', '190.0', '136.1', '192.3', '-56.2万 (-29.2%)'],
    ['重度拥挤 (万人·秒)', '113.0', '63.6', '126.0', '-62.4万 (-49.5%)'],
]
add_table(headers, rows)
para('去除排队感知后（NoWaitingTime），拥挤和排队指标大幅回退。t_wait分量独立贡献了重度拥挤减少49.5%。这证明了排队感知（而非密度-速度效应）是算法效果的主要来源。')

heading('3.4 出口分配差异（两算法路径分歧）', level=2)
para('Mode 1中7处差异，集中在L7和L2。Mode 4中28处差异，覆盖全部5条线路。核心模式：QA把IA集中在窄口（低容量闸机）的流量拆散到宽口（高容量闸机），并阻止跨线转移。')
para('典型案例：')
para('  L18_L16_transfer：IA有41.5%人走Gate_L2_N_East→Exit_L2_6（跨线去L2出口），QA将其100%收回Exit_L16_11_east（本线），消除了跨境走廊的穿行拥挤')
para('  Maglev_train1：IA有43.4%人走Exit_L2_6（跨线），QA将Maglev自有出口比例从52%提升至71%，大幅减少走廊穿行')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 4. Pathfinder仿真结果
# ══════════════════════════════════════════════════════════
heading('四、Pathfinder仿真结果', level=1)

heading('4.1 Mode 1 — 常规突发（2,187人）', level=2)
headers = ['指标', 'IA', 'QA', '差异']
rows = [
    ['T_s 总疏散时间 (s)', '350.6', '350.6', '0%'],
    ['T_av 平均疏散时间 (s)', '139.0', '135.5', '-2.5%'],
    ['T_ac 平均拥堵时间 (s)', '31.1', '26.2', '-15.7%'],
    ['P50 (s)', '122.8', '126.0', '+2.6%'],
    ['P90 (s)', '257.1', '233.9', '-9.0%'],
    ['严重拥堵占比 (>60s)', '19.6%', '17.4%', '-2.2pp'],
    ['StdDev (s)', '80.7', '74.8', '-7.3%'],
]
add_table(headers, rows)
para('QA方案在Pathfinder中产生了明显改善：平均拥堵时间低15.7%，P90快23秒。尾部人群（P90-P100区间）受益最显著。')

heading('4.2 Mode 4 — 双向满载（17,905人）', level=2)
headers = ['指标', 'IA', 'QA', '差异']
rows = [
    ['T_s 总疏散时间 (s)', '1612.4', '1627.6', '+0.9%'],
    ['T_av 平均疏散时间 (s)', '458.1', '451.2', '-1.5%'],
    ['T_ac 平均拥堵时间 (s)', '331.4', '328.1', '-1.0%'],
    ['P50 (s)', '393.7', '387.3', '-1.6%'],
    ['P90 (s)', '934.7', '918.1', '-1.8%'],
    ['严重拥堵占比 (>60s)', '80.1%', '79.2%', '-0.9pp'],
]
add_table(headers, rows)
para('饱和态下T_s持平（0.9%差异，在Pathfinder随机性范围内），改善集中在P50/P90区间（1-2%）。改善幅度从Mode 1的15-20%缩水到Mode 4的1-2%，反映饱和态下瓶颈通行能力成为主导约束。')
para('Exit 6（Exit_L2_6）是最显著的变化：IA有1,038人通过，QA仅633人（-39%）。这是QA将L18_L16_transfer等跨境人流收回本线后的直接效果。')

heading('4.3 出口分布跨平台对比（Mode 4）', level=2)
headers = ['出口', 'Python-IA', 'Python-QA', 'PF-IA', 'PF-QA', '方向']
rows = [
    ['Exit_L2_2', '1880', '1959', '1851', '1970', 'QA↑ 一致'],
    ['Exit_L2_3', '1423', '996', '1359', '996', 'QA↓ 一致'],
    ['Exit_L2_4', '2335', '2545', '2422', '2548', 'QA↑ 一致'],
    ['Exit_L2_6', '1037', '624', '1038', '633', 'QA↓ 一致'],
    ['Exit_L7_7', '1801', '1711', '1808', '1726', 'QA↓ 一致'],
    ['Exit_L7_8/9', '1966', '2077', '1962', '2062', 'QA↑ 一致'],
    ['Exit_L16_10', '293', '326', '146', '149', 'QA↑ 一致'],
    ['Exit_L16_11_west', '1320', '1165', '1337', '1258', 'QA↓ 一致'],
    ['Exit_L16_11_east', '1154', '1366', '1242', '1385', 'QA↑ 一致'],
    ['Exit_L18_12', '2341', '1873', '2345', '1873', 'QA↓ 一致'],
    ['Exit_L18_17', '1372', '1893', '1368', '1893', 'QA↑ 一致'],
    ['Exit_Maglev_19', '441', '737', '441', '739', 'QA↑ 一致'],
    ['Exit_Maglev_20', '252', '331', '252', '332', 'QA↑ 一致'],
]
add_table(headers, rows)
para('15/15个出口的变化方向在两个平台间完全一致（含Maglev_18和Maglev_21共15个）。这是两平台唯一可以直接对比的指标，因为出口分配的定义在两个平台中完全相同。')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 5. 两套指标的不可比性
# ══════════════════════════════════════════════════════════
heading('五、Python与Pathfinder两套指标的不可比性', level=1)

heading('5.1 Python产出的指标', level=2)
headers = ['指标', '定义', '单位']
rows = [
    ['T50/T100', '50%/100%总人数到达出口的仿真时钟时间', '秒'],
    ['排队时间', '∑(每个排队节点上的人数 × Δt)，空间聚合', '人·秒'],
    ['拥挤暴露', '∑(密度>3.0 p/m²的节点上的人数 × Δt)', '人·秒'],
    ['重度拥挤', '∑(密度>5.0 p/m²的节点上的人数 × Δt)', '人·秒'],
    ['出口Gini', '出口分配人数的Gini系数', '—'],
]
add_table(headers, rows)
para('聚合方式：空间聚合（在节点上逐个求和）。密度来源：Fruin公式，基于节点人数÷节点有效面积。')

heading('5.2 Pathfinder产出的指标', level=2)
headers = ['指标', '定义', '单位']
rows = [
    ['T_s', '最后一人到达出口的仿真时间', '秒'],
    ['T_av', '所有个体到达时间的算术平均', '秒'],
    ['T_ac', '每个个体拥堵时间的算术平均（Pathfinder内部判定逻辑）', '秒'],
    ['P50/P90', '个体到达时间的分位数', '秒'],
    ['严重拥堵占比', '个体拥堵时间>60s的比例', '%'],
]
add_table(headers, rows)
para('聚合方式：个体聚合（先算每个agent的指标，再取均值/分位）。拥堵判定：Pathfinder内部逻辑，与Python的定义完全不同。')

heading('5.3 核心矛盾', level=2)
para('两套指标在多维度上不可对比：', bold=True)
para('  1. 聚合逻辑不同：Python在空间（节点）上求和，Pathfinder在个体上平均')
para('  2. "拥堵"的定义不同：Python基于节点密度阈值，Pathfinder基于个体速度阈值')
para('  3. 物理模型不同：网络流 vs Agent物理仿真，对"排队"的建模方式完全不同')
para('  4. 量纲不同：Python用"人·秒"，Pathfinder用"秒/人"')
para('因此，不能将两套指标放入同一张"一致性汇总表"。能跨平台对比的只有出口分布——因为"哪个人走哪个出口"在两个平台中定义完全一致。')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 6. 待讨论的问题
# ══════════════════════════════════════════════════════════
heading('六、待与导师讨论的核心问题', level=1)

para('问题1：论文中两套指标的定位', bold=True)
para('理解：Python指标用于"解释"算法机制（证明t_wait减少了瓶颈排队累积），Pathfinder指标用于"验证"算法效果（在更真实的物理仿真中出口分配是否同样减轻了拥堵）。两者在不同表中报告，在Discussion中做定性关联。这样是否正确？')

para('问题2：Pathfinder验证的粒度', bold=True)
para('目前只控制了出口分配（修改behavior字段），Pathfinder内部的中间门选择仍是Locally Quickest动态决定的。Pathfinder结果是否只能视为"保守估计"？在论文中是否需要说明这个限制？')

para('问题3：Mode 4总时间持平的归因', bold=True)
para('IA和QA在Mode 4中T_s相差0.9%。原因是：(a)饱和态下瓶颈容量成为主导约束，(b)Pathfinder自身的随机性。是否需要多次运行Pathfinder做统计检验？如果差异确实不显著，论文中怎么说？')

para('问题4：三层论证结构是否成立', bold=True)
para('拟定的论证线索：(1)算法层——Python消融证明t_wait减少了拥挤，(2)验证层——Pathfinder证明出口分配在物理仿真中也改善了拥堵，(3)发现层——讨论饱和态下边际效益递减的物理必然性。这个框架是否有说服力？')

para('问题5：Python高负荷时间不可信的问题', bold=True)
para('已查明原因：Python使用点排队模型（Jin 2014），缺少spillback建模（Guo et al. 2011），高负荷下系统性低估（~850s vs Pathfinder ~1600s）。Limitations中讨论是否够？')

para('问题6：目标期刊和框架可行性', bold=True)
para('拟投Tunnelling and Underground Space Technology (TUST)，框架见第七部分。是否合理？')

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 7. 论文框架
# ══════════════════════════════════════════════════════════
heading('七、拟定的论文框架（TUST投稿）', level=1)

para('1. Introduction', bold=True)
para('  地铁站疏散的重要性 → 现有路径规划算法的局限 → 本文三点贡献（排队感知算法、消融分析、跨平台验证）')

para('2. Methodology', bold=True)
para('  2.1 龙阳路站网络流模型（拓扑结构、节点类型、边容量计算）')
para('  2.2 Improved A*基线（Meng et al. 2022，混合量纲 + 二元截断）')
para('  2.3 Adaptive Queue-Aware A*（四项改进 + 成本函数对比表）')
para('  2.4 消融变体（NoWaitingTime）')

para('3. Experimental Setup', bold=True)
para('  3.1 场景（Mode 1 & Mode 4）')
para('  3.2 Python仿真配置（指标定义表）')
para('  3.3 Pathfinder验证配置（注入方法 + 指标定义表）')

para('4. Results', bold=True)
para('  4.1 Python仿真（Mode 1 + Mode 4综合指标表）')
para('  4.2 消融实验（Mode 4消融表）')
para('  4.3 出口分配差异（代表性source_group的出口变化）')
para('  4.4 Pathfinder验证 — Mode 1')
para('  4.5 Pathfinder验证 — Mode 4')
para('  4.6 出口分布跨平台对比（两平台出口人数方向一致性）')

para('5. Discussion', bold=True)
para('  5.1 排队感知在非饱和态的有效性（Mode 1两平台一致改善）')
para('  5.2 饱和态边际效益递减的物理解释')
para('  5.3 拥挤安全 vs 出口均衡的trade-off')
para('  5.4 两套指标的不可比性与各平台的角色定位')
para('  5.5 对地铁疏散管理的工程启示')
para('  5.6 Limitations（点排队局限、出口级验证粒度、单站案例）')

para('6. Conclusion', bold=True)

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 8. 附录
# ══════════════════════════════════════════════════════════
heading('八、附录：关键数据', level=1)

heading('A. Python指标定义', level=2)
headers = ['指标', '公式/定义']
rows = [
    ['Txx', 'evacuation_curve达到xx%总人数的仿真时间'],
    ['排队时间', '对每个queue节点（gate/stair/escalator/passageway≤8m）逐时间步累加 (人数×Δt)'],
    ['拥挤暴露', '节点密度>3.0 p/m²或capacity_ratio>1.0时，累加 (人数×Δt)'],
    ['重度拥挤', '节点密度>5.0 p/m²或capacity_ratio>2.0时，累加 (人数×Δt)'],
    ['出口Gini', '基于各出口到达人数的标准Gini系数'],
]
add_table(headers, rows)

heading('B. Pathfinder设置', level=2)
headers = ['参数', '值']
rows = [
    ['版本', '2023.3.1206'],
    ['模式', 'Steering'],
    ['速度Profile', 'Fruin "Average All"（正态分布，均值~1.2 m/s）'],
    ['注入方式', '修改PTH文件occupant的behavior字段（指向预设Goto Exit目标）'],
    ['控制粒度', '出口级（Exit-level），中间路径由Locally Quickest决定'],
]
add_table(headers, rows)

heading('C. 算法改进对照表', level=2)
headers = ['改进项', 'Improved A* (Meng 2022)', 'AQA-AStar (本文)']
rows = [
    ['成本量纲', '混合（α·l + β·l/v）', '统一纯秒（t_travel + t_wait）'],
    ['排队感知', '无', 't_wait = N/μ（瓶颈节点）'],
    ['密度处理', '二元截断（>3.0阻断）', '连续Fruin（min 0.1 m/s）'],
    ['启发函数', 'h = γ·d（γ=0.10）', 'h = d/v_free（可采纳）'],
    ['溢出处理', '无', '溢出重分配到未饱和边'],
    ['权重更新', '每节点单独计算', '每步统一计算一次'],
]
add_table(headers, rows)

doc.save(r'C:\Users\帅美婷sweet baby\Desktop\network\与导师讨论文档.docx')
print('Done: 与导师讨论文档.docx')
